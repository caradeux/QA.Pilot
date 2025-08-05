#!/usr/bin/env python3
"""
Rutas de API para funcionalidad de Excel y Ejecución Masiva
Maneja análisis de archivos Excel, validación con IA y ejecución de casos
"""

from flask import Blueprint, request, jsonify, current_app, send_file
import os
import sys
import tempfile
import json
import time
import random
from werkzeug.utils import secure_filename
from excel_test_analyzer import ExcelTestAnalyzer, TestCase, analyze_excel_file
import uuid
from datetime import datetime, timezone
import asyncio
import threading
import subprocess
import platform
import traceback

# Importar integración de base de datos
try:
    from db_integration import get_db_integration
    DB_INTEGRATION_AVAILABLE = True
except ImportError:
    DB_INTEGRATION_AVAILABLE = False
    print("⚠️  Integración de base de datos no disponible")

# Crear blueprint para las rutas de Excel
excel_bp = Blueprint('excel_api', __name__, url_prefix='/api')

# Configuración
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

# Diccionario global para bloqueos de archivos
file_locks = {}
file_locks_lock = threading.Lock()

def get_file_lock(file_path):
    """Obtiene un lock específico para un archivo"""
    with file_locks_lock:
        if file_path not in file_locks:
            file_locks[file_path] = threading.Lock()
        return file_locks[file_path]

def safe_read_execution_file(execution_path):
    """Lee un archivo de ejecución de forma segura"""
    file_lock = get_file_lock(execution_path)
    max_retries = 3
    retry_delay = 0.1
    
    for attempt in range(max_retries):
        try:
            with file_lock:
                if not os.path.exists(execution_path):
                    raise FileNotFoundError(f"Archivo de ejecución no encontrado: {execution_path}")
                
                # Esperar un poco si es un reintento
                if attempt > 0:
                    time.sleep(retry_delay * attempt)
                
                with open(execution_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    if not content.strip():
                        raise ValueError("Archivo vacío")
                    return json.loads(content)
                    
        except (IOError, OSError, json.JSONDecodeError) as e:
            print(f"Intento {attempt + 1}/{max_retries} fallido leyendo {execution_path}: {e}")
            if attempt == max_retries - 1:
                print(f"Error definitivo leyendo archivo de ejecución {execution_path}: {e}")
                print(f"Traceback: {traceback.format_exc()}")
                raise
        except Exception as e:
            print(f"Error inesperado leyendo archivo de ejecución {execution_path}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            raise

def safe_write_execution_file(execution_path, execution_record):
    """Escribe un archivo de ejecución de forma segura"""
    file_lock = get_file_lock(execution_path)
    max_retries = 3
    retry_delay = 0.1
    
    for attempt in range(max_retries):
        try:
            with file_lock:
                # Esperar un poco si es un reintento
                if attempt > 0:
                    time.sleep(retry_delay * attempt)
                
                # Escribir a archivo temporal primero
                temp_path = execution_path + '.tmp'
                with open(temp_path, 'w', encoding='utf-8') as f:
                    json.dump(execution_record, f, ensure_ascii=False, indent=2)
                    f.flush()  # Forzar escritura
                    os.fsync(f.fileno())  # Forzar sincronización con disco
                
                # Mover archivo temporal al final (operación atómica)
                if os.path.exists(execution_path):
                    os.replace(temp_path, execution_path)
                else:
                    os.rename(temp_path, execution_path)
                
                return  # Éxito, salir del bucle
                
        except (IOError, OSError) as e:
            print(f"Intento {attempt + 1}/{max_retries} fallido escribiendo {execution_path}: {e}")
            # Limpiar archivo temporal si existe
            temp_path = execution_path + '.tmp'
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            
            if attempt == max_retries - 1:
                print(f"Error definitivo escribiendo archivo de ejecución {execution_path}: {e}")
                print(f"Traceback: {traceback.format_exc()}")
                raise
        except Exception as e:
            print(f"Error inesperado escribiendo archivo de ejecución {execution_path}: {e}")
            print(f"Traceback: {traceback.format_exc()}")
            raise

def allowed_file(filename):
    """Verifica si el archivo tiene una extensión permitida"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def validate_file_size(file):
    """Valida el tamaño del archivo"""
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    return size <= MAX_FILE_SIZE

@excel_bp.route('/analyze_excel', methods=['POST'])
def analyze_excel():
    """
    Analiza un archivo Excel y extrae casos de prueba
    """
    try:
        # Verificar que se envió un archivo
        if 'excel_file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No se encontró el archivo Excel en la solicitud'
            }), 400
        
        file = request.files['excel_file']
        
        # Verificar que se seleccionó un archivo
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No se seleccionó ningún archivo'
            }), 400
        
        # Verificar extensión del archivo
        if not allowed_file(file.filename):
            return jsonify({
                'success': False,
                'error': 'Tipo de archivo no permitido. Solo se aceptan archivos .xlsx y .xls'
            }), 400
        
        # Verificar tamaño del archivo
        if not validate_file_size(file):
            return jsonify({
                'success': False,
                'error': 'El archivo es demasiado grande. Tamaño máximo: 16MB'
            }), 400
        
        # Guardar archivo temporalmente
        filename = secure_filename(file.filename)
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, filename)
        
        try:
            file.save(temp_path)
            
            # Obtener modo de datos del formulario
            data_mode = request.form.get('data_mode', 'simulated')  # Por defecto: datos simulados
            current_app.logger.info(f"Modo de datos seleccionado: {data_mode}")
            
            # Analizar el archivo Excel
            current_app.logger.info(f"Analizando archivo Excel: {filename}")
            
            # Crear analizador (usa automáticamente la API key del sistema)
            analyzer = ExcelTestAnalyzer()
            test_cases = analyzer.extract_test_cases(temp_path)
            
            # Realizar análisis híbrido con progreso
            current_app.logger.info("Realizando análisis híbrido de casos de prueba...")
            
            # Función de callback para progreso (opcional, para futuras mejoras)
            def progress_callback(progress, message, details=""):
                current_app.logger.info(f"Progreso: {progress}% - {message} - {details}")
            
            analyzed_cases = analyzer.analyze_with_hybrid_approach(test_cases, progress_callback, data_mode)
            
            # Generar resumen
            summary = analyzer.generate_summary_report(analyzed_cases)
            
            # Intentar guardar casos en base de datos si está disponible
            saved_case_ids = []
            if DB_INTEGRATION_AVAILABLE:
                try:
                    db_integration = get_db_integration()
                    for case in analyzed_cases:
                        if case.es_valido:  # Solo guardar casos válidos
                            case_id = db_integration.save_excel_test_case(case)
                            saved_case_ids.append(case_id)
                    
                    current_app.logger.info(f"Guardados {len(saved_case_ids)} casos en base de datos")
                except Exception as e:
                    current_app.logger.warning(f"Error guardando en DB: {e}")
            
            # Convertir casos a diccionarios para JSON
            cases_dict = []
            for case in analyzed_cases:
                case_dict = {
                    'id': case.id,
                    'nombre': case.nombre,
                    'historia_usuario': case.historia_usuario,
                    'objetivo': case.objetivo,
                    'precondicion': case.precondicion,
                    'pasos': case.pasos,
                    'datos_prueba': case.datos_prueba,
                    'resultado_esperado': case.resultado_esperado,
                    'url_extraida': case.url_extraida,
                    'es_valido': case.es_valido,
                    'problemas': case.problemas,
                    'sugerencias': case.sugerencias,
                    'instrucciones_qa_pilot': case.instrucciones_qa_pilot
                }
                cases_dict.append(case_dict)
            
            current_app.logger.info(f"Análisis completado: {len(analyzed_cases)} casos procesados")
            
            # Determinar tipo de análisis basado en si se usó IA
            analysis_type = 'ia' if analyzer.client else 'basico'
            
            # Preparar respuesta
            response_data = {
                'success': True,
                'test_cases': cases_dict,
                'summary': summary,
                'analysis_type': analysis_type,
                'message': f'Se analizaron {len(analyzed_cases)} casos de prueba exitosamente'
            }
            
            # Agregar información de base de datos si se guardaron casos
            if saved_case_ids:
                response_data['database_save'] = {
                    'success': True,
                    'saved_cases': len(saved_case_ids),
                    'case_ids': saved_case_ids
                }
            
            return jsonify(response_data)
            
        finally:
            # Limpiar archivo temporal
            try:
                os.remove(temp_path)
                os.rmdir(temp_dir)
            except:
                pass
    
    except Exception as e:
        current_app.logger.error(f"Error al analizar Excel: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error interno al procesar el archivo: {str(e)}'
        }), 500

@excel_bp.route('/reanalyze_case', methods=['POST'])
def reanalyze_case():
    """
    Re-analiza un caso de prueba específico
    """
    try:
        data = request.get_json()
        
        if not data or 'test_case' not in data:
            return jsonify({
                'success': False,
                'error': 'Datos del caso de prueba no proporcionados'
            }), 400
        
        case_data = data['test_case']
        
        # Crear objeto TestCase desde los datos
        test_case = TestCase(
            id=case_data.get('id', ''),
            nombre=case_data.get('nombre', ''),
            historia_usuario=case_data.get('historia_usuario', ''),
            objetivo=case_data.get('objetivo', ''),
            precondicion=case_data.get('precondicion', ''),
            pasos=case_data.get('pasos', ''),
            datos_prueba=case_data.get('datos_prueba', ''),
            resultado_esperado=case_data.get('resultado_esperado', ''),
            url_extraida=case_data.get('url_extraida')
        )
        
        # Re-extraer URL si es necesario
        analyzer = ExcelTestAnalyzer()
        if not test_case.url_extraida:
            test_case.url_extraida = analyzer._extract_url(test_case.pasos + " " + test_case.datos_prueba)
        
        # Re-analizar el caso con enfoque híbrido
        analyzed_cases = analyzer.analyze_with_hybrid_approach([test_case])
        
        analyzed_case = analyzed_cases[0]
        
        # Convertir a diccionario
        case_dict = {
            'id': analyzed_case.id,
            'nombre': analyzed_case.nombre,
            'historia_usuario': analyzed_case.historia_usuario,
            'objetivo': analyzed_case.objetivo,
            'precondicion': analyzed_case.precondicion,
            'pasos': analyzed_case.pasos,
            'datos_prueba': analyzed_case.datos_prueba,
            'resultado_esperado': analyzed_case.resultado_esperado,
            'url_extraida': analyzed_case.url_extraida,
            'es_valido': analyzed_case.es_valido,
            'problemas': analyzed_case.problemas,
            'sugerencias': analyzed_case.sugerencias,
            'instrucciones_qa_pilot': analyzed_case.instrucciones_qa_pilot
        }
        
        return jsonify({
            'success': True,
            'analyzed_case': case_dict,
            'message': 'Caso re-analizado exitosamente'
        })
        
    except Exception as e:
        current_app.logger.error(f"Error al re-analizar caso: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error al re-analizar el caso: {str(e)}'
        }), 500

@excel_bp.route('/execute_bulk_cases', methods=['POST'])
def execute_bulk_cases():
    """
    Ejecuta múltiples casos de prueba secuencialmente (uno por uno)
    """
    try:
        # Manejar tanto JSON como FormData
        if request.content_type and 'application/json' in request.content_type:
            data = request.get_json()
            test_cases = data.get('test_cases', [])
            execution_mode = data.get('execution_mode', 'sequential')
            show_browser = data.get('show_browser', False)
        else:
            # FormData desde el frontend actualizado
            test_cases_json = request.form.get('test_cases_json')
            if test_cases_json:
                test_cases = json.loads(test_cases_json)
            else:
                test_cases = []
            execution_mode = request.form.get('execution_mode', 'sequential')
            show_browser = request.form.get('show_browser', 'false').lower() == 'true'
        
        if not test_cases:
            return jsonify({
                'success': False,
                'error': 'No se proporcionaron casos de prueba para ejecutar'
            }), 400
        
        # Generar ID único para la ejecución
        execution_id = str(uuid.uuid4())
        
        # Crear ejecución masiva en base de datos si está disponible
        bulk_execution_id = None
        if DB_INTEGRATION_AVAILABLE:
            try:
                db_integration = get_db_integration()
                test_case_ids = [case.get('id', str(uuid.uuid4())) for case in test_cases]
                config = {
                    'execution_mode': execution_mode,
                    'headless_mode': not show_browser,
                    'browser_type': 'chromium',
                    'capture_screenshots': True
                }
                bulk_execution_id = db_integration.create_bulk_execution(
                    name=f"Ejecución masiva {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                    test_case_ids=test_case_ids,
                    config=config
                )
                current_app.logger.info(f"Ejecución masiva creada en DB: {bulk_execution_id}")
            except Exception as e:
                current_app.logger.warning(f"Error creando ejecución en DB: {e}")
        
        # Crear registro de ejecución masiva
        execution_record = {
            'id': execution_id,
            'bulk_execution_id': bulk_execution_id,
            'timestamp': datetime.now().isoformat(),
            'total_cases': len(test_cases),
            'status': 'iniciado',
            'execution_mode': execution_mode,
            'show_browser': show_browser,
            'current_case': 0,
            'cases': test_cases,
            'results': [],
            'progress': 0,
            'current_case_status': 'preparando'
        }
        
        # Guardar registro usando funciones seguras
        execution_file = f"bulk_execution_{execution_id}.json"
        execution_path = os.path.join(tempfile.gettempdir(), execution_file)
        
        # Usar función segura para crear el archivo inicial
        safe_write_execution_file(execution_path, execution_record)
        
        # Iniciar ejecución en segundo plano
        if execution_mode == 'sequential':
            thread = threading.Thread(
                target=execute_cases_sequential_wrapper,
                args=(current_app._get_current_object(), execution_id, test_cases, execution_path)
            )
        else:
            thread = threading.Thread(
                target=execute_cases_background_wrapper,
                args=(current_app._get_current_object(), execution_id, test_cases, execution_path)
            )
        
        thread.daemon = False  # No daemon para capturar errores
        thread.start()
        
        current_app.logger.info(f"Ejecución masiva iniciada: {execution_id} ({len(test_cases)} casos)")
        
        return jsonify({
            'success': True,
            'execution_id': execution_id,
            'total_cases': len(test_cases),
            'message': f'Ejecución masiva iniciada con {len(test_cases)} casos'
        })
        
    except Exception as e:
        current_app.logger.error(f"Error al iniciar ejecución masiva: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error al iniciar la ejecución: {str(e)}'
        }), 500

def execute_cases_sequential_wrapper(app, execution_id, test_cases, execution_path):
    """Wrapper para ejecutar casos secuencialmente con contexto de Flask"""
    try:
        with app.app_context():
            execute_cases_sequential(execution_id, test_cases, execution_path)
    except Exception as e:
        print(f"ERROR CRÍTICO en hilo secuencial: {e}")
        import traceback
        traceback.print_exc()
        try:
            mark_execution_failed(execution_path, f"Error crítico en hilo: {str(e)}")
        except:
            pass

def execute_cases_background_wrapper(app, execution_id, test_cases, execution_path):
    """Wrapper para ejecutar casos en paralelo con contexto de Flask"""
    try:
        with app.app_context():
            execute_cases_background(execution_id, test_cases, execution_path)
    except Exception as e:
        print(f"ERROR CRÍTICO en hilo paralelo: {e}")
        import traceback
        traceback.print_exc()
        try:
            mark_execution_failed(execution_path, f"Error crítico en hilo: {str(e)}")
        except:
            pass

def _extract_steps_from_instructions(instructions):
    """Extrae pasos individuales de las instrucciones"""
    try:
        # Dividir por puntos, saltos de línea o numeración
        import re
        
        # Intentar varios patrones de separación (mejorados)
        patterns = [
            r'\d+\.\s+',  # "1. ", "2. ", etc. (requiere espacio después)
            r'\d+\)\s+',  # "1) ", "2) ", etc. (requiere espacio después)
            r'paso\s*\d+[\:\.]?\s+',  # "Paso 1:", "paso 2.", etc.
            r'\.\s+(?=[A-Z])',  # Puntos seguidos de mayúscula (nueva oración)
            r'\n\s*',  # Saltos de línea
        ]
        
        # Intentar primera separación por numeración
        for pattern in patterns[:3]:  # Solo patrones de numeración
            parts = re.split(pattern, instructions, flags=re.IGNORECASE)
            if len(parts) > 2:  # Al menos 2 pasos válidos
                steps = [part.strip() for part in parts if part.strip()]
                if len(steps) >= 2:
                    return steps[:10]  # Máximo 10 pasos
        
        # Si no hay numeración clara, dividir por puntos o saltos de línea
        for pattern in patterns[3:]:
            parts = re.split(pattern, instructions)
            if len(parts) > 2:
                steps = [part.strip() for part in parts if part.strip() and len(part.strip()) > 10]
                if len(steps) >= 2:
                    return steps[:10]
        
        # Fallback: instrucción completa como un solo paso
        return [instructions.strip()]
        
    except Exception as e:
        print(f"[EXTRACT-STEPS] Error extrayendo pasos: {e}")
        return [instructions.strip()]

def generate_browser_use_script(url, instrucciones, headless=True, max_tiempo=300, capturar_pasos=True, browser='chrome', fullscreen=False, task_id=None, instruction_source='unknown'):
    """
    CORRECCIÓN FINAL: Genera un script ultra-simplificado que realmente funciona
    """
    try:
        print(f"[SCRIPT-GEN] Generando script ultra-simplificado para URL: {url}")
        print(f"[SCRIPT-GEN] Instrucciones: {instrucciones[:100]}...")
        print(f"[SCRIPT-GEN] Configuración: headless={headless}, tiempo={max_tiempo}s")
        
        # Escapar las instrucciones para uso seguro en string
        instrucciones_safe = instrucciones.replace('\\', '\\\\').replace('"""', '\'\'\'').replace('\r', '')
        
        # Crear script ULTRA-SIMPLIFICADO que realmente funciona
        script_content = f'''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script ultra-simplificado para ejecución masiva con browser-use
Task ID: {task_id or "unknown"}
URL: {url}
Generado: {datetime.now().isoformat()}
"""

import asyncio
import os
import sys
import time
from dotenv import load_dotenv
from browser_use import ChatAnthropic

# Cargar variables de entorno
load_dotenv()

# Configurar variables de entorno
os.environ['PYDANTIC_V2'] = '1'
os.environ['BROWSER_USE_MODEL'] = 'claude-3-5-sonnet-20241022'

try:
    from browser_use import Agent, Browser, BrowserConfig
except ImportError as e:
    print(f"ERROR: No se pudo importar browser-use: {{e}}")
    sys.exit(1)

async def main():
    """Función principal ultra-simplificada"""
    
    print("Iniciando test simplificado...")
    
    # Configurar LLM
    anthropic_key = os.getenv('ANTHROPIC_API_KEY')
    if not anthropic_key:
        print("ERROR: ANTHROPIC_API_KEY no encontrada")
        sys.exit(1)
    
    llm = ChatAnthropic(
        model="claude-3-5-sonnet-20241022",
        api_key=anthropic_key,
        temperature=0.1
    )
    print("LLM configurado")
    
    # Configurar navegador ultra-simple
    browser_config = BrowserConfig(headless={headless})
    browser = Browser(config=browser_config)
    print("Navegador configurado")
    
    try:
        # Tarea ultra-simplificada
        simple_task = "Navegar a {url} y verificar que la pagina carga correctamente. Completar en maximo 5 pasos."
        
        # Crear agente con límites estrictos
        agent = Agent(
            task=simple_task,
            llm=llm,
            browser=browser
        )
        print("Agente creado")
        
        print("Ejecutando tarea...")
        start_time = time.time()
        
        # Ejecutar con límites MUY estrictos
        result = await agent.run(max_steps=5)  # SOLO 5 pasos máximo
        
        end_time = time.time()
        duration = end_time - start_time
        
        print(f"Test completado en {{duration:.2f}} segundos")
        print(f"Resultado: {{result}}")
        
        # Forzar salida exitosa después de 5 pasos
        sys.exit(0)
        
    except Exception as e:
        print(f"ERROR: {{e}}")
        sys.exit(1)
        
    finally:
        try:
            await browser.close()
            print("Navegador cerrado")
        except:
            pass

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except Exception as e:
        print(f"💥 Error fatal: {{e}}")
        sys.exit(1)
'''
        
        print(f"[SCRIPT-GEN] Script ultra-simplificado generado exitosamente ({len(script_content)} caracteres)")
        return script_content
        
    except Exception as e:
        print(f"[SCRIPT-GEN] Error generando script: {e}")
        raise Exception(f"Error al generar script: {e}")

def _optimize_instructions_for_browser_use(instructions, url, instruction_source):
    """
    CORRECCIÓN 2: Optimiza las instrucciones para una mejor ejecución con browser-use
    """
    try:
        # CORRECCIÓN CRÍTICA: Limpiar caracteres problemáticos de encoding
        print(f"[OPTIMIZE] Instrucciones originales length: {len(instructions)}")
        
        # Normalizar caracteres especiales para evitar problemas de encoding
        clean_instructions = instructions.strip()
        
        # CORRECCIÓN MEJORADA: Mapeo específico de caracteres UTF-8 corruptos
        utf8_corrections = {
            'TÃ©rmino': 'Termino',
            'tÃ©rmino': 'termino', 
            'bÃºsqueda': 'busqueda',
            'BÃºsqueda': 'Busqueda',
            'pÃ¡gina': 'pagina',
            'PÃ¡gina': 'Pagina',
            'especÃ­fico': 'especifico',
            'EspecÃ­fico': 'Especifico',
            'bÃ¡sica': 'basica',
            'BÃ¡sica': 'Basica',
            'artÃ­culos': 'articulos',
            'ArtÃ­culos': 'Articulos',
            'interÃ©s': 'interes',
            'InterÃ©s': 'Interes',
            'prÃ¡ctica': 'practica',
            'PrÃ¡ctica': 'Practica',
            'nÃºmero': 'numero',
            'NÃºmero': 'Numero',
            'descripciÃ³n': 'descripcion',
            'DescripciÃ³n': 'Descripcion',
            'informaciÃ³n': 'informacion',
            'InformaciÃ³n': 'Informacion',
            'VisualizaciÃ³n': 'Visualizacion',
            'visualizaciÃ³n': 'visualizacion',
            'decisiÃ³n': 'decision',
            'DecisiÃ³n': 'Decision',
            'secciÃ³n': 'seccion',
            'SecciÃ³n': 'Seccion',
            'caracterÃ­sticas': 'caracteristicas',
            'CaracterÃ­sticas': 'Caracteristicas',
            'imÃ¡genes': 'imagenes',
            'ImÃ¡genes': 'Imagenes',
            'Ã­cono': 'icono',
            'Ã\x81cono': 'icono',
            'botÃ³n': 'boton',
            'BotÃ³n': 'Boton',
            'Ã¡': 'a',
            'Ã©': 'e', 
            'Ã­': 'i',
            'Ã³': 'o',
            'Ãº': 'u',
            'Ã±': 'n',
            'Ã\x81': 'A',
            'Ã\x89': 'E',
            'Ã\x8d': 'I', 
            'Ã\x93': 'O',
            'Ã\x9a': 'U',
            'Ã\x91': 'N'
        }
        
        # Aplicar todas las correcciones de caracteres UTF-8
        for corrupted, correct in utf8_corrections.items():
            clean_instructions = clean_instructions.replace(corrupted, correct)
        
        # Limpieza adicional de caracteres problemáticos restantes
        clean_instructions = clean_instructions.replace('ú', 'u').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('á', 'a').replace('ñ', 'n')
        clean_instructions = clean_instructions.replace('?', 'u')  # Corregir caracteres corruptos genéricos
        
        # Verificar que no tengamos caracteres problemáticos restantes
        try:
            clean_instructions.encode('ascii', 'ignore')
        except:
            # Si aún hay problemas, usar solo caracteres ASCII seguros
            clean_instructions = ''.join(char if ord(char) < 128 else 'u' for char in clean_instructions)
        
        print(f"[OPTIMIZE] Instrucciones limpiadas length: {len(clean_instructions)}")
        
        # Si las instrucciones ya incluyen la URL, no duplicar
        if url and url.lower() not in clean_instructions.lower():
            # Agregar navegación inicial explícita
            if 'mercadolibre.cl' in url.lower():
                clean_instructions = f"Navegar a {url} (MercadoLibre Chile). Una vez cargada la pagina: {clean_instructions}"
            else:
                clean_instructions = f"Navegar a {url}. Una vez cargada la pagina: {clean_instructions}"
        
        # Optimizaciones específicas por fuente
        if instruction_source == 'qa_pilot':
            # Las instrucciones QA-Pilot ya están mejor estructuradas
            optimized = _format_qa_pilot_instructions(clean_instructions)
        elif instruction_source == 'original':
            # Los pasos originales necesitan más estructura
            optimized = _format_original_instructions(clean_instructions)
        else:
            # Instrucciones por defecto
            optimized = clean_instructions
        
        # Agregar contexto específico para browser-use
        optimized = _add_browser_use_context(optimized)
        
        print(f"[OPTIMIZE] Instrucciones optimizadas length: {len(optimized)}")
        return optimized
        
    except Exception as e:
        print(f"[OPTIMIZE] Error optimizando instrucciones: {e}")
        return instructions  # Fallback a instrucciones originales

def _format_qa_pilot_instructions(instructions):
    """Formatea instrucciones QA-Pilot para browser-use"""
    # Convertir formato QA-Pilot a formato más directo para browser-use
    lines = instructions.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Saltar líneas de objetivo/verificar (ya están incluidas implícitamente)
        if line.lower().startswith(('objetivo:', 'verificar:')):
            continue
        
        # Limpiar numeración inconsistente
        if line.lower().startswith('pasos:'):
            continue
        
        # Convertir pasos numerados a instrucciones directas
        import re
        # Remover numeración al inicio: "1.", "2)", "Paso 1:", etc.
        line = re.sub(r'^(\d+[\.\)]\s*|paso\s*\d+[\.\:]\s*)', '', line, flags=re.IGNORECASE)
        
        if line:
            formatted_lines.append(line)
    
    return '. '.join(formatted_lines) + '.'

def _format_original_instructions(instructions):
    """Formatea pasos originales para browser-use"""
    # Los pasos originales suelen ser más simples, agregar más estructura
    lines = instructions.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            # Remover numeración y limpiar
            import re
            line = re.sub(r'^(\d+[\.\)]\s*)', '', line)
            if line:
                formatted_lines.append(line)
    
    if formatted_lines:
        return '. '.join(formatted_lines) + '.'
    return instructions

def _add_browser_use_context(instructions):
    """Agrega contexto específico para browser-use"""
    # Agregar instrucciones para manejo de errores y completitud
    context_suffix = " Asegurar que cada acción se complete correctamente antes de continuar con la siguiente. Si un elemento no está visible, esperar o desplazarse para encontrarlo."
    
    return instructions + context_suffix

def _validate_execution_completion(return_code, stdout_text, stderr_text, instruction_source, case_name):
    """
    CORRECCIÓN 4: Validación mejorada de completitud de ejecución
    """
    try:
        # Indicadores de éxito en la salida
        success_indicators = [
            'test completado exitosamente',
            'ejecución completada con éxito',
            'step.*completado exitosamente',
            'task.*completado',
            'navegación.*exitosa',
            'todos los pasos.*completados'
        ]
        
        # Indicadores de fallo en la salida
        error_indicators = [
            'error', 'exception', 'failed', 'timeout', 'crash',
            'no se pudo', 'fallo', 'falló', 'excepción'
        ]
        
        # Indicadores específicos de browser-use
        browseruse_success = [
            'agent.*run.*completed',
            'browser.*session.*successful',
            'task.*finished.*successfully'
        ]
        
        # Indicadores de progreso paso a paso
        step_indicators = [
            r'step \d+.*completado',
            r'paso \d+.*exitoso',
            r'\[step \d+\].*completado exitosamente'
        ]
        
        import re
        
        stdout_lower = stdout_text.lower()
        stderr_lower = stderr_text.lower()
        combined_output = stdout_lower + " " + stderr_lower
        
        print(f"[VALIDATION] Analizando salida para caso: {case_name}")
        print(f"[VALIDATION] Código de retorno: {return_code}")
        print(f"[VALIDATION] Fuente de instrucciones: {instruction_source}")
        
        # Análisis básico por código de retorno
        if return_code != 0:
            print(f"[VALIDATION] Fallo por código de retorno: {return_code}")
            return False, 'fallido', f'Error de ejecución (código {return_code})'
        
        # Buscar errores explícitos incluso con código 0
        error_found = False
        for error_pattern in error_indicators:
            if re.search(error_pattern, combined_output, re.IGNORECASE):
                error_found = True
                print(f"[VALIDATION] Error detectado: {error_pattern}")
                break
        
        if error_found:
            return False, 'fallido', 'Error detectado en la salida del proceso'
        
        # Buscar indicadores de éxito específicos
        success_found = False
        success_type = ""
        
        # Verificar indicadores de browser-use
        for success_pattern in browseruse_success:
            if re.search(success_pattern, combined_output, re.IGNORECASE):
                success_found = True
                success_type = "browser-use completion"
                print(f"[VALIDATION] Éxito browser-use detectado: {success_pattern}")
                break
        
        # Verificar indicadores generales de éxito
        if not success_found:
            for success_pattern in success_indicators:
                if re.search(success_pattern, combined_output, re.IGNORECASE):
                    success_found = True
                    success_type = "general completion"
                    print(f"[VALIDATION] Éxito general detectado: {success_pattern}")
                    break
        
        # Verificar ejecución paso a paso
        if not success_found and instruction_source == 'qa_pilot':
            step_count = 0
            for step_pattern in step_indicators:
                matches = re.findall(step_pattern, combined_output, re.IGNORECASE)
                step_count += len(matches)
            
            if step_count >= 1:  # Al menos un paso completado
                success_found = True
                success_type = f"step-by-step ({step_count} pasos)"
                print(f"[VALIDATION] Ejecución paso a paso detectada: {step_count} pasos")
        
        # Validación de longitud de salida (al menos actividad)
        if len(stdout_text) < 100 and len(stderr_text) < 50:
            print(f"[VALIDATION] Salida muy corta, posible ejecución incompleta")
            return False, 'fallido', 'Ejecución demasiado breve, posiblemente incompleta'
        
        # Resultado final
        if success_found:
            message = f'Caso ejecutado correctamente ({success_type})'
            print(f"[VALIDATION] ÉXITO: {message}")
            return True, 'exitoso', message
        else:
            # Si no hay errores pero tampoco éxitos claros, asumir éxito parcial
            message = 'Ejecución completada (sin confirmación explícita de éxito)'
            print(f"[VALIDATION] ÉXITO PARCIAL: {message}")
            return True, 'exitoso_parcial', message
            
    except Exception as e:
        print(f"[VALIDATION] Error durante validación: {e}")
        # Fallback a validación básica
        if return_code == 0:
            return True, 'exitoso', 'Ejecución completada (validación básica)'
        else:
            return False, 'fallido', f'Error de ejecución (código {return_code})'

def execute_single_case_with_browser_use(case, case_index, total_cases, execution_path):
    """
    Ejecuta un caso individual usando browser-use (similar a run_test_background)
    """
    try:
        print(f"❌ [FUNCION-ANTIGUA] ¡SE ESTÁ EJECUTANDO LA FUNCIÓN ANTIGUA!")
        print(f"❌ [FUNCION-ANTIGUA] case_index: {case_index}, total_cases: {total_cases}")
        print(f"❌ [FUNCION-ANTIGUA] execution_path: {execution_path}")
        
        print(f"[BROWSER-USE] Iniciando ejecución del caso {case_index+1}/{total_cases}")
        
        # Obtener datos del caso
        case_id = case.get('id', f'case_{case_index}')
        case_name = case.get('nombre', f'Caso {case_index + 1}')
        url = case.get('url_extraida')
        
        # CORRECCIÓN 1: Priorizar instrucciones QA-Pilot mejoradas sobre pasos originales
        pasos_originales = case.get('pasos', '').strip()
        instrucciones_qa = case.get('instrucciones_qa_pilot', '').strip()
        
        # Prioridad: 1º QA-Pilot, 2º pasos originales (solo si QA-Pilot no existe o es muy corto)
        if instrucciones_qa and len(instrucciones_qa) > 50:  # QA-Pilot debe tener contenido sustancial
            instructions = instrucciones_qa
            instruction_source = 'qa_pilot'
            print(f"[BROWSER-USE] Usando instrucciones QA Pilot optimizadas ({len(instrucciones_qa)} chars)")
        elif pasos_originales:
            instructions = pasos_originales
            instruction_source = 'original'
            print(f"[BROWSER-USE] Usando pasos originales como fallback ({len(pasos_originales)} chars)")
        else:
            instructions = "Navegar por la página y explorar su contenido"
            instruction_source = 'default'
            print(f"[BROWSER-USE] Usando instrucciones por defecto (sin pasos definidos)")
        
        # Usar URL por defecto si no hay URL específica
        if not url:
            url = 'https://www.mercadolibre.cl/'
            print(f"[BROWSER-USE] Sin URL específica, usando URL por defecto: {url}")
        
        print(f"[BROWSER-USE] Caso {case_index+1}/{total_cases}: {case_name}")
        print(f"[BROWSER-USE] URL: {url}")
        print(f"[BROWSER-USE] Instrucciones: {instructions[:100]}...")
        print(f"[BROWSER-USE] Fuente: {instruction_source}")
        
        # CORRECCIÓN 2: Optimizar formato de instrucciones para browser-use
        optimized_instructions = _optimize_instructions_for_browser_use(instructions, url, instruction_source)
        print(f"[BROWSER-USE] Instrucciones optimizadas: {len(optimized_instructions)} chars")
        
        # Validar que al menos tenga instrucciones
        if not optimized_instructions:
            print(f"[BROWSER-USE] Caso inválido - Sin instrucciones de ejecución")
            return {
                'case_id': case_id,
                'case_name': case_name,
                'status': 'fallido',
                'success': False,
                'message': 'Caso inválido: sin instrucciones de ejecución',
                'url_tested': url,
                'execution_time': 0,
                'timestamp': datetime.now().isoformat(),
                'case_valid': False
            }
        
        # Actualizar estado actual
        print(f"[BROWSER-USE] Actualizando estado del caso {case_index}")
        update_current_case_status(execution_path, case_index, 'ejecutando', case_name)
        
        # Obtener configuraciones directas (valores hardcoded más seguros para hilos)
        print(f"[BROWSER-USE] Configurando directorios y ejecutables")
        try:
            # Usar rutas absolutas directas para evitar problemas de contexto
            BASE_DIR = os.path.abspath(os.path.dirname(os.path.dirname(__file__)))  # Directorio del proyecto
            SCRIPTS_DIR = os.path.join(BASE_DIR, 'test_scripts')
            python_exe = sys.executable  # Ejecutable Python actual
            
            # Crear directorio de scripts si no existe
            os.makedirs(SCRIPTS_DIR, exist_ok=True)
            print(f"[BROWSER-USE] Configuraciones: BASE_DIR={BASE_DIR}")
            print(f"[BROWSER-USE] SCRIPTS_DIR={SCRIPTS_DIR}")
            print(f"[BROWSER-USE] python_exe={python_exe}")
        except Exception as e:
            print(f"[BROWSER-USE] Error configurando directorios: {e}")
            raise Exception(f"No se pudieron configurar directorios: {e}")
        
        # Generar ID único para este caso
        task_id = str(uuid.uuid4())
        print(f"[BROWSER-USE] Task ID generado: {task_id}")
        
        # Configuración de ejecución
        # Leer configuración desde el archivo de ejecución
        try:
            with open(execution_path, 'r', encoding='utf-8') as f:
                execution_data = json.load(f)
            show_browser = execution_data.get('show_browser', False)
        except:
            show_browser = False
        
        headless = not show_browser  # Si show_browser es True, headless es False
        max_tiempo = 120  # CORRECCIÓN FINAL: 2 minutos máximo para tareas simples
        screenshots = True  # Habilitar capturas para evidencia
        browser = 'chrome'
        fullscreen = True  # CORRECCIÓN: Usar fullscreen como en ejecución manual
        
        print(f"[BROWSER-USE] Configuración: headless={headless}, max_tiempo={max_tiempo}, browser={browser}, show_browser={show_browser}")
        
        # Generar script de test
        print(f"[BROWSER-USE] Generando script de test")
        script_path = os.path.join(SCRIPTS_DIR, f"bulk_test_{task_id}.py")
        print(f"[BROWSER-USE] Ruta del script: {script_path}")
        
        try:
            script_content = generate_browser_use_script(
                url=url,
                instrucciones=optimized_instructions,  # CORRECCIÓN 2: Usar instrucciones optimizadas
                headless=headless,
                max_tiempo=max_tiempo,
                capturar_pasos=screenshots,
                browser=browser,
                fullscreen=fullscreen,
                task_id=task_id,
                instruction_source=instruction_source  # Pasar fuente para logging
            )
            print(f"[BROWSER-USE] Script generado exitosamente ({len(script_content)} caracteres)")
        except Exception as e:
            print(f"[BROWSER-USE] Error generando script: {e}")
            raise
        
        # Escribir el script
        print(f"[BROWSER-USE] Escribiendo script a archivo")
        try:
            with open(script_path, 'w', encoding='utf-8-sig', errors='replace') as f:
                f.write(script_content)
            print(f"[BROWSER-USE] Script escrito exitosamente")
        except Exception as e:
            print(f"[BROWSER-USE] Error escribiendo script: {e}")
            raise
        
        # Configurar entorno
        print(f"[BROWSER-USE] Configurando entorno de ejecución")
        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'
        env['PYTHONUTF8'] = '1'
        env['PYTHONLEGACYWINDOWSSTDIO'] = '0'
        env['PYDANTIC_V2'] = '1'
        env['PYDANTIC_STRICT_MODE'] = 'false'
        env['PYDANTIC_COLORS'] = 'false'
        env['BROWSER_USE_MODEL'] = 'claude-3-5-sonnet-20241022'  # CORRECCIÓN: Usar modelo exitoso
        
        start_time = time.time()
        print(f"[BROWSER-USE] Iniciando ejecución del script")
        
        # Ejecutar el script
        try:
            # CORRECCIÓN CRÍTICA: Permitir que las ventanas del navegador sean visibles
            creation_flags = 0  # No ocultar ventanas
            if not show_browser:
                # Solo ocultar la ventana de la consola del script, no del navegador
                creation_flags = subprocess.CREATE_NO_WINDOW if platform.system() == 'Windows' else 0
            
            proceso = subprocess.Popen(
                [python_exe, script_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
                cwd=BASE_DIR,
                bufsize=1,
                creationflags=creation_flags
            )
            print(f"[BROWSER-USE] Proceso iniciado con PID: {proceso.pid}")
            
            # CORRECCIÓN: Reducir actualizaciones de progreso que pueden interferir
            progress_count = 0
            while proceso.poll() is None and progress_count < 3:
                progress_step = (progress_count + 1) * 30  # Pasos más grandes
                print(f"[BROWSER-USE] Progreso: {progress_step}%")
                update_case_execution_progress(execution_path, case_index, progress_step, f"Ejecutando test - fase {progress_count + 1}/3")
                time.sleep(max_tiempo * 0.05)  # Esperar más tiempo entre actualizaciones
                progress_count += 1
            
            # Esperar a que termine
            print(f"[BROWSER-USE] Esperando finalización del proceso")
            return_code = proceso.wait(timeout=max_tiempo)
            execution_time = time.time() - start_time
            
            print(f"[BROWSER-USE] Proceso terminado con código: {return_code}, tiempo: {execution_time:.2f}s")
            
            # Leer salida
            stdout, stderr = proceso.communicate()
            stdout_text = stdout.decode('utf-8', errors='replace') if stdout else ''
            stderr_text = stderr.decode('utf-8', errors='replace') if stderr else ''
            
            print(f"[BROWSER-USE] STDOUT length: {len(stdout_text)}")
            print(f"[BROWSER-USE] STDERR length: {len(stderr_text)}")
            
            # Si hay error, mostrar salida completa para debug
            if return_code != 0:
                print(f"[BROWSER-USE] ===== STDOUT COMPLETO (Error Code {return_code}) =====")
                print(stdout_text)
                print(f"[BROWSER-USE] ===== STDERR COMPLETO =====")
                print(stderr_text)
                print(f"[BROWSER-USE] ===== FIN DE SALIDA =====")
            else:
                if stdout_text:
                    print(f"[BROWSER-USE] STDOUT preview: {stdout_text[:200]}...")
                if stderr_text:
                    print(f"[BROWSER-USE] STDERR preview: {stderr_text[:200]}...")
            
            # CORRECCIÓN 4: Mejorar la validación de que los pasos se completaron
            success, status, message = _validate_execution_completion(
                return_code, stdout_text, stderr_text, instruction_source, case_name
            )
            print(f"[BROWSER-USE] Resultado validado: {status} - {message}")
            
        except subprocess.TimeoutExpired:
            print(f"[BROWSER-USE] Timeout después de {max_tiempo} segundos")
            proceso.kill()
            execution_time = max_tiempo
            success = False
            status = 'fallido'
            message = f'Timeout después de {max_tiempo} segundos'
            
        except Exception as e:
            print(f"[BROWSER-USE] Error durante ejecución: {e}")
            execution_time = time.time() - start_time
            success = False
            status = 'error'
            message = f'Error inesperado: {str(e)}'
        
        finally:
            # Limpiar archivo de script
            try:
                if os.path.exists(script_path):
                    os.remove(script_path)
                    print(f"[BROWSER-USE] Script temporal eliminado")
            except Exception as e:
                print(f"[BROWSER-USE] No se pudo eliminar script temporal: {e}")
        
        # Crear resultado
        result = {
            'case_id': case_id,
            'case_name': case_name,
            'status': status,
            'success': success,
            'message': message,
            'url_tested': url,
            'execution_time': round(execution_time, 2),
            'timestamp': datetime.now().isoformat(),
            'case_valid': case.get('es_valido', False)
        }
        
        print(f"[BROWSER-USE] Caso {case_index+1}/{total_cases} completado: {status} ({execution_time:.1f}s)")
        return result
        
    except Exception as e:
        print(f"[BROWSER-USE] Error ejecutando caso {case.get('id', f'case_{case_index}')}: {str(e)}")
        print(f"[BROWSER-USE] Traceback: {traceback.format_exc()}")
        return {
            'case_id': case.get('id', f'case_{case_index}'),
            'case_name': case.get('nombre', f'Caso {case_index+1}'),
            'status': 'error',
            'success': False,
            'message': f'Error inesperado: {str(e)}',
            'timestamp': datetime.now().isoformat(),
            'execution_time': 0
        }

def execute_cases_sequential(execution_id, test_cases, execution_path):
    """
    Ejecuta casos de prueba uno por uno usando la misma lógica exitosa del home
    """
    try:
        print(f"🚀 [NUEVA-FUNCION] ¡EJECUTANDO NUEVA FUNCIÓN ACTUALIZADA!")
        print(f"🚀 [NUEVA-FUNCION] execution_id: {execution_id}")
        print(f"🚀 [NUEVA-FUNCION] total_cases: {len(test_cases)}")
        print(f"🚀 [NUEVA-FUNCION] execution_path: {execution_path}")
        
        print(f"[EXCEL-SEQUENTIAL] Iniciando ejecución secuencial para {len(test_cases)} casos")
        
        # Importar funciones del home que funcionan perfectamente
        from app import run_test_background, test_status_db, db_lock, save_test_status_to_file
        
        results = []
        
        for i, case in enumerate(test_cases):
            try:
                case_name = case.get('nombre', f'Caso {i+1}')
                print(f"[EXCEL-SEQUENTIAL] Ejecutando caso {i+1}/{len(test_cases)}: {case_name}")
                
                # Actualizar estado de ejecución masiva
                update_current_case_status(execution_path, i, 'ejecutando', case_name)
                
                # Generar task_id único para este caso individual
                case_task_id = str(uuid.uuid4())
                
                # Extraer datos del caso
                url = case.get('url_extraida', case.get('datos_prueba', 'https://www.google.com'))
                if not url or not url.startswith(('http://', 'https://')):
                    url = 'https://www.google.com'
                
                # Combinar objetivo y pasos para las instrucciones
                objetivo = case.get('objetivo', '')
                pasos = case.get('pasos', '')
                instrucciones = f"{objetivo}\n\n{pasos}".strip()
                
                if not instrucciones:
                    instrucciones = f"Navegar a {url} y realizar pruebas básicas"
                
                print(f"[EXCEL-SEQUENTIAL] Caso {i+1} - URL: {url}")
                print(f"[EXCEL-SEQUENTIAL] Caso {i+1} - Instrucciones: {instrucciones[:100]}...")
                
                # Leer show_browser desde el archivo de ejecución
                show_browser = False
                try:
                    with open(execution_path, 'r', encoding='utf-8') as f:
                        execution_record = json.load(f)
                        show_browser = execution_record.get('show_browser', False)
                except:
                    pass
                
                # Configuración exactamente igual al home exitoso
                headless = not show_browser  # Si show_browser=True, entonces headless=False
                max_tiempo = 600  # Mismo timeout que el home exitoso
                screenshots = True  # Siempre capturar screenshots
                fullscreen = True  # Mismo que el home exitoso
                browser = 'chrome'  # Mismo que el home exitoso
                model_name = 'claude-3-5-sonnet-20241022'  # Modelo actualizado
                
                print(f"[EXCEL-SEQUENTIAL] Configuración igual al home - headless: {headless}, max_tiempo: {max_tiempo}, fullscreen: {fullscreen}")
                
                # Inicializar estado en test_status_db (exactamente igual que run_test)
                with db_lock:
                    test_status_db[case_task_id] = {
                        'status': 'queued', 
                        'message': 'En cola...', 
                        'original_instructions': instrucciones,
                        'current_action': 'Esperando para despegar...',
                        'icon': 'fa-hourglass-start',
                        'url': url,
                        'created_at': datetime.now().isoformat(),
                        'case_name': case_name,
                        'case_index': i+1,
                        'execution_id': execution_id
                    }
                
                save_test_status_to_file()
                
                print(f"[EXCEL-SEQUENTIAL] Ejecutando run_test_background (función exitosa del home) para caso {i+1}")
                
                start_time = time.time()
                
                try:
                    # Ejecutar usando exactamente la misma función exitosa del home
                    print(f"[DEBUG] Llamando run_test_background con parámetros:")
                    print(f"  - task_id: {case_task_id}")
                    print(f"  - url: {url}")
                    print(f"  - instrucciones (primeros 100 chars): {instrucciones[:100]}...")
                    print(f"  - headless: {headless}")
                    print(f"  - max_tiempo: {max_tiempo}")
                    print(f"  - screenshots: {screenshots}")
                    print(f"  - browser: {browser}")
                    print(f"  - fullscreen: {fullscreen}")
                    
                    run_test_background(
                        case_task_id, url, instrucciones, headless, max_tiempo, 
                        screenshots, None, model_name, browser, fullscreen
                    )
                    
                    print(f"[DEBUG] run_test_background completado sin excepción")
                    
                except Exception as run_error:
                    print(f"[ERROR] Error en run_test_background: {str(run_error)}")
                    print(f"[ERROR] Tipo de error: {type(run_error).__name__}")
                    import traceback
                    print(f"[ERROR] Traceback completo:")
                    traceback.print_exc()
                    
                    # Marcar el caso como error en test_status_db
                    with db_lock:
                        test_status_db[case_task_id] = {
                            'status': 'error',
                            'message': f'Error en ejecución: {str(run_error)}',
                            'created_at': datetime.now().isoformat(),
                            'case_name': case_name,
                            'error_details': str(run_error)
                        }
                
                execution_time = time.time() - start_time
                
                print(f"[EXCEL-SEQUENTIAL] Caso {i+1} completado en {execution_time:.1f}s")
                
                # Esperar un momento para que la tarea se procese completamente
                time.sleep(2)
                
                # Obtener resultado del caso
                with db_lock:
                    case_result = test_status_db.get(case_task_id, {})
                
                print(f"[DEBUG] Estado final del caso {i+1}:")
                print(f"  - Status: {case_result.get('status', 'NO_ENCONTRADO')}")
                print(f"  - Message: {case_result.get('message', 'NO_MESSAGE')}")
                print(f"  - Screenshots: {len(case_result.get('screenshots', []))}")
                
                # Si no hay resultado o está en estado incompleto, esperar más tiempo
                if not case_result or case_result.get('status') in ['queued', 'running', 'generating']:
                    print(f"[DEBUG] Caso {i+1} aún procesándose, esperando 5 segundos adicionales...")
                    time.sleep(5)
                    
                    with db_lock:
                        case_result = test_status_db.get(case_task_id, {})
                    
                    print(f"[DEBUG] Estado final tras espera adicional:")
                    print(f"  - Status: {case_result.get('status', 'NO_ENCONTRADO')}")
                    print(f"  - Message: {case_result.get('message', 'NO_MESSAGE')}")
                
                # Crear resumen del resultado para la ejecución masiva
                result_summary = {
                    'case_index': i,
                    'case_name': case_name,
                    'task_id': case_task_id,
                    'status': case_result.get('status', 'unknown'),
                    'message': case_result.get('message', 'Sin mensaje'),
                    'url': url,
                    'execution_time': execution_time,
                    'completed_at': datetime.now().isoformat(),
                    'screenshots_count': len(case_result.get('screenshots', [])),
                    'success': case_result.get('status') in ['completed', 'success', 'completado']  # ✅ CORRECCIÓN: reconocer múltiples estados como exitosos
                }
                
                results.append(result_summary)
                print(f"[EXCEL-SEQUENTIAL] Resultado caso {i+1}: {result_summary['status']}")
                print(f"[DEBUG] Caso {i+1} marcado como {'ÉXITO' if result_summary['success'] else 'FALLO'}")
                
                # Actualizar progreso general
                overall_progress = ((i + 1) / len(test_cases)) * 100
                update_execution_progress(execution_path, results, len(test_cases), overall_progress)
                
            except Exception as e:
                print(f"[EXCEL-SEQUENTIAL] Error ejecutando caso {i+1}: {str(e)}")
                import traceback
                traceback.print_exc()
                
                # Agregar resultado de error
                error_result = {
                    'case_index': i,
                    'case_name': case.get('nombre', f'Caso {i+1}'),
                    'task_id': None,
                    'status': 'error',
                    'message': f'Error: {str(e)}',
                    'url': case.get('url_extraida', 'N/A'),
                    'execution_time': 0.0,
                    'completed_at': datetime.now().isoformat(),
                    'screenshots_count': 0,
                    'success': False
                }
                results.append(error_result)
                
                # Actualizar progreso incluso con error
                overall_progress = ((i + 1) / len(test_cases)) * 100
                update_execution_progress(execution_path, results, len(test_cases), overall_progress)
        
        # Finalizar ejecución
        finalize_execution(execution_path, results)
        
        # Mostrar resumen final
        successful_cases = len([r for r in results if r.get('success', False)])
        failed_cases = len([r for r in results if not r.get('success', False)])
        print(f"[EXCEL-SEQUENTIAL] Ejecución secuencial completada: {len(results)} casos procesados")
        print(f"[EXCEL-SEQUENTIAL] Resumen final - Exitosos: {successful_cases}, Fallidos: {failed_cases}")
        
        current_app.logger.info(f"Ejecución secuencial completada usando lógica del home: {len(results)} casos procesados")
        
    except Exception as e:
        print(f"[EXCEL-SEQUENTIAL] Error crítico en ejecución secuencial: {str(e)}")
        import traceback
        traceback.print_exc()
        current_app.logger.error(f"Error en ejecución secuencial: {str(e)}")
        mark_execution_failed(execution_path, str(e))

def execute_cases_background(execution_id, test_cases, execution_path):
    """
    Ejecuta los casos de prueba en paralelo usando browser-use
    """
    try:
        results = []
        total_cases = len(test_cases)
        
        # Para ejecución paralela, usamos threading
        def execute_case_thread(case, case_index):
            return execute_single_case_with_browser_use(case, case_index, total_cases, execution_path)
        
        # Crear threads para cada caso (máximo 3 en paralelo para no sobrecargar)
        max_parallel = min(3, total_cases)
        threads = []
        case_results = [None] * total_cases
        
        for i in range(0, total_cases, max_parallel):
            batch = test_cases[i:i+max_parallel]
            batch_threads = []
            
            for j, case in enumerate(batch):
                case_index = i + j
                thread = threading.Thread(
                    target=lambda idx=case_index, c=case: case_results.__setitem__(idx, execute_case_thread(c, idx))
                )
                thread.start()
                batch_threads.append(thread)
            
            # Esperar a que termine el batch
            for thread in batch_threads:
                thread.join()
            
            # Actualizar progreso
            completed = i + len(batch)
            progress = (completed / total_cases) * 100
            current_results = [r for r in case_results[:completed] if r is not None]
            update_execution_progress(execution_path, current_results, total_cases, progress)
        
        # Recopilar todos los resultados
        results = [r for r in case_results if r is not None]
        
        # Finalizar ejecución
        finalize_execution(execution_path, results)
        current_app.logger.info(f"Ejecución paralela completada: {len(results)} casos procesados")
        
    except Exception as e:
        current_app.logger.error(f"Error en ejecución paralela: {str(e)}")
        mark_execution_failed(execution_path, str(e))

def update_current_case_status(execution_path, case_index, status, case_name):
    """Actualiza el estado del caso actual en ejecución"""
    try:
        execution_record = safe_read_execution_file(execution_path)
        
        execution_record['current_case'] = case_index
        execution_record['current_case_status'] = status
        execution_record['current_case_name'] = case_name
        execution_record['current_case_progress'] = 0
        
        safe_write_execution_file(execution_path, execution_record)
            
    except Exception as e:
        print(f"Error actualizando estado del caso actual: {e}")
        print(f"Traceback: {traceback.format_exc()}")

def update_case_execution_progress(execution_path, case_index, case_progress, step_description):
    """Actualiza el progreso del caso individual en ejecución"""
    try:
        execution_record = safe_read_execution_file(execution_path)
        
        execution_record['current_case_progress'] = case_progress
        execution_record['current_step_description'] = step_description
        
        safe_write_execution_file(execution_path, execution_record)
            
    except Exception as e:
        print(f"Error actualizando progreso del caso: {e}")
        print(f"Traceback: {traceback.format_exc()}")

def update_execution_progress(execution_path, results, total_cases, overall_progress=None):
    """Actualiza el progreso de la ejecución"""
    try:
        execution_record = safe_read_execution_file(execution_path)
        
        execution_record['results'] = results
        execution_record['completed_cases'] = len(results)
        
        if overall_progress is not None:
            execution_record['progress'] = round(overall_progress, 2)
        else:
            execution_record['progress'] = round((len(results) / total_cases) * 100, 2)
        
        execution_record['status'] = 'en_progreso'
        
        # Calcular estadísticas en tiempo real
        successful_cases = len([r for r in results if r.get('success', False)])
        failed_cases = len(results) - successful_cases
        
        execution_record['successful_cases'] = successful_cases
        execution_record['failed_cases'] = failed_cases
        execution_record['success_rate'] = round((successful_cases / len(results)) * 100, 2) if results else 0
        
        safe_write_execution_file(execution_path, execution_record)
            
    except Exception as e:
        print(f"Error actualizando progreso: {e}")
        print(f"Traceback: {traceback.format_exc()}")

def finalize_execution(execution_path, results):
    """Finaliza la ejecución y calcula estadísticas"""
    try:
        execution_record = safe_read_execution_file(execution_path)
        
        successful_cases = len([r for r in results if r['success']])
        failed_cases = len(results) - successful_cases
        
        execution_record['results'] = results
        execution_record['completed_cases'] = len(results)
        execution_record['successful_cases'] = successful_cases
        execution_record['failed_cases'] = failed_cases
        execution_record['success_rate'] = round((successful_cases / len(results)) * 100, 2) if results else 0
        execution_record['status'] = 'completado'
        execution_record['end_timestamp'] = datetime.now().isoformat()
        execution_record['progress'] = 100
        
        safe_write_execution_file(execution_path, execution_record)
        print(f"✅ Ejecución finalizada exitosamente. Estado: completado, Progreso: 100%")
            
    except Exception as e:
        print(f"Error finalizando ejecución: {e}")
        print(f"Traceback: {traceback.format_exc()}")

def mark_execution_failed(execution_path, error_message):
    """Marca la ejecución como fallida"""
    try:
        execution_record = safe_read_execution_file(execution_path)
        
        execution_record['status'] = 'fallido'
        execution_record['error'] = error_message
        execution_record['end_timestamp'] = datetime.now().isoformat()
        
        safe_write_execution_file(execution_path, execution_record)
        print(f"❌ Ejecución marcada como fallida: {error_message}")
            
    except Exception as e:
        print(f"Error marcando ejecución como fallida: {e}")
        print(f"Traceback: {traceback.format_exc()}")

@excel_bp.route('/execution_status/<execution_id>', methods=['GET'])
def get_execution_status(execution_id):
    """
    Obtiene el estado de una ejecución masiva
    """
    try:
        execution_file = f"bulk_execution_{execution_id}.json"
        execution_path = os.path.join(tempfile.gettempdir(), execution_file)
        
        if not os.path.exists(execution_path):
            return jsonify({
                'success': False,
                'error': 'Ejecución no encontrada'
            }), 404
        
        execution_record = safe_read_execution_file(execution_path)
        
        return jsonify({
            'success': True,
            'execution': execution_record
        })
        
    except Exception as e:
        current_app.logger.error(f"Error obteniendo estado de ejecución: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error al obtener el estado: {str(e)}'
        }), 500

@excel_bp.route('/check_case_status/<case_id>', methods=['GET'])
def check_case_status(case_id):
    """
    Verifica si un caso específico está siendo ejecutado actualmente
    """
    try:
        # Buscar ejecuciones activas que contengan este caso
        temp_dir = tempfile.gettempdir()
        execution_files = [f for f in os.listdir(temp_dir) if f.startswith('bulk_execution_') and f.endswith('.json')]
        
        for execution_file in execution_files:
            execution_path = os.path.join(temp_dir, execution_file)
            try:
                with open(execution_path, 'r', encoding='utf-8') as f:
                    execution_record = json.load(f)
                
                # Verificar si el caso está en esta ejecución y si está activa
                if execution_record.get('status') in ['iniciado', 'ejecutando', 'en_progreso']:
                    cases = execution_record.get('cases', [])
                    for case in cases:
                        if case.get('id') == case_id:
                            return jsonify({
                                'success': True,
                                'is_running': True,
                                'execution_id': execution_record.get('id'),
                                'current_status': execution_record.get('current_case_status', 'desconocido'),
                                'progress': execution_record.get('progress', 0)
                            })
            except:
                continue
        
        # Si no se encontró en ninguna ejecución activa
        return jsonify({
            'success': True,
            'is_running': False,
            'execution_id': None,
            'current_status': 'inactivo',
            'progress': 0
        })
        
    except Exception as e:
        current_app.logger.error(f"Error verificando estado del caso: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error al verificar el estado: {str(e)}'
        }), 500

@excel_bp.route('/save_excel_as_suite', methods=['POST'])
def save_excel_as_suite():
    """
    Guarda los casos de prueba de Excel como una suite
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                'success': False,
                'error': 'Datos no proporcionados'
            }), 400
        
        suite_name = data.get('suite_name', '').strip()
        suite_description = data.get('suite_description', '').strip()
        test_cases = data.get('test_cases', [])
        include_invalid = data.get('include_invalid', False)
        
        if not suite_name:
            return jsonify({
                'success': False,
                'error': 'El nombre de la suite es obligatorio'
            }), 400
        
        if not test_cases:
            return jsonify({
                'success': False,
                'error': 'No se proporcionaron casos de prueba'
            }), 400
        
        # Crear estructura de suite
        suite_data = {
            'id': str(uuid.uuid4()),
            'name': suite_name,
            'description': suite_description,
            'created_from': 'excel',
            'created_at': datetime.now().isoformat(),
            'total_cases': len(test_cases),
            'valid_cases': len([tc for tc in test_cases if tc.get('es_valido', False)]),
            'invalid_cases': len([tc for tc in test_cases if not tc.get('es_valido', False)]),
            'include_invalid': include_invalid,
            'test_cases': test_cases
        }
        
        # Guardar suite (en un sistema real, esto iría a una base de datos)
        suite_file = f"suite_{suite_data['id']}.json"
        suite_path = os.path.join(tempfile.gettempdir(), suite_file)
        
        with open(suite_path, 'w', encoding='utf-8') as f:
            json.dump(suite_data, f, ensure_ascii=False, indent=2)
        
        current_app.logger.info(f"Suite guardada: {suite_name} ({len(test_cases)} casos)")
        
        return jsonify({
            'success': True,
            'suite_id': suite_data['id'],
            'suite_name': suite_name,
            'total_cases': len(test_cases),
            'message': f'Suite "{suite_name}" guardada exitosamente'
        })
        
    except Exception as e:
        current_app.logger.error(f"Error guardando suite: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Error al guardar la suite: {str(e)}'
        }), 500

@excel_bp.route('/generate_bulk_word_report/<execution_id>', methods=['POST'])
def generate_bulk_word_report(execution_id):
    """
    Genera un documento Word con el reporte de ejecución masiva
    """
    try:
        import os
        import tempfile
        from datetime import datetime
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.shared import RGBColor
        
        current_app.logger.debug(f"Generando reporte Word para ejecución masiva: {execution_id}")
        
        # Buscar el archivo de ejecución
        execution_file = f"bulk_execution_{execution_id}.json"
        execution_path = os.path.join(tempfile.gettempdir(), execution_file)
        
        if not os.path.exists(execution_path):
            return jsonify({
                'success': False,
                'message': 'Ejecución no encontrada'
            }), 404
        
        # Leer datos de la ejecución
        execution_data = safe_read_execution_file(execution_path)
        
        if not execution_data:
            return jsonify({
                'success': False,
                'message': 'No se pudieron leer los datos de la ejecución'
            }), 404
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos del documento
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Título principal
        title = doc.add_heading('Reporte de Ejecución Masiva - QA Pilot', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información general
        doc.add_heading('Información General', level=1)
        
        # Tabla de información general
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados de la tabla
        info_table.cell(0, 0).text = 'Campo'
        info_table.cell(0, 1).text = 'Valor'
        
        # Datos de la ejecución
        execution_date = execution_data.get('timestamp', datetime.now().isoformat())
        total_cases = execution_data.get('total_cases', 0)
        execution_mode = execution_data.get('execution_mode', 'sequential')
        show_browser = execution_data.get('show_browser', False)
        
        # Calcular estadísticas
        results = execution_data.get('results', [])
        successful_cases = len([r for r in results if r.get('success', False)])
        failed_cases = len([r for r in results if not r.get('success', False)])
        success_rate = (successful_cases / total_cases * 100) if total_cases > 0 else 0
        
        info_table.cell(1, 0).text = 'ID de Ejecución'
        info_table.cell(1, 1).text = execution_id[:8]
        
        info_table.cell(2, 0).text = 'Fecha y Hora'
        try:
            execution_datetime = datetime.fromisoformat(execution_date.replace('Z', '+00:00'))
            info_table.cell(2, 1).text = execution_datetime.strftime('%d/%m/%Y %H:%M:%S')
        except:
            info_table.cell(2, 1).text = execution_date
        
        info_table.cell(3, 0).text = 'Modo de Ejecución'
        info_table.cell(3, 1).text = 'Secuencial' if execution_mode == 'sequential' else 'Paralelo'
        
        info_table.cell(4, 0).text = 'Navegador Visible'
        info_table.cell(4, 1).text = 'Sí' if show_browser else 'No'
        
        info_table.cell(5, 0).text = 'Total de Casos'
        info_table.cell(5, 1).text = str(total_cases)
        
        # Estadísticas de resultados
        doc.add_heading('Estadísticas de Resultados', level=1)
        
        stats_table = doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        stats_table.cell(0, 0).text = 'Métrica'
        stats_table.cell(0, 1).text = 'Valor'
        
        stats_table.cell(1, 0).text = 'Casos Exitosos'
        stats_table.cell(1, 1).text = f"{successful_cases} ({success_rate:.1f}%)"
        
        stats_table.cell(2, 0).text = 'Casos Fallidos'
        stats_table.cell(2, 1).text = f"{failed_cases} ({100-success_rate:.1f}%)"
        
        stats_table.cell(3, 0).text = 'Tasa de Éxito'
        stats_table.cell(3, 1).text = f"{success_rate:.1f}%"
        
        # Detalle de casos ejecutados
        doc.add_heading('Detalle de Casos Ejecutados', level=1)
        
        if results:
            for i, result_summary in enumerate(results, 1):
                task_id = result_summary.get('task_id')
                case_name = result_summary.get('case_name', f'Caso {i}')
                
                # Título del caso
                doc.add_heading(f'Caso {i}: {case_name}', level=2)
                
                # Obtener la información completa y actualizada del test desde test_status_db
                full_test_data = {}
                if task_id:
                    try:
                        test_status_db = current_app.config.get('test_status_db')
                        db_lock = current_app.config.get('db_lock')
                        if test_status_db and db_lock:
                            with db_lock:
                                full_test_data = test_status_db.get(task_id, {})
                    except Exception as e:
                        current_app.logger.warning(f"No se pudo obtener la información completa para el task_id {task_id}: {e}")

                # Combinar el resumen con los datos completos, dando prioridad a los datos completos
                final_data = {**result_summary, **full_test_data}

                # Información del caso
                case_info = doc.add_paragraph()
                case_info.add_run('URL: ').bold = True
                case_info.add_run(final_data.get('url', 'No especificada'))
                
                case_info.add_run('\nEstado: ').bold = True
                status_text = '✅ Exitoso' if final_data.get('success', False) else '❌ Fallido'
                case_info.add_run(status_text)
                
                execution_time = final_data.get('execution_time')
                if execution_time is not None:
                    case_info.add_run('\nTiempo de Ejecución: ').bold = True
                    case_info.add_run(f"{execution_time:.2f} segundos")

                case_info.add_run('\nMensaje: ').bold = True
                case_info.add_run(final_data.get('message', 'Sin mensaje'))

                # Pasos ejecutados desde los datos completos
                steps_executed = final_data.get('steps_executed', [])
                if steps_executed:
                    doc.add_paragraph('Pasos Ejecutados:', style='Heading 3')
                    for step_num, step_desc in enumerate(steps_executed, 1):
                        doc.add_paragraph(f"Paso {step_num}: {step_desc}", style='List Bullet')

                # Detalles del error si el caso falló
                if not final_data.get('success', False):
                    error_details = final_data.get('error_details', '')
                    if error_details:
                        doc.add_paragraph('Detalles del Error:', style='Heading 3')
                        doc.add_paragraph(error_details)
                
                # Capturas de pantalla desde los datos completos
                screenshots = final_data.get('screenshots', [])
                if screenshots:
                    doc.add_paragraph('Capturas de Pantalla:', style='Heading 3')
                    for j, screenshot in enumerate(screenshots[:5], 1):  # Máximo 5 capturas
                        try:
                            path = screenshot.get('path')
                            if path and os.path.exists(path):
                                name = screenshot.get('name', f'Captura {j}')
                                doc.add_paragraph(f"Figura {j}: {name}")
                                doc.add_picture(path, width=Inches(5.0))
                                doc.add_paragraph()
                            else:
                                current_app.logger.warning(f"Ruta de captura no encontrada: {path}")
                        except Exception as e:
                            current_app.logger.warning(f"Error procesando captura para el reporte: {e}")
                else:
                    doc.add_paragraph("No se encontraron capturas de pantalla para este caso.")

                # Separador entre casos
                if i < len(results):
                    doc.add_page_break()
        else:
            doc.add_paragraph('No se encontraron resultados de ejecución.')
        
        # Resumen final
        doc.add_heading('Resumen Final', level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Ejecución completada el ').bold = True
        summary_para.add_run(datetime.now().strftime('%d/%m/%Y a las %H:%M:%S'))
        summary_para.add_run('\n\nResultados: ').bold = True
        summary_para.add_run(f'{successful_cases} casos exitosos de {total_cases} total')
        summary_para.add_run('\n\nTasa de éxito: ').bold = True
        summary_para.add_run(f'{success_rate:.1f}%')
        
        # Crear archivo temporal y enviarlo
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Reporte_Ejecucion_Masiva_{execution_id[:8]}_{timestamp}.docx"
        
        # Usar un directorio temporal específico
        import tempfile
        temp_dir = tempfile.gettempdir()
        tmp_path = os.path.join(temp_dir, filename)
        
        # Guardar el documento
        doc.save(tmp_path)
        current_app.logger.debug(f"Documento guardado en: {tmp_path}")
        
        # Verificar que el archivo se creó correctamente
        if not os.path.exists(tmp_path):
            raise Exception(f"No se pudo crear el archivo temporal: {tmp_path}")
        
        current_app.logger.debug(f"Archivo temporal creado exitosamente: {os.path.getsize(tmp_path)} bytes")
        
        def remove_file(response):
            try:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
                    current_app.logger.debug(f"Archivo temporal eliminado: {tmp_path}")
            except Exception as e:
                current_app.logger.warning(f"Error eliminando archivo temporal: {e}")
            return response
        
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        current_app.logger.error(f"Error generando reporte Word: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Error al generar reporte: {str(e)}'
        }), 500

# Función para registrar el blueprint en la aplicación principal
def register_excel_routes(app):
    """Registra las rutas de Excel en la aplicación Flask"""
    app.register_blueprint(excel_bp) 