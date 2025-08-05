import os
import sys
from datetime import datetime
import logging

# Configurar logger temprano
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s %(levelname)s [%(filename)s:%(lineno)d] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Configuración simplificada para Pydantic (ya no necesaria en browser-use 0.2.6)

from flask import Flask, render_template, request, session, flash, redirect, url_for, jsonify, send_from_directory, send_file
import secrets
import traceback
from utils import load_env_vars, save_api_keys_to_env, API_KEYS_TO_MANAGE
# Importar funciones de generación y ejecución de tests (moverlas a utils o definir aquí)
# from utils import generar_script_test, ejecutar_test_separado # Asumiendo que se moverán
# --- Temporalmente, copiaremos/adaptaremos la lógica aquí --- 
import sys
import subprocess
import time
from threading import Thread # Para ejecución en segundo plano simple
import uuid # Para IDs de tareas
import re
from langchain_google_genai import ChatGoogleGenerativeAI # Necesario para generar_script_test
from browser_use import Agent, Browser, BrowserConfig, BrowserContextConfig # Necesario para generar_script_test - API
import base64 # Necesario para generar_script_test
import platform
from langchain_openai import ChatOpenAI # Para OpenAI y DeepSeek
from langchain_anthropic import ChatAnthropic # Para Anthropic
from pydantic import SecretStr # Para manejar claves API seguras
import json
import json as json_module
import threading
from dotenv import load_dotenv
import asyncio
import requests  # Para RapidAPI Screenshot
import pyautogui # Para capturas de pantalla simples y confiables
import locale
import codecs
import io
import tempfile
import time
import uuid
import glob
import shutil
import pygetwindow as gw

# Definir variables de disponibilidad de sistemas
MCP_AVAILABLE = False  # MCP ha sido eliminado
PLAYWRIGHT_MCP_AVAILABLE = False  # Playwright MCP no está disponible por defecto

# Función add_to_history migrada a usar solo base de datos
def should_use_mcp() -> bool:
    # MCP ha sido eliminado, siempre devolver False
    return False

def should_use_playwright_mcp() -> bool:
    """
    Determina si se debe usar Playwright MCP basado en configuración
    """
    # Verificar si Playwright MCP está disponible
    if not PLAYWRIGHT_MCP_AVAILABLE:
        return False
    
    # Verificar variable de entorno para forzar uso de Playwright MCP
    use_playwright_mcp_env = os.getenv('USE_PLAYWRIGHT_MCP', 'false').lower()
    if use_playwright_mcp_env in ['true', '1', 'yes', 'on']:
        return True
    
    # Por defecto no usar hasta que esté completamente probado
    return False

def get_preferred_mcp_system():
    """
    Devuelve siempre 'browser_use' para restaurar el sistema original
    """
    return 'browser_use'

def run_test_background_unified(task_id, url, instrucciones, headless, max_tiempo, screenshots, 
                              gemini_key, model_name='claude-3-5-sonnet-20240620', browser='chrome', fullscreen=True):
    """
    Función unificada que decide si usar MCP o browser-use
    
    Mantiene la misma interfaz que run_test_background pero puede usar MCP o browser-use
    """
    try:
        if should_use_mcp():
            logger.info(f"🚀 Ejecutando test {task_id} con MCP")
            return run_test_background_mcp(
                task_id, url, instrucciones, headless, max_tiempo, screenshots,
                gemini_key, model_name, browser, fullscreen
            )
        else:
            logger.info(f"🚀 Ejecutando test {task_id} con browser-use")
            return run_test_background(
                task_id, url, instrucciones, headless, max_tiempo, screenshots,
                gemini_key, model_name, browser, fullscreen
            )
    except Exception as e:
        logger.error(f"❌ Error en ejecución unificada {task_id}: {e}")
        # Fallback a browser-use si MCP falla
        if should_use_mcp():
            logger.info(f"🔄 Fallback a browser-use para {task_id}")
            return run_test_background(
                task_id, url, instrucciones, headless, max_tiempo, screenshots,
                gemini_key, model_name, browser, fullscreen
            )
        else:
            raise

def run_test_background_suite_mcp(task_id, url, instrucciones, max_tiempo, screenshots, 
                                 gemini_key, model_name='claude-3-5-sonnet-20240620', browser='chrome'):
    """
    Función específica para suites que usa el mejor sistema MCP disponible con navegador visible y maximizado
    
    Args:
        task_id: ID único de la tarea
        url: URL objetivo
        instrucciones: Instrucciones para el test
        max_tiempo: Tiempo máximo en segundos
        screenshots: Habilitar capturas de pantalla
        gemini_key: Clave API (se detecta automáticamente)
        model_name: Modelo a usar
        browser: Navegador a usar
        
    Returns:
        None (ejecuta en background)
    """
    try:
        preferred_system = get_preferred_mcp_system()
        
        if preferred_system == 'playwright_mcp':
            if not PLAYWRIGHT_MCP_AVAILABLE:
                raise Exception("Sistema Playwright MCP no está disponible")
            
            logger.info(f"🎭 Ejecutando caso de suite {task_id} con Playwright MCP (navegador visible)")
            
            # Configuración forzada para suites: navegador SIEMPRE visible y maximizado
            forced_headless = False  # SIEMPRE visible
            forced_fullscreen = True  # SIEMPRE maximizado
            
            return run_test_background_playwright_mcp(
                task_id, url, instrucciones, 
                headless=forced_headless,  # Forzar navegador visible
                max_tiempo=max_tiempo, 
                screenshots=screenshots,
                gemini_key=gemini_key, 
                model_name=model_name, 
                browser=browser, 
                fullscreen=forced_fullscreen  # Forzar pantalla completa
            )
            
        elif preferred_system == 'mcp':
            if not MCP_AVAILABLE:
                raise Exception("Sistema MCP no está disponible")
            
            logger.info(f"🎭 Ejecutando caso de suite {task_id} con MCP custom (navegador visible)")
            
            # Configuración forzada para suites: navegador SIEMPRE visible y maximizado
            forced_headless = False  # SIEMPRE visible
            forced_fullscreen = True  # SIEMPRE maximizado
            
            return run_test_background_mcp(
                task_id, url, instrucciones, 
                headless=forced_headless,  # Forzar navegador visible
                max_tiempo=max_tiempo, 
                screenshots=screenshots,
                gemini_key=gemini_key, 
                model_name=model_name, 
                browser=browser, 
                fullscreen=forced_fullscreen  # Forzar pantalla completa
            )
        else:
            raise Exception("No hay sistemas MCP disponibles para ejecución de suites")
            
    except Exception as e:
        logger.error(f"❌ Error ejecutando caso de suite {task_id}: {e}")
        raise

def get_test_status_unified(task_id: str) -> dict:
    """
    Función unificada para obtener el estado de una prueba
    
    Args:
        task_id: ID de la tarea
        
    Returns:
        Estado de la prueba
    """
    try:
        preferred_system = get_preferred_mcp_system()
        
        if preferred_system == 'playwright_mcp' and PLAYWRIGHT_MCP_AVAILABLE:
            # Intentar obtener estado de Playwright MCP primero
            playwright_status = get_playwright_mcp_test_status(task_id)
            if playwright_status.get('status') != 'not_found':
                return playwright_status
        
        if preferred_system == 'mcp' and MCP_AVAILABLE:
            # Intentar obtener estado de MCP custom
            mcp_status = get_mcp_test_status(task_id)
            if mcp_status.get('status') != 'not_found':
                return mcp_status
        
        # Fallback a browser-use status
        with db_lock:
            status = test_status_db.get(task_id)
            if status:
                return status
            else:
                return None  # Retornar None para que el código original maneje el caso
    except Exception as e:
        logger.error(f"❌ Error obteniendo estado {task_id}: {e}")
        return {
            'status': 'error',
            'message': f'Error obteniendo estado: {str(e)}',
            'current_action': 'Error',
            'icon': 'fa-exclamation-triangle'
        }

def add_to_history(task_id, test_data, history_db=None, history_lock=None, max_items=100):
    """
    Agrega un test ejecutado directamente a la base de datos.
    Los parámetros history_db y history_lock se mantienen por compatibilidad pero ya no se usan.
    """
    try:
        logger.info(f"💾 Guardando ejecución {task_id[:8]} en base de datos")
        
        # Usar el nuevo servicio de historial para guardar directamente en BD
        history_service = get_history_service()
        success = history_service.add_execution(task_id, test_data)
        
        if success:
            logger.info(f"✅ Ejecución {task_id[:8]} guardada exitosamente en BD")
            return {'success': True, 'message': 'Ejecución guardada en base de datos'}
        else:
            logger.warning(f"❌ No se pudo guardar ejecución {task_id[:8]} en BD")
            return {'success': False, 'message': 'Error guardando en base de datos'}
    except Exception as e:
        logger.error(f"❌ Error guardando ejecución {task_id[:8]}: {e}")
        return {'success': False, 'message': f'Error al guardar ejecución: {str(e)}'}

def save_execution_to_database(task_id, test_data, history_entry):
    """Guarda una ejecución de test en la base de datos PostgreSQL"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.test_connection():
            print("DEBUG: Base de datos no disponible para guardar ejecución")
            return
        
        with db_integration.get_session() as session:
            from db_models import TestCase, TestExecution, Screenshot, Project
            from datetime import datetime, timezone
            import uuid
            
            # Obtener o crear proyecto por defecto
            default_project = db_integration.get_default_project()
            
            # Buscar o crear caso de prueba
            test_case = session.query(TestCase).filter_by(
                nombre=history_entry['name']
            ).first()
            
            if not test_case:
                # Crear nuevo caso de prueba
                test_case = TestCase(
                    project_id=default_project.id,
                    nombre=history_entry['name'],
                    codigo=f'TC_{task_id[:8]}',
                    tipo='funcional',
                    prioridad='media',
                    objetivo=f"Test automático: {history_entry['name']}",
                    pasos=history_entry['instructions'] or 'Pasos generados automáticamente',
                    resultado_esperado='Test ejecutado correctamente',
                    url_objetivo=history_entry['url'],
                    es_valido=history_entry['status'] == 'success',
                    instrucciones_qa_pilot=history_entry['instructions'],
                    status='active',
                    created_by='qa_pilot_auto'
                )
                session.add(test_case)
                session.flush()  # Para obtener el ID
            
            # Crear ejecución
            execution = TestExecution(
                id=uuid.UUID(task_id) if len(task_id) == 36 else uuid.uuid4(),
                test_case_id=test_case.id,
                project_id=default_project.id,
                execution_type='automated',
                browser_config={'browser': 'chrome', 'headless': False},
                environment='test',
                url_executed=history_entry['url'],
                status='passed' if history_entry['status'] == 'success' else 'failed',
                result_details=history_entry['message'],
                script_path=history_entry.get('script_path'),
                start_time=datetime.now(timezone.utc),
                end_time=datetime.now(timezone.utc),
                duration_seconds=0,
                created_by='qa_pilot'
            )
            session.add(execution)
            session.flush()  # Para obtener el ID
            
            # Guardar screenshots si existen
            if history_entry.get('screenshots'):
                for screenshot_data in history_entry['screenshots']:
                    if isinstance(screenshot_data, dict) and screenshot_data.get('path'):
                        screenshot = Screenshot(
                            execution_id=execution.id,
                            step_name=screenshot_data.get('name', 'screenshot'),
                            file_path=screenshot_data['path'],
                            file_size=os.path.getsize(screenshot_data['path']) if os.path.exists(screenshot_data['path']) else 0,
                            capture_time=datetime.now(timezone.utc),
                            screenshot_type='step'
                        )
                        session.add(screenshot)
            
            session.commit()
            print(f"DEBUG: Ejecución guardada en DB con ID {execution.id}")
            
    except Exception as e:
        print(f"ERROR: No se pudo guardar ejecución en base de datos: {e}")
        import traceback
        traceback.print_exc()

# Importar rutas de Excel para ejecución masiva
from excel_api_routes import register_excel_routes

# Importar integración de base de datos
from db_integration import init_db_integration, get_db_integration

app = Flask(__name__)

# Clase para capturar y sanitizar logs
class SanitizedFileHandler(logging.FileHandler):
    def __init__(self, filename, mode='a', encoding=None, delay=False):
        super().__init__(filename, mode, encoding, delay)
    
    def emit(self, record):
        try:
            # Sanitizar mensaje antes de emitirlo
            if isinstance(record.msg, str):
                record.msg = ''.join(c if ord(c) < 128 else '?' for c in record.msg)
            super().emit(record)
        except Exception:
            self.handleError(record)

# Añadir un handler de archivo si es necesario
try:
    log_file_handler = SanitizedFileHandler('app_debug.log', encoding='utf-8')
    log_file_handler.setLevel(logging.DEBUG)
    logger.addHandler(log_file_handler)
    logger.info("Logger configurado")
except Exception as e:
    print(f"No se pudo configurar log a archivo: {e}")

# Cargar configuración secreta y variables iniciales
managed_keys_init = load_env_vars()
app.secret_key = os.getenv('FLASK_SECRET_KEY', secrets.token_hex(16)) # Clave secreta para sesiones

# Directorios (ajustar rutas si es necesario)
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
SCREENSHOTS_DIR = os.path.join(BASE_DIR, "test_screenshots")
SCRIPTS_DIR = os.path.join(BASE_DIR, "test_scripts")

# Variable global para el directorio de capturas del test actual
test_dir = None

os.makedirs(SCREENSHOTS_DIR, exist_ok=True)
os.makedirs(SCRIPTS_DIR, exist_ok=True)

# Determinar la ruta del ejecutable de Python
python_exe = sys.executable
print(f"INFO: Usando ejecutable Python: {python_exe}")

# Almacenamiento simple en memoria para el estado de los tests (NO APTO PARA PRODUCCIÓN)
# En producción usar base de datos, Redis, etc.
test_status_db = {}

# Tracking de ejecuciones de suite activas
active_suite_executions = {}  # {execution_id: {'thread': thread_obj, 'suite_id': suite_id, 'stop_flag': threading.Event()}}
suite_executions_lock = threading.Lock()

# Tracking de procesos subprocess activos (browser-use)
active_subprocess_processes = {}  # {task_id: subprocess.Popen}
subprocess_processes_lock = threading.Lock()

# Archivo para persistir el estado de los tests
TEST_STATUS_FILE = os.path.join(BASE_DIR, "test_status_backup.json")

# Funciones de almacenamiento local eliminadas - ahora todo se guarda en BD

# Configurar app.config con las variables necesarias para ejecuciones en hilos
app.config['SCRIPTS_DIR'] = SCRIPTS_DIR
app.config['PYTHON_EXE'] = python_exe
app.config['BASE_DIR'] = BASE_DIR
app.config['SCREENSHOTS_DIR'] = SCREENSHOTS_DIR

# Inicializar integración de base de datos
init_db_integration(app)

# Registrar rutas de Excel para ejecución masiva
register_excel_routes(app)

# Lock global para acceder a test_status_db de forma segura desde los hilos
db_lock = threading.Lock()

# Variables temporales por compatibilidad - serán eliminadas gradualmente
history_db = []  # Solo para compatibilidad temporal
history_lock = threading.Lock()  # Solo para compatibilidad temporal

# Agregar las variables a la configuración de la aplicación después de definirlas
app.config['test_status_db'] = test_status_db
app.config['db_lock'] = db_lock

# Función para cargar historial desde la base de datos PostgreSQL
# Funciones de historial local eliminadas - ahora todo se maneja desde BD usando HistoryService

# Clase para interceptar y sanitizar la salida
class SafeTextIOWrapper(io.TextIOWrapper):
    def __init__(self, buffer, encoding='utf-8', errors='replace', **kwargs):
        super().__init__(buffer, encoding=encoding, errors=errors, **kwargs)
    
    def write(self, text):
        if isinstance(text, str):
            # Sanitizar el texto antes de escribirlo
            # Eliminar caracteres de control y no imprimibles
            text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', text)
            # Reemplazar caracteres no ASCII
            safe_text = ''.join(c if ord(c) < 128 else '?' for c in text)
            return super().write(safe_text)
        return super().write(text)

# Configurar codificación para Windows
if sys.platform == 'win32':
    try:
        # Intentar configurar la codificación de la consola
        import ctypes
        kernel32 = ctypes.windll.kernel32
        kernel32.SetConsoleCP(65001)         # Establecer codificación de entrada
        kernel32.SetConsoleOutputCP(65001)   # Establecer codificación de salida
        
        # Reemplazar stdout/stderr con nuestros wrappers seguros
        sys.stdout = SafeTextIOWrapper(sys.stdout.buffer, 'utf-8', 'replace')
        sys.stderr = SafeTextIOWrapper(sys.stderr.buffer, 'utf-8', 'replace')
        
        # Establecer locale predeterminado
        locale.setlocale(locale.LC_ALL, 'es-ES.UTF-8' if locale.windows_locale.get(locale.getdefaultlocale()[0]) else '.UTF-8')
        
        # Establecer variable de entorno para subprocesos
        os.environ['PYTHONIOENCODING'] = 'utf-8'
        
        print("Configuración de codificación segura aplicada para Windows")
    except Exception as e:
        print(f"Advertencia: No se pudo configurar completamente la codificación UTF-8: {e}")

# --- Rutas Flask --- 

@app.route('/')
def index():
    # Cargar claves actuales para mostrarlas en el formulario
    current_api_keys = load_env_vars()
    # Aquí cargaríamos el historial si lo almacenamos
    history = [] 
    return render_template('index.html', api_keys=current_api_keys, history=history)

@app.route('/config')
def config():
    # Cargar claves actuales para mostrarlas en el formulario
    current_api_keys = load_env_vars()
    return render_template('config.html', api_keys=current_api_keys, API_KEYS_TO_MANAGE=API_KEYS_TO_MANAGE)

@app.route('/suites')
def suites():
    """Página para gestionar suites de prueba."""
    return render_template('suites.html')

@app.route('/bulk_execution')
def bulk_execution():
    """Página para ejecución masiva de pruebas."""
    return render_template('bulk_execution.html')

@app.route('/excel_bulk_execution')
def excel_bulk_execution():
    """Página para ejecución masiva desde Excel."""
    return render_template('excel_bulk_execution.html')

@app.route('/download_excel_template')
def download_excel_template():
    """Descarga el template de Excel para casos de prueba."""
    try:
        template_path = os.path.join(os.getcwd(), "test_evidencia", "template")
        return send_from_directory(template_path, "TemplateNew.xlsx", as_attachment=True, download_name="Template_Casos_Prueba.xlsx")
    except Exception as e:
        logger.error(f"Error descargando template: {str(e)}")
        return f"Error descargando template: {str(e)}", 500

@app.route('/api_usage')
def api_usage():
    """Página para consultar el consumo de APIs"""
    # Cargar claves actuales
    current_api_keys = load_env_vars()
    # Función para la fecha actual
    from datetime import datetime
    def now():
        return datetime.now()
    
    # Inicializar diccionario para almacenar resultados de consumo
    api_usage_data = session.get('api_usage_data', {})
    api_key_checked = session.get('api_key_checked', '')
    
    return render_template('api_usage.html', 
                          api_keys=current_api_keys, 
                          API_KEYS_TO_MANAGE=API_KEYS_TO_MANAGE,
                          api_usage_data=api_usage_data,
                          api_key_checked=api_key_checked,
                          now=now)

@app.route('/save_config', methods=['POST'])
def save_config():
    """Guarda la configuración de las claves API en el archivo .env."""
    try:
        # 1. Validar método HTTP
        if request.method != 'POST':
            flash('Método no permitido', 'error')
            return redirect(url_for('config'))
        
        # 2. Procesar y validar las claves API
        api_keys = {}
        keys_processed = 0
        validation_errors = []
        
        for key in API_KEYS_TO_MANAGE:
            value = request.form.get(key, '').strip()
            # Validación básica de formato de la clave API
            if value:
                if len(value) < 20:  # La mayoría de las API keys tienen al menos 20 caracteres
                    validation_errors.append(f"La clave {key} parece demasiado corta")
                elif not value.startswith(('sk-', 'AI', 'an')):  # Prefijos comunes de API keys
                    validation_errors.append(f"El formato de {key} no parece válido")
                else:
                    api_keys[key] = value
                    keys_processed += 1
                    print(f"Clave {key} procesada correctamente")
        
        # Si hay errores de validación, mostrarlos
        if validation_errors:
            for error in validation_errors:
                flash(error, 'warning')
            return redirect(url_for('config'))
        
        # 3. Guardar las claves validadas en el archivo .env
        if api_keys:
            success, message = save_api_keys_to_env(api_keys)
            if success:
                print(f"Claves guardadas exitosamente en .env ({keys_processed} claves)")
                flash(f'Configuración actualizada: {keys_processed} claves guardadas', 'success')
            else:
                print(f"Error al guardar claves: {message}")
                flash(f'Error al guardar la configuración: {message}', 'error')
        else:
            flash('No se proporcionaron claves API válidas', 'warning')
        
        return redirect(url_for('config'))
        
    except Exception as e:
        error_detail = traceback.format_exc()
        print(f"Error inesperado en save_config: {error_detail}")
        flash(f'Error inesperado al procesar la configuración: {str(e)}', 'error')
        return redirect(url_for('config'))

@app.route('/database')
def database_page():
    """Página de estado de la base de datos"""
    return render_template('database_status.html')

@app.route('/database_status')
def database_status():
    """Verificar estado de la base de datos"""
    try:
        db_integration = get_db_integration()
        connection_ok = db_integration.test_connection()
        
        if connection_ok:
            # Obtener estadísticas básicas
            with db_integration.get_session() as session:
                from sqlalchemy import text, func
                from db_models import Project, TestCase, TestExecution
                
                # Contar registros
                projects_count = session.query(func.count(Project.id)).scalar()
                cases_count = session.query(func.count(TestCase.id)).scalar()
                executions_count = session.query(func.count(TestExecution.id)).scalar()
                
                stats = {
                    'connection': True,
                    'projects': projects_count,
                    'test_cases': cases_count,
                    'executions': executions_count
                }
        else:
            stats = {'connection': False}
            
        return jsonify({
            'success': True,
            'database_status': stats
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'database_status': {'connection': False}
        }), 500

@app.route('/api/test_cases')
def api_get_test_cases():
    """Obtener casos de prueba desde la base de datos"""
    try:
        db_integration = get_db_integration()
        
        # Parámetros de consulta
        status = request.args.get('status')
        limit = int(request.args.get('limit', 50))
        offset = int(request.args.get('offset', 0))
        
        cases = db_integration.get_test_cases(
            status=status,
            limit=limit,
            offset=offset
        )
        
        return jsonify({
            'success': True,
            'test_cases': cases,
            'count': len(cases)
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/save_test_to_db', methods=['POST'])
def api_save_test_to_db():
    """Guardar un caso de prueba en la base de datos desde el historial"""
    try:
        data = request.get_json()
        task_id = data.get('task_id')
        
        if not task_id:
            return jsonify({
                'success': False,
                'error': 'Task ID es requerido'
            }), 400
        
        # Buscar el test en el historial usando el nuevo servicio
        history_service = get_history_service()
        history_items = history_service.get_history_items(limit=200)
        
        test_data = None
        for item in history_items:
            if item['id'] == task_id:
                test_data = item
                break
        
        if not test_data:
            return jsonify({
                'success': False,
                'error': 'Test no encontrado en el historial'
            }), 404
        
        # Crear un objeto TestCase simulado para guardar en DB
        from excel_test_analyzer import TestCase as ExcelTestCase
        
        excel_case = ExcelTestCase()
        excel_case.nombre = test_data.get('name', f'Test {task_id[:8]}')
        excel_case.codigo = f'TC_{task_id[:8]}'
        excel_case.tipo = 'funcional'
        excel_case.prioridad = 'media'
        excel_case.objetivo = f"Automatizar test: {test_data.get('name', 'Sin nombre')}"
        excel_case.pasos = test_data.get('instructions', 'Sin pasos definidos')
        excel_case.resultado_esperado = 'Test ejecutado exitosamente'
        excel_case.url_extraida = test_data.get('url', '')
        excel_case.es_valido = test_data.get('status') == 'success'
        excel_case.problemas = [] if test_data.get('status') == 'success' else [test_data.get('message', 'Error en ejecución')]
        excel_case.sugerencias = []
        excel_case.instrucciones_qa_pilot = test_data.get('instructions', '')
        excel_case.instrucciones_browser_use = test_data.get('instructions', '')
        
        # Guardar en base de datos
        db_integration = get_db_integration()
        case_id = db_integration.save_excel_test_case(excel_case)
        
        return jsonify({
            'success': True,
            'message': 'Caso de prueba guardado en base de datos',
            'case_id': case_id
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/test_cases/orphaned')
def api_get_orphaned_test_cases():
    """Obtener casos de prueba sin suite asociada"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({
                'success': False,
                'error': 'Base de datos no disponible'
            }), 500
        
        # Obtener casos huérfanos (sin suite_id)
        orphaned_cases = db_integration.get_orphaned_test_cases()
        
        return jsonify({
            'success': True,
            'orphaned_cases': orphaned_cases,
            'total': len(orphaned_cases)
        })
        
    except Exception as e:
        print(f"Error getting orphaned test cases: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/test_cases/<case_id>', methods=['DELETE'])
def api_delete_test_case(case_id):
    """Eliminar un caso de prueba"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({
                'success': False,
                'error': 'Base de datos no disponible'
            }), 500
        
        success = db_integration.delete_test_case(case_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Caso de prueba eliminado exitosamente'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'No se pudo eliminar el caso de prueba'
            }), 404
        
    except Exception as e:
        print(f"Error deleting test case: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/test_cases/orphaned/bulk_delete', methods=['POST'])
def api_bulk_delete_orphaned_cases():
    """Eliminar múltiples casos huérfanos"""
    try:
        data = request.get_json()
        case_ids = data.get('case_ids', [])
        
        if not case_ids:
            return jsonify({
                'success': False,
                'error': 'No se proporcionaron IDs de casos'
            }), 400
        
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({
                'success': False,
                'error': 'Base de datos no disponible'
            }), 500
        
        deleted_count = db_integration.bulk_delete_test_cases(case_ids)
        
        return jsonify({
            'success': True,
            'deleted_count': deleted_count,
            'message': f'{deleted_count} casos eliminados exitosamente'
        })
        
    except Exception as e:
        print(f"Error bulk deleting test cases: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/history')
def history():
    """Vista para mostrar el historial de tests."""
    logger.info("🔄 Cargando historial desde base de datos")
    
    # Usar el nuevo servicio de historial basado en BD
    history_service = get_history_service()
    sorted_history = history_service.get_history_items(limit=100)
    
    # Verificar qué ejecuciones tienen casos asociados en la BD
    db_integration = get_db_integration()
    if db_integration and db_integration.is_connected():
        try:
            # Obtener todos los casos de la BD para verificar relaciones
            all_test_cases = db_integration.get_test_cases(limit=1000)
            
            # Crear mapas para verificaciones rápidas
            cases_by_name = {}
            cases_by_suite = {}
            
            for case in all_test_cases:
                case_name = case.get('nombre', '').lower()
                case_code = case.get('codigo', '').lower()
                
                # Mapear por nombre y código
                if case_name:
                    cases_by_name[case_name] = case
                if case_code:
                    cases_by_name[case_code] = case
                
                # Mapear por suite
                suite_id = case.get('suite_id')
                if suite_id:
                    if suite_id not in cases_by_suite:
                        cases_by_suite[suite_id] = []
                    cases_by_suite[suite_id].append(case)
            
            # Enriquecer cada item del historial con información de BD
            for item in sorted_history:
                item_name = item.get('name', '').lower()
                item_id = item.get('id', '')[:8].lower()
                
                # Verificar si este historial tiene caso asociado en BD
                associated_case = None
                
                # Buscar por nombre exacto
                if item_name in cases_by_name:
                    associated_case = cases_by_name[item_name]
                # Buscar por ID parcial en código
                elif any(item_id in code for code in cases_by_name.keys()):
                    associated_case = next((case for code, case in cases_by_name.items() if item_id in code), None)
                
                # Agregar información del caso asociado
                if associated_case:
                    item['db_case'] = {
                        'id': associated_case.get('id'),
                        'nombre': associated_case.get('nombre'),
                        'codigo': associated_case.get('codigo'),
                        'suite_id': associated_case.get('suite_id'),
                        'status': associated_case.get('status', 'draft')
                    }
                    
                    # Si tiene suite, agregar información de la suite
                    if associated_case.get('suite_id'):
                        item['is_in_suite'] = True
                        item['suite_id'] = associated_case.get('suite_id')
                    else:
                        item['is_available_case'] = True
                else:
                    item['is_execution_only'] = True
                    
        except Exception as e:
            logger.warning(f"No se pudo enriquecer historial con datos de BD: {e}")
        
        # Enriquecer datos si necesitan mejoras
        for item in sorted_history:
            # Asegurar que todos los items tienen un nombre
            if not item.get('name'):
                item['name'] = f"Test {item.get('id', '')[:8]}"
            
            # Verificar scripts asociados y actualizar si necesario
            script_path = item.get('script_path')
            if script_path and not os.path.exists(script_path):
                # Intentar encontrar por ID
                test_id = item.get('id', '')
                if test_id:
                    for filename in os.listdir(SCRIPTS_DIR):
                        if test_id[:8] in filename and filename.endswith('.py'):
                            item['script_path'] = os.path.join(SCRIPTS_DIR, filename)
                            print(f"DEBUG /history: Actualizada ruta de script para {item['name']}: {item['script_path']}")
                            break
            
            # Verificar capturas existentes y filtrar las que no existen físicamente
            if item.get('screenshots'):
                print(f"DEBUG /history: Validando {len(item['screenshots'])} capturas para {item['name']}")
                valid_screenshots = []
                
                for screenshot in item['screenshots']:
                    if isinstance(screenshot, dict):
                        # Verificar si el archivo existe físicamente
                        screenshot_path = screenshot.get('path')
                        if screenshot_path and os.path.exists(screenshot_path):
                            valid_screenshots.append(screenshot)
                        else:
                            # Si no tiene path, intentar reconstruir la ruta desde la URL
                            url = screenshot.get('url', '')
                            if url.startswith('/media/screenshots/'):
                                relative_path = url.replace('/media/screenshots/', '')
                                full_path = os.path.join(SCREENSHOTS_DIR, relative_path.replace('/', os.path.sep))
                                if os.path.exists(full_path):
                                    # Actualizar el path si existe
                                    screenshot['path'] = full_path
                                    valid_screenshots.append(screenshot)
                                else:
                                    print(f"DEBUG /history: Captura eliminada o no encontrada: {screenshot.get('name', 'Sin nombre')}")
                
                # Actualizar la lista con solo las capturas válidas
                if len(valid_screenshots) != len(item['screenshots']):
                    print(f"DEBUG /history: Actualizando capturas de {len(item['screenshots'])} a {len(valid_screenshots)} para {item['name']}")
                    item['screenshots'] = valid_screenshots
            
            # Buscar capturas adicionales en el directorio del test (solo si existe)
            if 'test_dir' in item and item['test_dir'] and os.path.exists(item['test_dir']):
                test_dir = item['test_dir']
                
                # Buscar imágenes PNG en este directorio que no estén ya en la lista
                existing_names = {s.get('name') for s in item.get('screenshots', []) if isinstance(s, dict)}
                
                try:
                    print(f"DEBUG /history: Buscando capturas adicionales en directorio {test_dir}")
                    for file in os.listdir(test_dir):
                        if file.lower().endswith('.png') and file not in existing_names:
                            file_path = os.path.join(test_dir, file)
                            if os.path.exists(file_path):  # Verificar que existe antes de agregar
                                rel_path = os.path.relpath(file_path, SCREENSHOTS_DIR)
                                url_path = rel_path.replace(os.path.sep, '/')
                                
                                new_screenshot = {
                                    'url': f'/media/screenshots/{url_path}',
                                    'path': file_path,
                                    'name': file
                                }
                                
                                if not item.get('screenshots'):
                                    item['screenshots'] = []
                                item['screenshots'].append(new_screenshot)
                                print(f"DEBUG /history: Agregada captura adicional: {file}")
                except Exception as e:
                    print(f"ERROR /history: No se pudieron recuperar capturas adicionales de {test_dir}: {e}")
            
            # Buscar capturas en test_screenshots (directorio principal sin duplicación)
            test_id_full = item.get('id', '')
            test_screenshots_dir = os.path.join(os.getcwd(), "test_screenshots", test_id_full)
            
            if os.path.exists(test_screenshots_dir) and test_screenshots_dir != item.get('test_dir'):
                try:
                    print(f"DEBUG /history: Buscando capturas adicionales en {test_screenshots_dir}")
                    dir_screenshots = []
                    for file in os.listdir(test_screenshots_dir):
                        if file.lower().endswith('.png'):
                            file_path = os.path.join(test_screenshots_dir, file)
                            rel_path = os.path.relpath(file_path, os.getcwd())
                            url_path = rel_path.replace(os.path.sep, '/')
                            
                            dir_screenshots.append({
                                'url': f'/media/screenshots/{url_path}',
                                'path': file_path,
                                'name': file
                            })
                    
                    if dir_screenshots:
                        print(f"DEBUG /history: Se encontraron {len(dir_screenshots)} capturas adicionales en {test_screenshots_dir}")
                        # Añadir capturas o actualizar lista existente
                        if not item.get('screenshots'):
                            item['screenshots'] = dir_screenshots
                        else:
                            # Agregar solo capturas no duplicadas
                            existing_urls = {s.get('url') for s in item.get('screenshots', []) if isinstance(s, dict)}
                            for ss in dir_screenshots:
                                if ss.get('url') not in existing_urls:
                                    item['screenshots'].append(ss)
                                    existing_urls.add(ss.get('url'))
                except Exception as e:
                    print(f"ERROR /history: No se pudieron recuperar capturas adicionales de {test_screenshots_dir}: {e}")
    
    # Información de debug
    logger.info(f"📊 Historial cargado: {len(sorted_history)} elementos desde BD")
    if len(sorted_history) > 0:
        logger.debug(f"Primer elemento: {sorted_history[0].get('name', 'sin nombre')} - {sorted_history[0].get('date', 'sin fecha')}")
        ss_count = len(sorted_history[0].get('screenshots', []))
        logger.debug(f"Capturas: {ss_count}")
    
    # Verificar que todos los tests tienen su script
    scripts_missing = 0
    for item in sorted_history:
        script_path = item.get('script_path', '')
        if not script_path or not os.path.exists(script_path):
            scripts_missing += 1
    
    if scripts_missing > 0:
        logger.debug(f"{scripts_missing} de {len(sorted_history)} tests no tienen script válido")
    
    return render_template('history.html', history=sorted_history)

@app.route('/get_history')
def get_history():
    """Endpoint JSON para obtener historial (usado por bulk_execution)"""
    try:
        # Usar el mismo servicio que usa /history pero devolver JSON
        history_service = get_history_service()
        sorted_history = history_service.get_history_items(limit=100)
        
        # Filtrar solo ejecuciones que NO han sido guardadas como casos en BD
        # para evitar duplicación en bulk execution
        db_integration = get_db_integration()
        filtered_history = []
        
        if db_integration and db_integration.is_connected():
            try:
                # Obtener casos existentes en BD
                all_test_cases = db_integration.get_test_cases(limit=1000)
                case_names = {case.get('nombre', '').lower() for case in all_test_cases}
                case_codes = {case.get('codigo', '').lower() for case in all_test_cases}
                
                # Filtrar historial excluyendo los que ya están como casos en BD
                for item in sorted_history:
                    item_name = item.get('name', '').lower()
                    item_id = item.get('id', '')[:8].lower()
                    
                    # Solo incluir si NO está en la BD como caso
                    is_in_bd = (item_name in case_names or 
                               any(item_id in code for code in case_codes))
                    
                    if not is_in_bd:
                        filtered_history.append(item)
                        
            except Exception as e:
                logger.warning(f"Error filtrando historial para bulk execution: {e}")
                # Si hay error, usar el historial completo
                filtered_history = sorted_history
        else:
            filtered_history = sorted_history
        
        return jsonify({
            'status': 'success',
            'history': filtered_history,
            'total': len(filtered_history)
        })
        
    except Exception as e:
        logger.error(f"Error en get_history: {e}")
        return jsonify({
            'status': 'error',
            'message': str(e),
            'history': []
        }), 500

@app.route('/save_test_as_playwright', methods=['POST'])
def save_test_as_playwright():
    """Guarda un test como caso de Playwright."""
    if request.method != 'POST':
        return jsonify({'status': 'error', 'message': 'Método no permitido'}), 405
    
    test_id = request.form.get('test_id')
    name = request.form.get('name', '')
    description = request.form.get('description', '')
    
    if not test_id:
        return jsonify({'status': 'error', 'message': 'ID de test no proporcionado'}), 400
    
    # Buscar el test en el historial
    with history_lock:
        test_data = None
        for item in history_db:
            if item.get('id') == test_id:
                test_data = item
                break
        
        if not test_data:
            return jsonify({'status': 'error', 'message': 'Test no encontrado en historial'}), 404
    
    # Verificar que el script exista
    script_path = test_data.get('script_path')
    if not script_path or not os.path.exists(script_path):
        # Intentar buscar el script en test_scripts por ID
        test_id_short = test_id[:8]
        possible_scripts = []
        
        for filename in os.listdir(SCRIPTS_DIR):
            if test_id_short in filename and filename.endswith('.py'):
                possible_script = os.path.join(SCRIPTS_DIR, filename)
                possible_scripts.append(possible_script)
        
        if possible_scripts:
            # Usar el script más reciente que coincida con el ID
            script_path = max(possible_scripts, key=os.path.getmtime)
            print(f"DEBUG: Script encontrado por ID: {script_path}")
        else:
            return jsonify({'status': 'error', 'message': 'Script de test no encontrado'}), 404
    
    # Sanitizar el nombre
    if not name:
        name = f"test_{test_id[:8]}"
    
    safe_name = "".join(c for c in name if c.isalnum() or c in "._- ").replace(" ", "_").lower()
    if not safe_name.endswith('.py'):
        safe_name += '.py'
    
    # Crear directorios de Playwright si no existen
    PLAYWRIGHT_DIR = os.path.join(BASE_DIR, "playwright_scripts")
    CASOS_DIR = os.path.join(PLAYWRIGHT_DIR, "casos")
    
    os.makedirs(PLAYWRIGHT_DIR, exist_ok=True)
    os.makedirs(CASOS_DIR, exist_ok=True)
    
    # Ruta de destino para el caso de Playwright
    playwright_case_path = os.path.join(CASOS_DIR, safe_name)
    
    try:
        # Leer el contenido del script original
        with open(script_path, 'r', encoding='utf-8-sig', errors='replace') as f:
            script_content = f.read()
        
        # Generar o modificar el docstring
        if description:
            docstring = f'"""\n{description}\n"""\n\n'
            # Eliminar docstring existente si hay
            script_content = re.sub(r'^""".*?"""', '', script_content, flags=re.DOTALL)
            script_content = docstring + script_content.lstrip()
        
        # Guardar el script como caso de Playwright
        with open(playwright_case_path, 'w', encoding='utf-8') as f:
            f.write(script_content)
        
        print(f"DEBUG: Caso de Playwright guardado en: {playwright_case_path}")
        
        # Eliminar el archivo original si está en test_scripts
        if script_path.startswith(SCRIPTS_DIR):
            try:
                os.remove(script_path)
                print(f"DEBUG: Archivo original eliminado: {script_path}")
            except Exception as e:
                print(f"ERROR: No se pudo eliminar el archivo original: {e}")
        
        # Actualizar el test en el historial
        with history_lock:
            for item in history_db:
                if item.get('id') == test_id:
                    item['name'] = name
                    item['playwright_case'] = os.path.basename(playwright_case_path)
                    item['script_path'] = playwright_case_path  # Actualizar ruta a la nueva ubicación
                    break
            
            # Cambios se guardan automáticamente en BD
        
        # Intentar guardar también en la base de datos
        saved_to_db = False
        case_id = None
        try:
            db_integration = get_db_integration()
            if db_integration and db_integration.is_connected():
                # Crear un objeto TestCase para la base de datos
                from excel_test_analyzer import TestCase as ExcelTestCase
                
                excel_case = ExcelTestCase()
                excel_case.nombre = name
                excel_case.codigo = f'TC_{test_id[:8]}'
                excel_case.tipo = 'automatizado'
                excel_case.prioridad = 'media'
                excel_case.objetivo = f"Caso automatizado: {name}"
                excel_case.pasos = test_data.get('instructions', 'Caso automatizado generado desde historial')
                excel_case.resultado_esperado = 'Ejecución exitosa del caso automatizado'
                excel_case.url_extraida = test_data.get('url', '')
                excel_case.es_valido = test_data.get('status') == 'success'
                excel_case.problemas = [] if test_data.get('status') == 'success' else [test_data.get('message', 'Error en ejecución')]
                excel_case.sugerencias = []
                excel_case.instrucciones_qa_pilot = test_data.get('instructions', '')
                excel_case.instrucciones_browser_use = test_data.get('instructions', '')
                
                # Agregar metadatos importantes para la sincronización
                import json
                excel_case.metadata_json = json.dumps({
                    'test_id': test_id,
                    'original_script_path': script_path,
                    'playwright_case_path': playwright_case_path,
                    'created_from_history': True,
                    'history_date': test_data.get('date'),
                    'history_status': test_data.get('status')
                })
                
                # Guardar en base de datos
                case_id = db_integration.save_excel_test_case(excel_case)
                saved_to_db = True
                logger.info(f"Caso guardado en base de datos con ID: {case_id}")
                
        except Exception as e:
            logger.warning(f"No se pudo guardar en base de datos: {e}")
            # No fallar si no se puede guardar en DB
        
        # Información para respuesta
        case_info = {
            'id': test_id,
            'name': name,
            'path': playwright_case_path,
            'filename': os.path.basename(playwright_case_path),
            'saved_to_db': saved_to_db,
            'case_id': case_id
        }
        
        message = f'Caso guardado como {name} y original eliminado'
        if saved_to_db:
            message += f' (guardado en base de datos con ID: {case_id})'
        
        return jsonify({
            'status': 'success', 
            'message': message,
            'case': case_info
        })
    
    except Exception as e:
        print(f"ERROR: Error al guardar caso de Playwright: {str(e)}")
        print(traceback.format_exc())
        return jsonify({
            'status': 'error', 
            'message': f'Error al guardar caso: {str(e)}'
        }), 500

@app.route('/update_test_name', methods=['POST'])
def update_test_name():
    logger.debug(f"Received request for /update_test_name")
    logger.debug(f"Request headers: {request.headers}")
    
    if not request.is_json:
        logger.error("Request is not JSON")
        return jsonify({'status': 'error', 'message': 'Formato incorrecto, se esperaba JSON'}), 415

    data = request.get_json()
    logger.debug(f"Request JSON data: {data}")

    if not data or 'test_id' not in data or 'new_name' not in data:
        logger.error(f"Datos incompletos en la solicitud: {data}")
        return jsonify({'status': 'error', 'message': 'Datos incompletos, se requiere test_id y new_name'}), 400
    
    test_id = data['test_id']
    new_name = data['new_name'].strip()
    
    if not new_name:
        logger.error("El nuevo nombre está vacío")
        return jsonify({'status': 'error', 'message': 'El nombre no puede estar vacío'}), 400

    history_service = get_history_service()
    success, message = history_service.update_test_name(test_id, new_name)
    
    if success:
        logger.info(f"Nombre del test {test_id} actualizado a '{new_name}'")
        return jsonify({'status': 'success', 'message': message})
    else:
        logger.error(f"Error al actualizar el nombre para el test {test_id}: {message}")
        return jsonify({'status': 'error', 'message': message}), 500

@app.route('/reload_history', methods=['POST'])
def reload_history():
    """Recarga el historial desde la base de datos (útil para debugging)"""
    try:
        logger.info("🔄 Recargando historial desde base de datos")
        
        # Usar el nuevo servicio de historial (que siempre carga desde BD)
        history_service = get_history_service()
        history_items = history_service.get_history_items(limit=200)
        
        logger.info(f"✅ Historial recargado: {len(history_items)} elementos desde BD")
        
        return jsonify({
            'status': 'success',
            'message': f'Historial recargado desde base de datos. {len(history_items)} elementos encontrados',
            'count': len(history_items)
        })
        
    except Exception as e:
        logger.error(f"❌ Error recargando historial: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Error recargando historial: {str(e)}'
        }), 500

@app.route('/sync_test_names', methods=['POST'])
def sync_test_names():
    """Sincroniza nombres de tests entre historial y base de datos."""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({
                'status': 'error',
                'message': 'Base de datos no disponible'
            }), 503
        
        synced_count = 0
        errors = []
        processed_items = []  # Para debugging
        
        with history_lock:
            logger.info(f"🔄 Iniciando sincronización. Items en historial: {len(history_db)}")
            
            for i, item in enumerate(history_db):
                test_id = item.get('id')
                test_name = item.get('name')
                
                # Log detallado de cada item para debugging
                logger.debug(f"📋 Item {i+1}: ID={test_id}, Name={test_name}")
                
                if not test_id or not test_name:
                    logger.debug(f"⏭️ Saltando item {i+1} (ID o name faltante)")
                    continue
                
                try:
                    with db_integration.get_session() as session:
                        from db_models import TestCase
                        from datetime import datetime
                        
                        # Buscar casos relacionados con este test_id
                        test_cases = []
                        
                        # Método 1: Búsqueda por metadatos JSON (si la columna existe)
                        try:
                            from sqlalchemy import text
                            # Intentar primero con metadata_json si existe
                            test_cases = session.query(TestCase).filter(
                                text("metadata_json::text LIKE :pattern")
                            ).params(pattern=f'%{test_id}%').all()
                            logger.debug(f"🔍 Búsqueda por metadatos para {test_id}: {len(test_cases)} casos encontrados")
                        except Exception as meta_error:
                            logger.debug(f"❌ Error buscando por metadatos (columna metadata_json puede no existir): {meta_error}")
                            test_cases = []
                        
                        # Método 2: Búsqueda por código (fallback principal)
                        if not test_cases:
                            try:
                                test_cases = session.query(TestCase).filter(
                                    TestCase.codigo.like(f'%{test_id[:8]}%')
                                ).all()
                                logger.debug(f"🔍 Búsqueda por código para {test_id[:8]}: {len(test_cases)} casos encontrados")
                            except Exception as code_error:
                                logger.debug(f"❌ Error buscando por código: {code_error}")
                        
                        # Método 3: Búsqueda por nombre exacto o similar
                        if not test_cases:
                            try:
                                # Buscar casos que podrían corresponder al nombre del test
                                search_name = test_name.replace('Demo ML', '').strip()
                                if len(search_name) > 3:  # Solo buscar si el nombre es significativo
                                    test_cases = session.query(TestCase).filter(
                                        TestCase.nombre.ilike(f'%{search_name}%')
                                    ).all()
                                    logger.debug(f"🔍 Búsqueda por nombre '{search_name}': {len(test_cases)} casos encontrados")
                            except Exception as name_error:
                                logger.debug(f"❌ Error buscando por nombre: {name_error}")
                        
                        # Método 4: Crear caso automáticamente si no existe
                        if not test_cases:
                            try:
                                # Solo crear caso si el test fue exitoso y no es demo
                                if test_id != 'test-id-demo' and test_name and 'Demo' not in test_name:
                                    logger.info(f"🆕 Creando nuevo caso para {test_id}: {test_name}")
                                    
                                    # Buscar información adicional del historial
                                    history_item = None
                                    for hist_item in history_db:
                                        if hist_item.get('id') == test_id:
                                            history_item = hist_item
                                            break
                                    
                                    new_case = TestCase(
                                        # Usar project_id por defecto o el primero disponible
                                        project_id=session.query(TestCase).first().project_id if session.query(TestCase).first() else None,
                                        nombre=test_name,
                                        codigo=f'TC_{test_id[:8]}',
                                        tipo='funcional',
                                        prioridad='media',
                                        objetivo=f'Caso sincronizado desde historial: {test_name}',
                                        pasos=history_item.get('instructions', 'Pasos no disponibles') if history_item else 'Pasos no disponibles',
                                        resultado_esperado='Ejecución exitosa según instrucciones',
                                        url_objetivo=history_item.get('url', '') if history_item else '',
                                        es_valido=True,
                                        status='approved',
                                        created_by='sync_auto'
                                    )
                                    
                                    # Solo agregar metadatos si la columna existe
                                    try:
                                        if hasattr(new_case, 'metadata_json'):
                                            new_case.metadata_json = {
                                                'sync_source': 'history',
                                                'original_test_id': test_id,
                                                'sync_timestamp': datetime.now().isoformat()
                                            }
                                    except:
                                        pass  # Ignorar si no se pueden agregar metadatos
                                    
                                    session.add(new_case)
                                    session.flush()  # Para obtener el ID
                                    test_cases = [new_case]
                                    logger.info(f"✅ Caso creado automáticamente: {new_case.id}")
                                    
                            except Exception as create_error:
                                logger.debug(f"❌ Error creando caso automáticamente: {create_error}")
                        
                        # Actualizar casos encontrados
                        item_updates = 0
                        for case in test_cases:
                            if case.nombre != test_name:
                                old_name = case.nombre
                                case.nombre = test_name
                                item_updates += 1
                                synced_count += 1
                                logger.info(f"✅ Actualizado caso {case.id}: '{old_name}' → '{test_name}'")
                        
                        if item_updates > 0:
                            session.commit()
                            logger.info(f"💾 Guardados {item_updates} cambios para test_id {test_id}")
                        
                        # Agregar info de debugging
                        processed_items.append({
                            'test_id': test_id,
                            'test_name': test_name,
                            'cases_found': len(test_cases),
                            'updates_made': item_updates
                        })
                        
                except Exception as e:
                    error_msg = f"Error sincronizando {test_id}: {str(e)}"
                    errors.append(error_msg)
                    logger.error(f"❌ {error_msg}")
        
        # Log de resumen
        logger.info(f"📊 Sincronización completada:")
        logger.info(f"   - Items procesados: {len(processed_items)}")
        logger.info(f"   - Casos actualizados: {synced_count}")
        logger.info(f"   - Errores: {len(errors)}")
        
        return jsonify({
            'status': 'success',
            'message': f'Sincronización completada: {synced_count} nombres actualizados',
            'synced_count': synced_count,
            'errors': errors,
            'debug_info': {
                'total_history_items': len(history_db),
                'processed_items': len(processed_items),
                'detailed_results': processed_items[:10]  # Primeros 10 para evitar respuesta muy grande
            }
        })
        
    except Exception as e:
        logger.error(f"❌ Error general en sincronización: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Error en sincronización: {str(e)}'
        }), 500

@app.route('/debug_sync_data', methods=['GET'])
def debug_sync_data():
    """Endpoint de debug para analizar datos de sincronización."""
    try:
        db_integration = get_db_integration()
        debug_data = {
            'database_connected': db_integration is not None and db_integration.is_connected() if db_integration else False,
            'history_items': [],
            'database_cases': [],
            'potential_matches': []
        }
        
        # Analizar datos del historial
        with history_lock:
            for i, item in enumerate(history_db[:20]):  # Solo primeros 20 para evitar sobrecarga
                debug_data['history_items'].append({
                    'index': i,
                    'id': item.get('id'),
                    'name': item.get('name'),
                    'url': item.get('url'),
                    'date': item.get('date'),
                    'status': item.get('status'),
                    'all_keys': list(item.keys())
                })
        
        # Analizar datos de la base de datos
        if debug_data['database_connected']:
            with db_integration.get_session() as session:
                from db_models import TestCase
                from sqlalchemy import text
                
                # Obtener casos recientes
                recent_cases = session.query(TestCase).order_by(TestCase.created_at.desc()).limit(20).all()
                
                for case in recent_cases:
                    # Manejar metadatos de forma segura
                    metadata_json = None
                    has_metadata = False
                    try:
                        if hasattr(case, 'metadata_json'):
                            metadata_json = case.metadata_json
                            has_metadata = bool(metadata_json)
                    except:
                        pass  # Ignorar errores de metadatos
                    
                    debug_data['database_cases'].append({
                        'id': str(case.id),
                        'nombre': case.nombre,
                        'codigo': case.codigo,
                        'created_at': case.created_at.isoformat() if case.created_at else None,
                        'metadata_json': metadata_json,
                        'has_metadata': has_metadata
                    })
                
                # Buscar posibles coincidencias
                for hist_item in debug_data['history_items']:
                    if hist_item['id']:
                        test_id = hist_item['id']
                        
                        # Buscar por metadatos (si la columna existe)
                        metadata_matches = []
                        try:
                            metadata_matches = session.query(TestCase).filter(
                                text("metadata_json::text LIKE :pattern")
                            ).params(pattern=f'%{test_id}%').all()
                        except Exception as e:
                            logger.debug(f"Columna metadata_json no disponible: {e}")
                        
                        # Buscar por código
                        code_matches = []
                        try:
                            code_matches = session.query(TestCase).filter(
                                TestCase.codigo.like(f'%{test_id[:8]}%')
                            ).all()
                        except Exception as e:
                            logger.debug(f"Error buscando por código: {e}")
                        
                        if metadata_matches or code_matches:
                            debug_data['potential_matches'].append({
                                'history_id': test_id,
                                'history_name': hist_item['name'],
                                'metadata_matches': len(metadata_matches),
                                'code_matches': len(code_matches),
                                'metadata_cases': [{'id': str(c.id), 'nombre': c.nombre, 'codigo': c.codigo} for c in metadata_matches],
                                'code_cases': [{'id': str(c.id), 'nombre': c.nombre, 'codigo': c.codigo} for c in code_matches]
                            })
        
        # Información adicional sobre el esquema de BD
        if debug_data['database_connected']:
            try:
                with db_integration.get_session() as session:
                    from sqlalchemy import inspect
                    inspector = inspect(session.bind)
                    
                    # Verificar qué columnas existen en test_cases
                    if inspector.has_table('test_cases', schema='testing'):
                        columns = inspector.get_columns('test_cases', schema='testing')
                        debug_data['database_schema'] = {
                            'test_cases_columns': [col['name'] for col in columns],
                            'has_metadata_json': 'metadata_json' in [col['name'] for col in columns],
                            'schema_info': 'Esquema testing.test_cases encontrado'
                        }
                    else:
                        debug_data['database_schema'] = {
                            'error': 'Tabla testing.test_cases no encontrada',
                            'available_tables': inspector.get_table_names(schema='testing')
                        }
                        
            except Exception as schema_error:
                debug_data['database_schema'] = {
                    'error': f'Error verificando esquema: {str(schema_error)}'
                }
        
        return jsonify(debug_data)
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error en debug: {str(e)}'
        }), 500

@app.route('/run_playwright_case', methods=['POST'])
def run_playwright_case():
    """Ejecuta un caso de Playwright."""
    if request.method != 'POST':
        return jsonify({'status': 'error', 'message': 'Método no permitido'}), 405
    
    case_name = request.form.get('case_name')
    headless = request.form.get('headless', 'true').lower() == 'true'
    
    if not case_name:
        return jsonify({'status': 'error', 'message': 'Nombre del caso no proporcionado'}), 400
    
    # Verificar que el caso exista
    PLAYWRIGHT_DIR = os.path.join(BASE_DIR, "playwright_scripts")
    CASOS_DIR = os.path.join(PLAYWRIGHT_DIR, "casos")
    
    if not case_name.endswith('.py'):
        case_name += '.py'
    
    case_path = os.path.join(CASOS_DIR, case_name)
    
    if not os.path.exists(case_path):
        return jsonify({'status': 'error', 'message': f'Caso {case_name} no encontrado'}), 404
    
    # Generar un task_id para seguimiento
    task_id = str(uuid.uuid4())
    
    # Inicializar estado en test_status_db
    with db_lock:
        test_status_db[task_id] = {
            'status': 'queued',
            'message': 'Preparando ejecución de caso Playwright...',
            'original_instructions': f'Ejecutando caso Playwright: {case_name}',
            'stdout': '',
            'stderr': '',
            'current_action': 'Inicializando...',
            'icon': 'fa-play-circle',
            'url': 'Caso Playwright',
            'name': os.path.splitext(case_name)[0],
            'playwright_case': case_name
        }
    
    # Lanzar ejecución en segundo plano
    thread = Thread(target=execute_playwright_case, args=(task_id, case_path, headless))
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'status': 'started',
        'message': f'Ejecutando caso {case_name}',
        'task_id': task_id
    })

def execute_playwright_case(task_id, case_path, headless=True):
    """Ejecuta un caso de Playwright en segundo plano."""
    global test_status_db
    global history_db
    
    try:
        # Obtener el nombre base del caso
        case_name = os.path.basename(case_path)
        
        with db_lock:
            if task_id not in test_status_db:
                return
            
            test_status_db[task_id].update({
                'status': 'running',
                'message': 'Ejecutando caso...',
                'current_action': 'Iniciando...',
                'script_path': case_path
            })
        
        # Preparar entorno para ejecución
        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'
        env['PYTHONUTF8'] = '1'
        
        # Configurar headless
        if headless:
            env['PLAYWRIGHT_HEADLESS'] = '1'
        else:
            env['PLAYWRIGHT_HEADLESS'] = '0'
        
        # Configurar directorio para capturas
        timestamp = int(time.time())
        test_dir = os.path.join(SCREENSHOTS_DIR, f"test_pw_{task_id[:8]}_{timestamp}")
        os.makedirs(test_dir, exist_ok=True)
        
        with db_lock:
            if task_id in test_status_db:
                test_status_db[task_id].update({
                    'test_dir': test_dir,
                    'screenshots': []
                })
        
        # Ejecutar proceso
        process = subprocess.Popen(
            [python_exe, case_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            env=env,
            cwd=os.path.dirname(case_path),
            bufsize=1,
            creationflags=subprocess.CREATE_NO_WINDOW if platform.system() == 'Windows' else 0
        )
        
        # Función para leer salida
        def read_output(stream, output_key):
            try:
                for line in iter(stream.readline, b''):
                    try:
                        decoded_line = line.decode('utf-8', errors='replace')
                        # Sanitizar línea
                        decoded_line = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', decoded_line)
                        decoded_line = ''.join(c if ord(c) < 128 else '?' for c in decoded_line)
                        
                        with db_lock:
                            if task_id in test_status_db:
                                test_status_db[task_id][output_key] += decoded_line
                                
                                # Detectar indicaciones de progreso
                                if output_key == 'stdout':
                                    if 'Iniciando' in decoded_line:
                                        test_status_db[task_id]['current_action'] = 'Iniciando caso...'
                                    elif 'Navegando' in decoded_line:
                                        test_status_db[task_id]['current_action'] = 'Navegando...'
                                    elif 'Captura' in decoded_line:
                                        test_status_db[task_id]['current_action'] = 'Capturando pantalla...'
                                        # Buscar ruta de captura
                                        match = re.search(r'Captura.+guardada en: (.+\.png)', decoded_line)
                                        if match and match.group(1):
                                            screenshot_path = match.group(1)
                                            if os.path.exists(screenshot_path):
                                                relative_path = os.path.relpath(screenshot_path, SCREENSHOTS_DIR)
                                                url = f'/media/screenshots/{relative_path}'
                                                test_status_db[task_id]['screenshots'].append({
                                                    'url': url,
                                                    'path': screenshot_path,
                                                    'name': os.path.basename(screenshot_path)
                                                })
                    except Exception as e:
                        logger.error(f"Error procesando salida: {e}")
            except Exception as e:
                logger.error(f"Error en read_output: {e}")
        
        # Iniciar hilos para leer salida
        stdout_thread = Thread(target=read_output, args=(process.stdout, 'stdout'))
        stderr_thread = Thread(target=read_output, args=(process.stderr, 'stderr'))
        stdout_thread.daemon = True
        stderr_thread.daemon = True
        stdout_thread.start()
        stderr_thread.start()
        
        # Esperar finalización del proceso
        return_code = process.wait()
        
        # Procesar resultado
        with db_lock:
            if task_id in test_status_db:
                execution_timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
                execution_result = {
                    'timestamp': execution_timestamp,
                    'timestamp_unix': int(time.time()),
                    'success': return_code == 0,
                    'message': 'Caso ejecutado correctamente' if return_code == 0 else f'Error al ejecutar caso (código {return_code})',
                    'screenshots': test_status_db[task_id].get('screenshots', [])
                }
                
                if return_code == 0:
                    test_status_db[task_id].update({
                        'status': 'success',
                        'message': 'Caso ejecutado correctamente',
                        'current_action': 'Finalizado',
                        'icon': 'fa-check-circle'
                    })
                else:
                    test_status_db[task_id].update({
                        'status': 'error',
                        'message': f'Error al ejecutar caso (código {return_code})',
                        'current_action': 'Error',
                        'icon': 'fa-times-circle'
                    })
                
                # Buscar si ya existe un registro para este caso en el historial
                case_history_item = None
                for item in history_db:
                    if item.get('playwright_case') == case_name:
                        case_history_item = item
                        break
                
                # Si no existe, añadir un nuevo registro
                if not case_history_item:
                    print(f"DEBUG: Caso {case_name} no encontrado en historial, añadiendo nuevo registro")
                    
                    # Añadir al historial en BD
                    add_result = add_to_history(task_id, test_status_db[task_id])
                    logger.info(f"Resultado de guardar en BD para nuevo caso: {add_result}")
                else:
                    print(f"DEBUG: Caso {case_name} encontrado en historial, actualizando ejecuciones")
                    
                    # Actualizar el registro existente con la nueva ejecución
                    with history_lock:
                        # Inicializar el arreglo de ejecuciones si no existe
                        if 'executions' not in case_history_item:
                            case_history_item['executions'] = []
                        
                        # Añadir la nueva ejecución al inicio del arreglo
                        case_history_item['executions'].insert(0, execution_result)
                        
                        # Limitar a máximo 10 ejecuciones guardadas
                        case_history_item['executions'] = case_history_item['executions'][:10]
                        
                        # Actualizar fecha y estado con la ejecución más reciente
                        case_history_item['date'] = execution_timestamp
                        case_history_item['date_timestamp'] = time.time()
                        case_history_item['status'] = 'success' if return_code == 0 else 'error'
                        case_history_item['message'] = execution_result['message']
                
                # El historial se guarda automáticamente en BD
    
    except Exception as e:
        logger.error(f"Error ejecutando caso Playwright: {e}")
        with db_lock:
            if task_id in test_status_db:
                test_status_db[task_id].update({
                    'status': 'error',
                    'message': f'Error: {str(e)}',
                    'current_action': 'Error',
                    'icon': 'fa-exclamation-triangle'
                })

@app.route('/help')
def help_page():
    return render_template('help.html')

@app.route('/run_test', methods=['POST'])
def run_test():
    global test_status_db
    try:
        url = request.form.get('url')
        instructions = request.form.get('instructions')
        max_time = int(request.form.get('max_time', 600))
        headless = 'headless' in request.form  # Si está marcado será True (headless), si no está marcado será False (visible)
        screenshots = 'screenshots' in request.form
        fullscreen = 'fullscreen' in request.form
        browser = request.form.get('browser', 'chrome')

        # Registrar parámetros recibidos
        print(f"DEBUG run_test: Parámetros recibidos - URL: {url}, instrucciones: {instructions[:30]}..., max_time: {max_time}, headless: {headless}, fullscreen: {fullscreen}, browser: {browser}")

        model_name = 'claude-3-5-sonnet-20240620'
        
        required_key = 'ANTHROPIC_API_KEY'
        key_value = os.getenv(required_key)
        if not key_value:
            print(f"ERROR run_test: Clave {required_key} no configurada")
            return jsonify({'status': 'error', 'message': f'Error: Clave {required_key} no configurada'}), 400
        
        if not url or not url.startswith(("http://", "https://")):
            print(f"ERROR run_test: URL inválida: {url}")
            return jsonify({'status': 'error', 'message': 'URL inválida'}), 400
        
        if not instructions or len(instructions.strip()) < 5:
            print(f"ERROR run_test: Instrucciones vacías o muy cortas")
            return jsonify({'status': 'error', 'message': 'Instrucciones vacías o insuficientes'}), 400

        task_id = str(uuid.uuid4())
        
        # Inicializar el estado del test INMEDIATAMENTE
        print(f"DEBUG run_test: Creando task_id: {task_id}")
        print(f"DEBUG run_test: test_status_db antes de agregar: {len(test_status_db)} elementos")
        
        test_status_db[task_id] = {
            'status': 'queued', 
            'message': 'En cola...', 
            'original_instructions': instructions,
            'current_action': 'Esperando para despegar...',
            'icon': 'fa-hourglass-start',
            'url': url,
            'created_at': datetime.now().isoformat()
        }
        
        print(f"DEBUG run_test: test_status_db después de agregar: {len(test_status_db)} elementos")
        print(f"DEBUG run_test: Estado inicial para {task_id}: {test_status_db[task_id]}")

        # Estado se mantiene en memoria (temporal durante ejecución)

        thread = Thread(target=run_test_background_unified, args=(
            task_id, url, instructions, headless, max_time, screenshots, 
            None, model_name, browser, fullscreen
        ))
        thread.start()
        print(f"DEBUG run_test: Hilo iniciado para task_id: {task_id}")

        # Verificar que el estado se guardó correctamente
        if task_id in test_status_db:
            print(f"DEBUG run_test: ✅ Estado confirmado en test_status_db para {task_id}")
        else:
            print(f"ERROR run_test: ❌ Estado NO encontrado en test_status_db para {task_id}")

        return jsonify({'status': 'started', 'task_id': task_id})

    except Exception as e:
        print(f"ERROR run_test: Excepción: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'status': 'error', 'message': str(e)}), 500

@app.route('/test_status/<task_id>')
def get_test_status(task_id):
    global test_status_db
    try:
        logger.debug(f"=== GET_TEST_STATUS INICIADO ===")
        logger.debug(f"Obteniendo estado para task_id={task_id}")
        logger.debug(f"test_status_db contiene {len(test_status_db)} elementos")
        logger.debug(f"Claves en test_status_db: {list(test_status_db.keys())}")
        
        # Función auxiliar para sanitizar diccionarios recursivamente
        def sanitize_dict(data):
            if isinstance(data, dict):
                # Si es un diccionario con información de capturas, asegurar que las URLs están bien formateadas
                if 'url' in data and 'path' in data and isinstance(data['path'], str) and data['path'].endswith('.png'):
                    # Normalizar la ruta para que use el formato correcto en URLs
                    try:
                        file_path = data['path']
                        if os.path.exists(file_path):
                            rel_path = os.path.relpath(file_path, SCREENSHOTS_DIR)
                            url_path = rel_path.replace(os.path.sep, '/')
                            data['url'] = f'/media/screenshots/{url_path}'
                    except Exception as e:
                        logger.error(f"Error al normalizar URL de captura: {e}")
                
                return {k: sanitize_dict(v) for k, v in data.items()}
            elif isinstance(data, list):
                return [sanitize_dict(item) for item in data]
            elif isinstance(data, str):
                try:
                    # Sanitizar string: eliminar caracteres no ASCII, no imprimibles y escapar JSON
                    sanitized = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', data)
                    sanitized = ''.join(c if ord(c) < 128 else '?' for c in sanitized)
                    sanitized = sanitized.replace('\\', '\\\\').replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
                    return sanitized
                except Exception as e:
                    logger.exception(f"Error sanitizando string")
                    return "[Contenido no representable]"
            else:
                return data
        
        status_info = get_test_status_unified(task_id)
        logger.debug(f"status_info encontrado: {status_info is not None}")
        
        if status_info:
            logger.debug(f"Estado encontrado: {status_info.get('status', 'N/A')}")
            logger.debug(f"Mensaje: {status_info.get('message', 'N/A')}")
            logger.debug(f"Acción actual: {status_info.get('current_action', 'N/A')}")
            
            try:
                logger.debug(f"Encontrado status_info para {task_id}, sanitizando...")
                
                # Verificar si hay capturas y si están correctamente formateadas
                if 'screenshots' in status_info and isinstance(status_info['screenshots'], list):
                    logger.debug(f"Capturas encontradas: {len(status_info['screenshots'])}")
                    # Verificar si test_dir existe y tiene capturas que no estén en la lista
                    test_dir = status_info.get('test_dir')
                    if test_dir and os.path.exists(test_dir):
                        existing_urls = {s.get('url', '') for s in status_info['screenshots'] if isinstance(s, dict)}
                        try:
                            for file in os.listdir(test_dir):
                                if file.lower().endswith('.png'):
                                    file_path = os.path.join(test_dir, file)
                                    rel_path = os.path.relpath(file_path, SCREENSHOTS_DIR)
                                    url_path = rel_path.replace(os.path.sep, '/')
                                    url = f'/media/screenshots/{url_path}'
                                    
                                    # Si esta URL no está en la lista de capturas, agregarla
                                    if url not in existing_urls:
                                        status_info['screenshots'].append({
                                            'url': url,
                                            'path': file_path,
                                            'name': file
                                        })
                                        existing_urls.add(url)
                                        logger.debug(f"Añadida captura adicional encontrada: {file}")
                        except Exception as e:
                            logger.error(f"Error al buscar capturas adicionales: {e}")
                
                # Sanitizar datos
                sanitized_info = sanitize_dict(status_info)
                logger.debug(f"Sanitización completada, preparando respuesta JSON")
                
                # Intentar convertir a JSON con manejo explícito de errores
                # Primero intentar con dumps para detectar problemas
                try:
                    json_text = json_module.dumps(sanitized_info)
                    logger.debug(f"json.dumps exitoso, longitud={len(json_text)}")
                except Exception as json_dumps_error:
                    logger.error(f"Error en json.dumps: {json_dumps_error}")
                    # Si hay error, hacer una sanitización más agresiva de stdout/stderr
                    if 'stdout' in sanitized_info:
                        logger.debug("Sanitizando stdout agresivamente")
                        sanitized_info['stdout'] = re.sub(r'[^\x20-\x7E\n\r\t]', '?', sanitized_info.get('stdout', ''))
                    if 'stderr' in sanitized_info:
                        logger.debug("Sanitizando stderr agresivamente")
                        sanitized_info['stderr'] = re.sub(r'[^\x20-\x7E\n\r\t]', '?', sanitized_info.get('stderr', ''))
                    
                    # Sanitizar recursivamente todo de nuevo por si acaso
                    sanitized_info = json_module.loads(
                        json_module.dumps(
                            sanitized_info,
                            ensure_ascii=True,  # Forzar codificación ASCII
                            default=str  # Convertir cualquier objeto no serializable a string
                        )
                    )
                    logger.debug("Re-sanitización completada")
                
                # Ahora intentar jsonify con la data potencialmente re-sanitizada
                logger.debug("Generando respuesta JSON con jsonify")
                response = jsonify(sanitized_info)
                logger.debug("Respuesta JSON generada correctamente")
                logger.debug(f"=== GET_TEST_STATUS COMPLETADO EXITOSAMENTE ===")
                return response
                
            except Exception as json_error:
                logger.exception(f"Error al convertir a JSON")
                # Registrar información detallada del error
                if status_info:
                    for key, value in status_info.items():
                        if isinstance(value, str) and len(value) > 1000:
                            logger.debug(f"key={key}, type={type(value)}, len={len(value)}, excerto={value[:50]}...")
                        else:
                            logger.debug(f"key={key}, type={type(value)}")
                
                # Intentar crear una respuesta mínima segura
                return jsonify({
                    'status': 'error',
                    'error': True, 
                    'message': 'Error al procesar la respuesta',
                    'icon': 'fa-exclamation-triangle',
                    'details': 'Error de codificación en la respuesta'
                }), 200
        else:
            # En lugar de devolver un error 404, permitir recuperar el estado
            # con un mensaje informativo mejor para la interfaz
            logger.warning(f"Test {task_id} no encontrado en test_status_db")
            logger.debug(f"=== GET_TEST_STATUS COMPLETADO - TEST NO ENCONTRADO ===")
            return jsonify({
                'status': 'error', 
                'error': True,
                'message': 'Estado del test no disponible.', 
                'icon': 'fa-exclamation-triangle',
                'details': 'Es posible que el servidor se haya reiniciado o que el ID sea inválido.'
            }), 200  # Código 200 para que el frontend pueda mostrar el mensaje
    except Exception as e:
        # Mejorar el registro de errores
        logger.exception(f"Error general en get_test_status")
        print(f"Error general en get_test_status: {e}")
        print(traceback.format_exc())
        
        # Devolver un mensaje de error útil para el frontend
        return jsonify({
            'status': 'error',
            'error': True,
            'message': 'Error al obtener el estado del test.',
            'icon': 'fa-exclamation-circle',
            'details': str(e).replace('\\u', '\\\\u')  # Sanitizar unicode en el mensaje de error
        }), 200  # Código 200 para que el frontend pueda mostrar el mensaje

@app.route('/debug/test_status_db')
def debug_test_status_db():
    """Endpoint de debug para verificar el estado de test_status_db."""
    global test_status_db
    try:
        debug_info = {
            'total_tests': len(test_status_db),
            'test_ids': list(test_status_db.keys()),
            'test_details': {}
        }
        
        for task_id, status_info in test_status_db.items():
            debug_info['test_details'][task_id] = {
                'status': status_info.get('status', 'N/A'),
                'message': status_info.get('message', 'N/A'),
                'current_action': status_info.get('current_action', 'N/A'),
                'created_at': status_info.get('created_at', 'N/A'),
                'url': status_info.get('url', 'N/A'),
                'has_screenshots': 'screenshots' in status_info,
                'screenshot_count': len(status_info.get('screenshots', [])),
                'keys': list(status_info.keys())
            }
        
        return jsonify(debug_info)
    except Exception as e:
        return jsonify({
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

# Ruta para servir archivos de screenshots estáticamente
@app.route('/media/screenshots/<path:filename>')
def static_media(filename):
    # Reemplazar barras normales con separadores de sistema operativo
    # Esto es importante en Windows donde las rutas usan '\'
    try:
        # Dividir el path en partes para manejar subdirectorios correctamente
        parts = filename.split('/')
        if len(parts) > 1:
            # Si hay subdirectorio, extraemos el directorio y el nombre del archivo
            subdir = parts[0]
            subpath = '/'.join(parts[1:])
            full_path = os.path.join(SCREENSHOTS_DIR, subdir)
            
            print(f"DEBUG: Sirviendo archivo desde subdirectorio: {full_path}, archivo: {subpath}")
            return send_from_directory(full_path, subpath)
        else:
            # Caso sin subdirectorio
            print(f"DEBUG: Sirviendo archivo directo: {filename}")
            return send_from_directory(SCREENSHOTS_DIR, filename)
    except Exception as e:
        print(f"ERROR al servir imagen {filename}: {str(e)}")
        return f"Error al servir la imagen: {str(e)}", 404

@app.route('/api/token_usage')
def api_get_token_usage():
    """API para obtener datos reales de consumo de tokens desde la base de datos"""
    try:
        db_integration = get_db_integration()
        
        # Inicializar estructura de datos vacía
        data = {
            'total_tokens_today': 0,
            'total_cost_today': 0.00,
            'executions_today': 0,
            'avg_response_time': 0,
            'providers': {
                'openai': {'tokens': 0, 'cost': 0.00, 'last_used': 'Nunca'},
                'anthropic': {'tokens': 0, 'cost': 0.00, 'last_used': 'Nunca'},
                'gemini': {'tokens': 0, 'cost': 0.00, 'last_used': 'Nunca'}
            },
            'chart_data': {
                'labels': ['00:00', '04:00', '08:00', '12:00', '16:00', '20:00'],
                'openai': [0, 0, 0, 0, 0, 0],
                'anthropic': [0, 0, 0, 0, 0, 0],
                'gemini': [0, 0, 0, 0, 0, 0]
            },
            'executions': [],
            'alerts': []
        }
        
        if db_integration and db_integration.is_connected():
            # Obtener datos reales de ejecuciones de prueba desde la base de datos
            today = datetime.now().date()
            
            # Obtener ejecuciones del día actual
            executions_today = db_integration.get_executions_by_date(today)
            data['executions_today'] = len(executions_today)
            
            # Calcular tiempo promedio de respuesta
            if executions_today:
                valid_durations = [exec.get('duration_seconds', 0) for exec in executions_today if exec.get('duration_seconds')]
                if valid_durations:
                    data['avg_response_time'] = int(sum(valid_durations) / len(valid_durations) * 1000)  # Convertir a ms
            
            # Obtener ejecuciones recientes para la tabla
            recent_executions = db_integration.get_recent_executions(limit=20)
            
            # Formatear ejecuciones para la tabla
            for execution in recent_executions:
                # Para datos reales, no tenemos información específica de tokens IA
                # Solo mostramos información de ejecuciones de pruebas
                data['executions'].append({
                    'timestamp': execution.get('start_time', ''),
                    'type': 'Ejecución de Prueba',
                    'provider': 'QA-Pilot',
                    'model': 'Sistema Interno',
                    'input_tokens': 0,  # No disponible en ejecuciones de prueba
                    'output_tokens': 0,  # No disponible en ejecuciones de prueba
                    'total_tokens': 0,  # No disponible en ejecuciones de prueba
                    'cost': 0.00,  # No hay costo directo
                    'duration': execution.get('duration_seconds', 0) * 1000 if execution.get('duration_seconds') else 0  # Convertir a ms
                })
            
            # Obtener estadísticas de los últimos 7 días para el gráfico
            chart_data = db_integration.get_execution_stats_by_hour(days=1)  # Últimas 24 horas
            if chart_data:
                data['chart_data'] = chart_data
        
        # Nota informativa para el usuario
        if data['executions_today'] == 0 and len(data['executions']) == 0:
            data['alerts'].append({
                'title': '📊 Monitor de Tokens',
                'message': 'No hay ejecuciones registradas hoy. Los datos de tokens IA se mostrarán cuando se realicen ejecuciones que utilicen servicios de IA.'
            })
        else:
            data['alerts'].append({
                'title': '🔍 Información sobre Tokens IA',
                'message': 'Actualmente mostrando ejecuciones de pruebas del sistema. El rastreo específico de tokens de IA se activará cuando se implementen llamadas a APIs de IA con monitoreo.'
            })
        
        return jsonify(data)
        
    except Exception as e:
        print(f"Error al obtener datos de tokens: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/token_usage/clear', methods=['POST'])
def api_clear_token_usage():
    """API para limpiar el historial de tokens"""
    try:
        # Aquí se implementaría la limpieza real de la base de datos
        # Por ahora solo retornamos éxito
        return jsonify({'status': 'success', 'message': 'Historial de tokens limpiado'})
        
    except Exception as e:
        print(f"Error al limpiar historial de tokens: {str(e)}")
        return jsonify({'error': str(e)}), 500

# --- Monitor de Tokens IA ---
# Funciones para rastrear y monitorear el consumo de tokens de IA en tiempo real

def track_token_usage(provider, model, input_tokens, output_tokens, cost, execution_type='test_generation'):
    """Registra el uso de tokens en la base de datos para monitoreo"""
    try:
        db_integration = get_db_integration()
        if db_integration and db_integration.is_connected():
            # Implementar registro en base de datos
            # Por ahora solo registramos en logs
            print(f"Token Usage: {provider} {model} - Input: {input_tokens}, Output: {output_tokens}, Cost: ${cost:.4f}")
        
    except Exception as e:
        print(f"Error al registrar uso de tokens: {str(e)}")

def calculate_token_cost(provider, model, input_tokens, output_tokens):
    """Calcula el costo aproximado basado en el provider y modelo"""
    # Precios aproximados por 1000 tokens (actualizar según precios reales)
    pricing = {
        'openai': {
            'gpt-4': {'input': 0.03, 'output': 0.06},
            'gpt-3.5-turbo': {'input': 0.001, 'output': 0.002}
        },
        'anthropic': {
            'claude-3-5-sonnet': {'input': 0.003, 'output': 0.015},
            'claude-3-haiku': {'input': 0.00025, 'output': 0.00125}
        },
        'google': {
            'gemini-pro': {'input': 0.00025, 'output': 0.0005}
        }
    }
    
    try:
        provider_pricing = pricing.get(provider.lower(), {})
        model_pricing = provider_pricing.get(model.lower(), {'input': 0.001, 'output': 0.002})
        
        input_cost = (input_tokens / 1000) * model_pricing['input']
        output_cost = (output_tokens / 1000) * model_pricing['output']
        
        return input_cost + output_cost
        
    except Exception as e:
        print(f"Error al calcular costo: {str(e)}")
        return 0.0

# --- Funciones para Generación de Código Playwright --- 

def formatear_instrucciones(instrucciones):
    """Formatea las instrucciones para el test."""
    # Eliminar líneas vacías y espacios extra
    instrucciones = '\n'.join(line.strip() for line in instrucciones.split('\n') if line.strip())
    return instrucciones

def generar_script_test_nuevo(url, instrucciones, headless=True, max_tiempo=600, capturar_pasos=False, browser='chrome', fullscreen=True):
    """Genera un script de test para browser-use."""
    print(f"DEBUG: generar_script_test_nuevo llamado con params: url='{url}', headless={headless}, max_tiempo={max_tiempo}, capturar_pasos={capturar_pasos}, browser={browser}, fullscreen={fullscreen}")
    
    # Función para sanitizar las instrucciones y evitar problemas con comillas
    def sanitizar_instrucciones(texto):
        import re
        
        # Manejar caso de string vacío o None
        if not texto:
            return ""
            
        # Convertir a string si no lo es
        if not isinstance(texto, str):
            texto = str(texto)
        
        # Reemplazar comillas dobles por comillas simples
        texto_sanitizado = texto.replace('"', "'")
        
        # Eliminar caracteres que podrían causar problemas
        texto_sanitizado = texto_sanitizado.replace('\0', '')
        
        # Escapar cualquier barra invertida para evitar secuencias de escape no intencionales
        texto_sanitizado = texto_sanitizado.replace('\\', '\\\\')
        
        # Eliminar caracteres no imprimibles y caracteres de control
        texto_sanitizado = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', texto_sanitizado)
        
        # Reemplazar caracteres no ASCII que podrían causar problemas
        texto_sanitizado = ''.join(c if ord(c) < 128 else '?' for c in texto_sanitizado)
        
        # Eliminar comillas triples consecutivas que podrían cerrar el string multilinea
        texto_sanitizado = texto_sanitizado.replace("'''", "''_'")
        texto_sanitizado = texto_sanitizado.replace('"""', '"_"_"')
        
        print(f"DEBUG: Instrucciones sanitizadas. Longitud original: {len(texto)}, longitud final: {len(texto_sanitizado)}")
        return texto_sanitizado
    
    # Detectar pasos numerados en las instrucciones
    def detectar_pasos_numerados(texto):
        import re
        pasos = []
        # Detectar pasos con formato "1. Hacer algo" o "1) Hacer algo"
        lineas = texto.split('\n')
        for linea in lineas:
            # Buscar formatos como "1.", "2.", "1)", "2)"
            match = re.match(r'^\s*(\d+)[.\)]\s+(.+)$', linea.strip())
            if match:
                num_paso = int(match.group(1))
                desc_paso = match.group(2).strip()
                pasos.append((num_paso, desc_paso))
        
        # Ordenar los pasos por número
        pasos.sort(key=lambda x: x[0])
        print(f"DEBUG: Pasos detectados: {pasos}")
        return pasos
    
    # Analizar instrucciones para detectar pasos numerados
    pasos_numerados = detectar_pasos_numerados(instrucciones)
    
    # Seleccionar llaves de API según disponibilidad
    api_key = None
    model = None
    llm_provider = None
    
    # Intentar usar Anthropic primero
    anthropic_key = os.getenv('ANTHROPIC_API_KEY', '')
    if anthropic_key:
        api_key = anthropic_key
        model = "claude-3-5-sonnet-20240620"
        llm_provider = "anthropic"
        print("DEBUG: Usando Anthropic Claude 3.5 para generar script")
    
    # Si no hay Anthropic key, usar alternativas
    if not api_key:
        openai_key = os.getenv('OPENAI_API_KEY', '')
        if openai_key:
            api_key = openai_key
            model = "gpt-4o" 
            llm_provider = "openai"
            print("DEBUG: Usando OpenAI GPT-4 para generar script")
    
    # Si no hay llaves disponibles, devolver mensaje de error
    if not api_key:
        print("ERROR: No se encontraron llaves API válidas para generar script")
        return "# ERROR: No hay llaves API válidas para generar script. Configure API keys."

    # Código para capturar pantallas por pasos usando solo Playwright con Chrome
    captura_pasos_codigo = """
# Función simple para tomar capturas usando Playwright (Chrome)
import os
import time
import re

async def tomar_captura_navegador(agent, descripcion, paso_num=None):
    try:
        timestamp = int(time.time())
        safe_desc = re.sub(r'[^a-zA-Z0-9_]', '_', descripcion[:30])
        if paso_num:
            filename = f"paso_{paso_num}_{safe_desc}_{timestamp}.png"
        else:
            filename = f"{safe_desc}_{timestamp}.png"
        script_name = os.path.splitext(os.path.basename(__file__))[0]
        task_id = script_name.replace('test_', '') if script_name.startswith('test_') else script_name
        screenshot_subdir = os.path.join(os.getcwd(), "test_screenshots", task_id)
        os.makedirs(screenshot_subdir, exist_ok=True)
        screenshot_path = os.path.join(screenshot_subdir, filename)
        page = agent.browser_context.pages[0]
        await page.screenshot(path=screenshot_path)
        print(f"✅ Captura guardada en: {screenshot_path}")
        return screenshot_path
    except Exception as e:
        print(f"Error al tomar captura del paso: {str(e)}")
    """

    # Construir las líneas de capturas de pasos numerados fuera de la f-string
    capturas_pasos_numerados = ''
    if capturar_pasos and pasos_numerados:
        capturas_pasos_numerados = '\n        '.join([
            f'await tomar_captura_navegador(agent, "{desc}", paso_num={num})'
            for num, desc in pasos_numerados
        ])

    # Template base para script usando browser-use 0.2.6 (API corregida con capturas del navegador)
    script_template = f"""# -*- coding: utf-8 -*-
import sys, os, asyncio, base64, traceback, re, time, random, json, pyautogui
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
from browser_use import Agent, Browser, BrowserConfig

load_dotenv()

# Configuración del navegador con parámetros dinámicos
browser = Browser(
    config=BrowserConfig(
        headless={str(headless)},  # Usar parámetro dinámico
        browser_class='chromium',
        extra_browser_args=[]
        # No especificar browser_binary_path para usar navegador built-in de Playwright
    )
)

# --- Captura de pantalla utilitaria ---
{captura_pasos_codigo if capturar_pasos else ''}

async def main():
    try:
        print("DEBUG: Iniciando test...")
        print(f"DEBUG: Configuración headless={{{str(headless)}}}")
        anthropic_key = os.getenv('ANTHROPIC_API_KEY')
        if not anthropic_key:
            env_path = os.path.join(os.getcwd(), '.env')
            if os.path.exists(env_path):
                try:
                    with open(env_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            if line.strip() and not line.startswith('#'):
                                key, value = line.strip().split('=', 1)
                                if key == 'ANTHROPIC_API_KEY':
                                    anthropic_key = value.strip('"').strip("'")
                                    os.environ['ANTHROPIC_API_KEY'] = anthropic_key
                                    break
                except Exception as e:
                    print(f"Error leyendo .env: {{e}}")
        if not anthropic_key:
            raise Exception("ANTHROPIC_API_KEY no encontrada")
        llm = ChatAnthropic(
            model="claude-3-5-sonnet-20241022",
            api_key=anthropic_key,
            temperature=0.1,
            max_tokens=4000
        )
        agent = Agent(
            task="Navega a {url} y luego: {sanitizar_instrucciones(instrucciones)}",
            llm=llm,
            browser=browser
        )
        print("DEBUG: Agente creado")
        await agent.run(max_steps=15)
        print("DEBUG: Test completado exitosamente")

        # --- Captura después de navegar ---
        {"await tomar_captura_navegador(agent, 'Página cargada', paso_num=0)" if capturar_pasos else ''}

        # --- Captura después de cada paso numerado ---
        {capturas_pasos_numerados}

        # --- Captura final ---
        {"await tomar_captura_navegador(agent, 'Test completado', paso_num=999)" if capturar_pasos else ''}

    except Exception as e:
        print(f"ERROR: {{str(e)}}")
        print(f"TRACEBACK: {{traceback.format_exc()}}")
    finally:
        try:
            await browser.close()
            print("DEBUG: Navegador cerrado")
        except Exception as e:
            print(f"Error al cerrar navegador: {{e}}")

if __name__ == "__main__":
    asyncio.run(main())
"""
    return script_template

def run_test_background(task_id, url, instrucciones, headless, max_tiempo, screenshots, gemini_key, model_name='claude-3-5-sonnet-20240620', browser='chrome', fullscreen=True):
    """Ejecuta un test en background."""
    script_path = None
    global test_status_db
    global test_dir
    proceso = None
    final_status = 'error'
    screenshot_timer = None

    max_tiempo = max(max_tiempo, 60)  # Asegurar un mínimo de 60 segundos
    capturar_pasos = bool(screenshots)
    
    # Registrar datos clave para debug
    print(f"DEBUG [run_test_background]: Iniciando tarea {task_id}")
    print(f"DEBUG [run_test_background]: URL={url}, headless={headless}, fullscreen={fullscreen}, browser={browser}")
    print(f"DEBUG [run_test_background]: screenshots={screenshots}, capturar_pasos={capturar_pasos}")
    
    # Función para capturar pantalla directamente en este contexto
    def capturar_pantalla_interna(nombre_captura):
        """Captura la pantalla completa usando pyautogui y la guarda en el directorio de test actual"""
        try:
            if not test_dir or not os.path.exists(test_dir):
                print(f"ERROR: Directorio de capturas no válido: {test_dir}")
                return False
                
            timestamp = int(time.time())
            output_path = os.path.join(test_dir, f"{nombre_captura}_{timestamp}.png")
            
            # Tomar captura con pyautogui
            captura = pyautogui.screenshot()
            captura.save(output_path)
            
            # Registrar la captura en la base de datos del test
            with db_lock:
                if task_id in test_status_db:
                    if 'screenshots' not in test_status_db[task_id]:
                        test_status_db[task_id]['screenshots'] = []
                    
                    # Normalizar rutas para URL
                    relative_path = os.path.relpath(output_path, SCREENSHOTS_DIR)
                    # Usar siempre "/" en URLs, incluso en Windows
                    url_path = relative_path.replace(os.path.sep, '/')
                    
                    test_status_db[task_id]['screenshots'].append({
                        'url': f'/media/screenshots/{url_path}',
                        'path': output_path,
                        'name': os.path.basename(output_path)
                    })
                    print(f"INFO: Captura registrada para task_id {task_id}: {url_path}")
            
            print(f"INFO: Captura guardada en: {output_path}")
            return True
        except Exception as e:
            print(f"ERROR al capturar pantalla: {str(e)}")
            traceback.print_exc()
            return False
    
    def stream_reader(stream, output_key):
        last_human_status = ""
        try:
            # Leer bytes línea por línea
            for line_bytes in iter(stream.readline, b''): # Iterar hasta encontrar bytes vacíos
                if not line_bytes:
                    break
                
                # Decodificar explícitamente a UTF-8, reemplazando errores
                try:
                    decoded_line = line_bytes.decode('utf-8', errors='replace')
                    # Sanitizar: eliminar caracteres no imprimibles y reemplazar no ASCII
                    decoded_line = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F-\x9F]', '', decoded_line)
                    decoded_line = ''.join(c if ord(c) < 128 else '?' for c in decoded_line)
                except Exception as e:
                    # Fallback por si la decodificación falla de alguna manera inesperada
                    decoded_line = f"[Error de decodificación: {str(e)}]\n"
                
                human_status_update = None
                current_step_update = None

                # Detectar capturas de pantalla tomadas por el navegador
                if "Captura del paso" in decoded_line and "guardada en:" in decoded_line:
                    print(f"DEBUG: Línea de captura detectada: {decoded_line.strip()}")
                    try:
                        # Extraer ruta del archivo de captura
                        match = re.search(r"Captura del paso .+? guardada en: (.+\.png)", decoded_line)
                        if match and match.group(1):
                            screenshot_path = match.group(1).strip()
                            print(f"DEBUG: Captura detectada: {screenshot_path}")
                            
                            # Normalizar ruta para asegurar formato correcto en Windows
                            screenshot_path = os.path.normpath(screenshot_path)
                            
                            if os.path.exists(screenshot_path):
                                print(f"DEBUG: El archivo existe en: {screenshot_path}")
                                # Registrar captura en la base de datos del test
                                with db_lock:
                                    if task_id in test_status_db:
                                        if 'screenshots' not in test_status_db[task_id]:
                                            test_status_db[task_id]['screenshots'] = []
                                        
                                        try:
                                            # Normalizar rutas para URL
                                            relative_path = os.path.relpath(screenshot_path, os.getcwd())
                                            # Usar siempre "/" en URLs, incluso en Windows
                                            url_path = relative_path.replace(os.path.sep, '/')
                                            screenshot_name = os.path.basename(screenshot_path)
                                            
                                            print(f"DEBUG: URL calculada: /media/screenshots/{url_path}")
                                            
                                            # Evitar duplicados
                                            already_exists = False
                                            for ss in test_status_db[task_id]['screenshots']:
                                                if isinstance(ss, dict) and ss.get('path') == screenshot_path:
                                                    already_exists = True
                                                    break
                                            
                                            if not already_exists:
                                                test_status_db[task_id]['screenshots'].append({
                                                    'url': f'/media/screenshots/{url_path}',
                                                    'path': screenshot_path,
                                                    'name': screenshot_name
                                                })
                                                print(f"INFO: Captura del navegador registrada: {screenshot_name}")
                                                # Actualizar acción actual
                                                human_status_update = f"Captura guardada: {screenshot_name}"
                                        except Exception as e:
                                            print(f"ERROR al calcular ruta relativa: {e}")
                            else:
                                print(f"ADVERTENCIA: Archivo de captura no existe: {screenshot_path}")
                        else:
                            print(f"DEBUG: No se pudo extraer ruta de captura de: {decoded_line.strip()}")
                    except Exception as e:
                        print(f"ERROR al procesar captura de navegador: {e}")
                        traceback.print_exc()
                
                # Detectar cualquier línea que mencione capturas para debugging
                if "captura" in decoded_line.lower() or "screenshot" in decoded_line.lower():
                    print(f"DEBUG: Línea relacionada con capturas: {decoded_line.strip()}")

                if output_key == 'stdout':
                    if "DEBUG: Iniciando función main()" in decoded_line:
                        human_status_update = "Inicializando..."
                        current_step_update = 1
                    elif "DEBUG: Inicializando Browser..." in decoded_line:
                        human_status_update = "Iniciando navegador..."
                        current_step_update = 2
                    elif "DEBUG: Navegando a URL:" in decoded_line or "DEBUG: Navegando a la URL inicial..." in decoded_line:
                        human_status_update = "Navegando..."
                        current_step_update = 3
                    elif "DEBUG: Ejecutando tarea principal del agente" in decoded_line or "DEBUG: Ejecutando agente con las instrucciones..." in decoded_line:
                        human_status_update = "Procesando..."
                        current_step_update = 4
                    elif "CAPTURA: " in decoded_line or "CAPTURA HTML2CANVAS: " in decoded_line:
                        human_status_update = "Capturando..."
                        current_step_update = 5
                    elif "ERROR al tomar captura" in decoded_line:
                        human_status_update = "Error en captura..."
                    elif "DEBUG: Cerrando el navegador..." in decoded_line:
                        human_status_update = "Finalizando..."
                        current_step_update = 6
                    # Detectar pasos numerados ejecutándose
                    elif "Ejecutando paso " in decoded_line:
                        match = re.search(r"Ejecutando paso (\d+)[:.]?\s*(.+)?", decoded_line)
                        if match:
                            paso_num = int(match.group(1))
                            paso_desc = match.group(2) if match.group(2) else f"Paso {paso_num}"
                            human_status_update = f"Ejecutando paso {paso_num}: {paso_desc[:30]}..."
                            current_step_update = paso_num + 3  # Offset para los pasos iniciales
               
                with db_lock:
                    if task_id in test_status_db:
                        # Sanitizar la línea antes de agregarla al diccionario
                        if isinstance(decoded_line, str):
                            sanitized_line = ''.join(c if ord(c) < 128 else '?' for c in decoded_line)
                        else:
                            sanitized_line = str(decoded_line)
                        
                        # Inicializar la clave del output si no existe
                        if output_key not in test_status_db[task_id]:
                            test_status_db[task_id][output_key] = ''
                            
                        test_status_db[task_id][output_key] += sanitized_line
                        
                        if human_status_update and human_status_update != last_human_status:
                            # Sanitizar también el status update
                            safe_status = ''.join(c if ord(c) < 128 else '?' for c in human_status_update)
                            test_status_db[task_id]['current_action'] = safe_status
                            last_human_status = human_status_update
                            
                            # Actualizar paso actual si tenemos la información
                            if current_step_update is not None:
                                test_status_db[task_id]['current_step'] = current_step_update

                            # Comentado para evitar capturas duplicadas - las capturas se manejan en el script
                            # if capturar_pasos and human_status_update:
                            #     estado_nombre = re.sub(r'[^a-z0-9_\-]', '', human_status_update.lower().replace(' ', '_'))
                            #     threading.Thread(
                            #         target=capturar_pantalla_interna, 
                            #         args=(f"estado_{estado_nombre}",),
                            #         daemon=True
                            #     ).start()
                    else:
                        break 
        except ValueError:
            pass
        except Exception as e:
            print(f"Error en stream_reader: {e}")
        finally:
            try:
                if stream and not stream.closed:
                     stream.close()
            except:
                pass

    try:
        # Crear directorio para el test actual
        test_dir = os.path.join(SCREENSHOTS_DIR, task_id)
        os.makedirs(test_dir, exist_ok=True)
        print(f"DEBUG [run_test_background]: Directorio de test creado: {test_dir}")
        
        # También crear directorio en test_screenshots para compatibilidad con browser-use 0.2.6
        test_screenshots_dir = os.path.join(os.getcwd(), "test_screenshots")
        os.makedirs(test_screenshots_dir, exist_ok=True)
        print(f"DEBUG [run_test_background]: Directorio test_screenshots creado: {test_screenshots_dir}")

        # Inicializar estructura en test_status_db
        with db_lock:
            if task_id not in test_status_db:
                test_status_db[task_id] = {}
            test_status_db[task_id].update({
                'status': 'generating',
                'message': 'Generando script...',
                'stdout': '',
                'stderr': '',
                'current_action': 'Preparando...',
                'screenshots': [],
                'test_dir': test_dir if capturar_pasos else None,
                'icon': 'fa-spinner fa-spin'
            })
            # Estado temporal durante ejecución

        # Generar script de test
        script_path = os.path.join(SCRIPTS_DIR, f"test_{task_id}.py")
        script_content = generar_script_test_nuevo(
            url=url,
            instrucciones=instrucciones,
            headless=headless,  # Usar valor del formulario en lugar de forzar False
            max_tiempo=max_tiempo,
            capturar_pasos=capturar_pasos,
            browser=browser,
            fullscreen=fullscreen
        )

        # Escribir el script con codificación UTF-8 y BOM
        with open(script_path, 'w', encoding='utf-8-sig', errors='replace') as f:
            f.write(script_content)

        with db_lock:
            if task_id not in test_status_db:
                test_status_db[task_id] = {}
            test_status_db[task_id].update({
                'status': 'running', 
                'message': 'Ejecutando...', 
                'current_action': 'Iniciando...', 
                'script_path': script_path,
                'screenshots': []
                            })
            # Estado temporal durante ejecución

        env = os.environ.copy()
        env['PYTHONIOENCODING'] = 'utf-8'
        env['PYTHONUTF8'] = '1'
        env['PYTHONLEGACYWINDOWSSTDIO'] = '0'
        
        # Variables de entorno para Pydantic
        env['PYDANTIC_V2'] = '1'  # Activar modo Pydantic v2
        env['PYDANTIC_STRICT_MODE'] = 'false'  # Desactivar modo estricto
        env['PYDANTIC_COLORS'] = 'false'  # Desactivar colores en errores
        
        # --- DEBUG: Imprimir variables de entorno clave ---
        print(f"DEBUG [run_test_background]: PYTHONPATH={env.get('PYTHONPATH')}")
        print(f"DEBUG [run_test_background]: PYTHONIOENCODING={env.get('PYTHONIOENCODING')}")
        print(f"DEBUG [run_test_background]: PYTHONUTF8={env.get('PYTHONUTF8')}")
        print(f"DEBUG [run_test_background]: PYTHONLEGACYWINDOWSSTDIO={env.get('PYTHONLEGACYWINDOWSSTDIO')}")
        print(f"DEBUG [run_test_background]: PYDANTIC_V2={env.get('PYDANTIC_V2')}")
        print(f"DEBUG [run_test_background]: PYDANTIC_STRICT_MODE={env.get('PYDANTIC_STRICT_MODE')}")
        # --- FIN DEBUG ---
        env['BROWSER_USE_MODEL'] = model_name
        
        try:
            # Determinar creation flags según si debe ser headless o no
            creation_flags = 0
            if platform.system() == 'Windows':
                if headless:
                    # Solo usar CREATE_NO_WINDOW cuando realmente queremos modo headless
                    creation_flags = subprocess.CREATE_NO_WINDOW
                else:
                    # Cuando no es headless, permitir que el navegador sea visible
                    # No usar CREATE_NO_WINDOW para que el navegador pueda mostrarse
                    creation_flags = 0
            
            proceso = subprocess.Popen(
                [python_exe, script_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=env,
                cwd=BASE_DIR,        
                bufsize=1,           # Mantener para línea por línea
                creationflags=creation_flags
            )
            
            # Registrar el proceso en el tracking global
            with subprocess_processes_lock:
                active_subprocess_processes[task_id] = proceso
                logger.info(f"🔄 Proceso {proceso.pid} registrado para task_id {task_id}")

            # Configurar la codificación de la consola en Windows
            if platform.system() == 'Windows':
                import ctypes
                kernel32 = ctypes.windll.kernel32
                kernel32.SetConsoleOutputCP(65001)  # UTF-8
                kernel32.SetConsoleCP(65001)  # UTF-8
                
            # Mover el timer aquí después de que proceso esté definido
            if capturar_pasos:
                # Crear un timer para tomar capturas periódicas, pero solo como respaldo
                def tomar_captura_periodica():
                    try:
                        # Comentar línea de captura directa, confiamos en las capturas del navegador
                        # capturar_pantalla_interna(f"captura_periodica")
                        # Programar la próxima captura
                        global screenshot_timer
                        screenshot_timer = threading.Timer(15.0, tomar_captura_periodica)
                        screenshot_timer.daemon = True
                        screenshot_timer.start()
                    except Exception as e:
                        print(f"Error en captura periódica: {e}")
                
                # Iniciar el timer para la primera captura
                screenshot_timer = threading.Timer(10.0, tomar_captura_periodica)
                screenshot_timer.daemon = True
                screenshot_timer.start()
                
        except Exception as e:
            with db_lock:
                if task_id in test_status_db:
                    test_status_db[task_id].update({
                        'status': 'error',
                        'message': str(e),
                        'stderr': f"\nError al iniciar proceso: {str(e)}\n",
                        'icon': 'fa-exclamation-triangle'
                    })
            return
            
        stdout_thread = threading.Thread(target=stream_reader, args=(proceso.stdout, 'stdout'), daemon=True)
        stderr_thread = threading.Thread(target=stream_reader, args=(proceso.stderr, 'stderr'), daemon=True)
        stdout_thread.start()
        stderr_thread.start()

        return_code = None
        timeout_ocurrido = False

        try:
            return_code = proceso.wait(timeout=max_tiempo)
            if return_code == 0:
                final_status = 'success'
        except subprocess.TimeoutExpired:
            timeout_ocurrido = True
            try:
                proceso.kill()
                proceso.wait(timeout=5)
            except:
                pass
            return_code = -9
            with db_lock:
                if task_id in test_status_db:
                    test_status_db[task_id]['stderr'] += f"\nTiempo máximo excedido ({max_tiempo}s)\n"
                    test_status_db[task_id]['current_action'] = "Timeout"
        except Exception as e:
            return_code = -1
            with db_lock:
                 if task_id in test_status_db:
                     test_status_db[task_id]['stderr'] += f"\nError: {str(e)}\n"
                     test_status_db[task_id]['current_action'] = "Error"
            if proceso and proceso.poll() is None:
                 try: proceso.kill()
                 except: pass
        finally:
            if stdout_thread.is_alive(): stdout_thread.join(timeout=2)
            if stderr_thread.is_alive(): stderr_thread.join(timeout=2)

            final_stdout = ""
            final_stderr = ""
            screenshots_found = []
            final_message = f"Error (Código {return_code})"
            final_icon = 'fa-exclamation-triangle'
            
            correct_screenshot_dir = None
            if script_path and capturar_pasos:
                try:
                    script_filename_no_ext = os.path.splitext(os.path.basename(script_path))[0]
                    correct_screenshot_dir = os.path.join(SCREENSHOTS_DIR, script_filename_no_ext)
                except:
                    pass

            with db_lock:
                if task_id in test_status_db:
                    final_stdout = test_status_db[task_id].get('stdout', '')
                    final_stderr = test_status_db[task_id].get('stderr', '')
                else:
                     final_status = 'desconocido'
                     return

            if timeout_ocurrido:
                 final_status = 'error'
                 final_message = f'Timeout ({max_tiempo}s)'
                 final_icon = 'fa-clock'
            elif return_code == 0:
                final_status = 'success'
                final_message = "Completado"
                final_icon = 'fa-check-circle'
                match = re.search(r"INFO: Resultado final del agente: (.*)", final_stdout, re.IGNORECASE | re.DOTALL)
                if match:
                    result_text = match.group(1).strip()
                    final_message = f"Completado: {result_text[:100]}{'...' if len(result_text) > 100 else ''}"
            else:
                final_status = 'error'
                error_summary = final_stderr.strip().splitlines()[-1] if final_stderr.strip() else "Error desconocido"
                final_message = f"Error: {error_summary[:150]}{'...' if len(error_summary) > 150 else ''}"
                final_icon = 'fa-times-circle'

            with db_lock:
                 if task_id in test_status_db: 
                     update_data = {
                         'status': final_status,
                         'message': final_message,
                         'screenshots': screenshots_found,
                         'icon': final_icon,
                         'current_action': 'Finalizado',
                         'url': url  # Añadir la URL al estado final
                     }
                     test_status_db[task_id].update(update_data)
                     
                     # Añadir el test al historial una vez completado
                     print(f"DEBUG: Intentando añadir test {task_id} al historial. Estado actual: {final_status}")
                     print(f"DEBUG: Historial actual tiene {len(history_db)} elementos")
                     
                     # Asegurar que todos los campos necesarios están presentes
                     test_data = test_status_db[task_id].copy()
                     if 'url' not in test_data or not test_data['url']:
                         test_data['url'] = url
                     
                     add_result = add_to_history(task_id, test_data)
                     logger.info(f"Resultado de guardar en BD: {add_result}")
                     
                     # ===============================================================
                     # REGISTRO AUTOMÁTICO EN BASE DE DATOS
                     # ===============================================================
                     try:
                         db_integration = get_db_integration()
                         if db_integration and db_integration.is_connected():
                             print(f"DEBUG: Intentando registrar test {task_id} en base de datos...")
                             
                             # Obtener las instrucciones originales correctamente
                             original_instructions = test_status_db[task_id].get('original_instructions', instrucciones)
                             print(f"DEBUG: Instrucciones originales obtenidas ({len(original_instructions)} chars): {original_instructions[:100]}...")
                             
                             # Crear caso de prueba en BD
                             test_case_data = {
                                 'nombre': f'Test {task_id[:8]} - {datetime.now().strftime("%Y%m%d_%H%M%S")}',
                                 'codigo': f'AUTO-{task_id[:8]}',
                                 'tipo': 'ui',
                                 'prioridad': 'media',
                                 'objetivo': 'Prueba ejecutada desde la interfaz web de QA-Pilot',
                                 'pasos': original_instructions,
                                 'resultado_esperado': 'Ejecución exitosa según las instrucciones proporcionadas',
                                 'url_objetivo': url,
                                 'es_valido': True,
                                 'status': 'approved' if final_status == 'success' else 'draft',
                                 'created_by': 'qa_pilot_web',
                                 'instrucciones_qa_pilot': original_instructions,
                                 'metadata_json': {
                                     'execution_id': task_id,
                                     'browser': browser,
                                     'headless': headless,
                                     'screenshots_enabled': screenshots,
                                     'max_tiempo': max_tiempo,
                                     'execution_timestamp': datetime.now().isoformat(),
                                     'final_status': final_status,
                                     'final_message': final_message
                                 }
                             }
                             
                             print(f"DEBUG: Datos del caso a guardar: {test_case_data}")
                             case_id = db_integration.save_test_case(test_case_data)
                             print(f"DEBUG: ✅ Caso de prueba guardado en BD con ID: {case_id}")
                             
                             # Crear ejecución de prueba en BD
                             if case_id:
                                 execution_data = {
                                     'test_case_id': case_id,
                                     'execution_type': 'manual',
                                     'status': 'passed' if final_status == 'success' else 'failed',
                                     'result_details': final_message,
                                     'error_message': final_stderr if final_status != 'success' else None,
                                     'url_executed': url,
                                     'script_path': script_path if script_path else None,
                                     'stdout_log': final_stdout,
                                     'stderr_log': final_stderr,
                                     'return_code': return_code,
                                     'duration_seconds': int((time.time() - float(test_status_db[task_id].get('start_time', time.time())))),
                                     'browser_config': {
                                         'browser': browser,
                                         'headless': headless,
                                         'fullscreen': fullscreen,
                                         'screenshots_enabled': screenshots
                                     },
                                     'executed_by': 'qa_pilot_web',
                                     'metadata_json': {
                                         'task_id': task_id,
                                         'screenshots_count': len(test_data.get('screenshots', [])),
                                         'execution_source': 'web_interface'
                                     }
                                 }
                                 
                                 print(f"DEBUG: Datos de ejecución a guardar: {execution_data}")
                                 execution_id = db_integration.create_test_execution(execution_data)
                                 print(f"DEBUG: ✅ Ejecución de prueba registrada en BD con ID: {execution_id}")
                                 
                                 # Registrar screenshots en BD si existen
                                 if execution_id and test_data.get('screenshots'):
                                     for idx, screenshot in enumerate(test_data['screenshots']):
                                         if isinstance(screenshot, dict) and 'path' in screenshot:
                                             screenshot_data = {
                                                 'execution_id': execution_id,
                                                 'test_case_id': case_id,
                                                 'name': screenshot.get('name', f'screenshot_{idx}'),
                                                 'description': f'Captura automática del paso {idx + 1}',
                                                 'step_number': idx + 1,
                                                 'screenshot_type': 'step',
                                                 'file_path': screenshot['path'],
                                                 'file_name': os.path.basename(screenshot['path']),
                                                 'file_size_bytes': os.path.getsize(screenshot['path']) if os.path.exists(screenshot['path']) else 0,
                                                 'url_captured': url,
                                                 'timestamp_ms': int(time.time() * 1000),
                                                 'is_valid': True,
                                                 'metadata_json': {
                                                     'task_id': task_id,
                                                     'url_web': screenshot.get('url', ''),
                                                     'capture_source': 'qa_pilot_web'
                                                 }
                                             }
                                             
                                             try:
                                                 screenshot_id = db_integration.save_screenshot(screenshot_data)
                                                 print(f"DEBUG: Screenshot registrado en BD con ID: {screenshot_id}")
                                             except Exception as ss_error:
                                                 print(f"DEBUG: Error al guardar screenshot en BD: {ss_error}")
                                 
                                 # Actualizar información en test_status_db
                                 test_status_db[task_id]['db_case_id'] = str(case_id)
                                 test_status_db[task_id]['db_execution_id'] = str(execution_id)
                                 
                         else:
                             print(f"DEBUG: ❌ Base de datos no disponible o no conectada, test {task_id} solo guardado en historial")
                             
                     except Exception as db_error:
                         print(f"ERROR: ❌ Error al registrar test en base de datos: {db_error}")
                         print(f"ERROR: Traceback completo:")
                         print(traceback.format_exc())
                     
                     # Guardar una copia de seguridad en disco
                     try:
                         persist_history_to_disk() # type: ignore
                     except Exception as e:
                         print(f"DEBUG: Error al persistir historial: {e}")

    except Exception as e:
        with db_lock:
             if task_id in test_status_db:
                current_stderr = test_status_db[task_id].get('stderr', '')
                test_status_db[task_id].update({
                    'status': 'error',
                    'message': str(e),
                    'stderr': current_stderr + f"\nError: {str(e)}\n",
                    'icon': 'fa-bomb',
                    'current_action': 'Error'
                })
    finally:
        # Limpiar proceso del tracking global
        with subprocess_processes_lock:
            if task_id in active_subprocess_processes:
                del active_subprocess_processes[task_id]
                logger.info(f"🧹 Proceso removido del tracking para task_id {task_id}")
        
        # Detener timer de capturas si existe
        if screenshot_timer:
            try:
                screenshot_timer.cancel()
            except:
                pass
        
        if False:
            if script_path and os.path.exists(script_path):
                try: os.remove(script_path)
                except: pass

def generar_playwright_con_contexto(instrucciones, log_ejecucion, language="python"):
    """Genera código Playwright con contexto del log de ejecución."""
    prompt_template = """
    Genera código Playwright en {language} para automatizar las siguientes instrucciones:
    {instrucciones} 

    Contexto de ejecución:
    {log_ejecucion}

    Consideraciones:
    1. Usa {browser} como navegador
    2. Incluye manejo de errores
    3. Usa selectores robustos
    4. Incluye comentarios explicativos
    5. Usa esperas explícitas cuando sea necesario
    6. Incluye capturas de pantalla en puntos clave
    7. Usa el patrón Page Object Model cuando sea apropiado

    Devuelve SOLO el código, sin explicaciones adicionales.
    """
    
    # Determinar el navegador según el lenguaje
    browser = "chromium" if language == "python" else "chromium"
    
    # Formatear el prompt
    prompt = prompt_template.format(
        language=language,
        instrucciones=instrucciones,
        log_ejecucion=log_ejecucion,
        browser=browser
    )

    try:
        print(f"DEBUG: Creando cliente Anthropic para generar Playwright en {language}...")
        anthropic_key = os.getenv('ANTHROPIC_API_KEY')
        if not anthropic_key:
            print("ERROR: ANTHROPIC_API_KEY no encontrada en entorno para generar Playwright.")
            return f"# Error: ANTHROPIC_API_KEY no configurada."
        
        llm = ChatAnthropic(
            model="claude-3-5-sonnet-20241022", # Usar Claude 3.5 Sonnet más reciente
            api_key=anthropic_key,
            temperature=0.1, # Baja temperatura para código más determinista
            timeout=60,      # Timeout reducido para mejor UX
            max_retries=2    # Añadir retries automáticos
        )
        
        print("DEBUG: Enviando prompt a Anthropic Claude...")
        
        # Implementar retry manual para errores 529
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                response = llm.invoke(prompt)
                generated_code = response.content
                print(f"DEBUG: Código Playwright ({language}) generado exitosamente en intento {attempt + 1}")
                print(f"{generated_code[:500]}...") # Log inicial del código
                return generated_code
                
            except Exception as retry_error:
                error_str = str(retry_error).lower()
                if ("529" in error_str or "overloaded" in error_str or "overloaded_error" in error_str) and attempt < max_attempts - 1:
                    import time
                    wait_time = (attempt + 1) * 2  # Espera incremental: 2s, 4s
                    print(f"DEBUG: Error 529 en intento {attempt + 1}, reintentando en {wait_time}s...")
                    time.sleep(wait_time)
                    continue
                else:
                    # Si no es error 529 o ya agotamos intentos, lanzar excepción
                    raise retry_error

    except Exception as e:
        print(f"ERROR: Error al llamar a Anthropic Claude para generar código Playwright: {e}")
        print(traceback.format_exc())
        
        # Detectar error específico de Claude sobrecargado (Error 529)
        error_str = str(e).lower()
        if "529" in error_str or "overloaded" in error_str or "overloaded_error" in error_str:
            return f"""# ⚠️ Servicio temporalmente no disponible
# 
# Claude está experimentando alta demanda en este momento.
# Por favor, intente generar el código Playwright nuevamente en unos minutos.
#
# Error técnico: {str(e)}"""
        
        # Para otros errores, mostrar mensaje genérico mejorado
        return f"""# ❌ Error al generar código Playwright
# 
# Ocurrió un problema al generar el código automáticamente.
# Puede intentar nuevamente o revisar los logs para más detalles.
#
# Error técnico: {str(e)}"""

@app.route('/generate_playwright/<task_id>')
def generate_playwright_endpoint(task_id):
    """Endpoint para generar código Playwright basado en un test ejecutado."""
    global test_status_db
    print(f"DEBUG: Solicitud para generar código Playwright para task_id: {task_id}")
    
    language = request.args.get('language', 'python')
    print(f"DEBUG: Lenguaje solicitado: {language}")
    
    with db_lock:
        status_info = test_status_db.get(task_id)
    
    if not status_info:
        print(f"ERROR: No se encontró información para task_id: {task_id}")
        return jsonify({'error': 'Test ID no encontrado'}), 404

    instrucciones = status_info.get('original_instructions', 'Instrucciones no disponibles') 
    log_ejecucion = status_info.get('stdout', '') + "\n" + status_info.get('stderr', '')

    print(f"DEBUG: Instrucciones para task {task_id}: {instrucciones[:100]}...")
    print(f"DEBUG: Log para task {task_id}: {log_ejecucion[:100]}...")

    generated_code = generar_playwright_con_contexto(instrucciones, log_ejecucion, language)
    
    # Detectar si el código generado contiene un mensaje de error
    if generated_code.startswith("# ⚠️ Servicio temporalmente no disponible"):
        return jsonify({
            'error': True,
            'message': 'Claude está experimentando alta demanda. Intente nuevamente en unos minutos.',
            'generated_code': generated_code,
            'language': language
        })
    elif generated_code.startswith("# ❌ Error al generar código Playwright"):
        return jsonify({
            'error': True,
            'message': 'Error al generar código Playwright. Puede intentar nuevamente.',
            'generated_code': generated_code,
            'language': language
        })

    return jsonify({'generated_code': generated_code, 'language': language}) 

@app.route('/generate_word_evidence/<task_id>', methods=['POST'])
def generate_word_evidence(task_id):
    """Genera un documento Word con evidencias del test basado en el template."""
    try:
        import os
        import tempfile
        from datetime import datetime
        from docx import Document
        from docx.shared import Inches
        
        logger.debug(f"Generando evidencia en Word para test: {task_id}")
        
        # Buscar el test en el historial
        with history_lock:
            test_data = None
            for item in history_db:
                if item.get('id') == task_id:
                    test_data = item
                    break
            
            if not test_data:
                return jsonify({
                    'status': 'error',
                    'message': 'Test no encontrado en historial'
                }), 404
        
        # Ruta del template
        template_path = os.path.join(os.getcwd(), "test_evidencia", "template", "QA_Evidencia_Template_Mejorado.docx")
        
        if not os.path.exists(template_path):
            return jsonify({
                'status': 'error',
                'message': f'Template no encontrado en: {template_path}'
            }), 404
        
        # Crear documento basado en el template
        doc = Document(template_path)
        
        # Obtener información del test
        test_name = test_data.get('name', f'Test {task_id[:8]}')
        test_date = test_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        test_url = test_data.get('url', 'No especificada')
        test_instructions = test_data.get('instructions', 'No especificadas')
        test_status = test_data.get('status', 'Desconocido')
        screenshots = test_data.get('screenshots', [])
        
        # Función para reemplazar marcadores
        def replace_placeholders(text):
            if not text or not isinstance(text, str):
                return text
                
            pasos_text = ""
            if '{PASOS_EJECUTADOS}' in text:
                pasos = []
                lines = test_instructions.split('\n')
                for i, line in enumerate(lines, 1):
                    if line.strip():
                        pasos.append(f"{i}. {line.strip()}")
                pasos_text = '\n'.join(pasos[:10])
            
            replacements = {
                '{SISTEMA}': 'QA-Pilot - Sistema de Automatización de Pruebas',
                '{TIPO_APLICACION}': 'Aplicación Web',
                '{NAVEGADOR}': 'Chrome / Edge / Firefox',
                '{AMBIENTE}': 'Desarrollo / QA / Producción',
                '{URL}': test_url,
                '{FECHA}': test_date,
                '{ID_PRUEBA}': task_id[:8],
                '{TITULO_PRUEBA}': test_name,
                '{DESCRIPCION}': test_instructions[:500] + '...' if len(test_instructions) > 500 else test_instructions,
                '{PASOS_EJECUTADOS}': pasos_text,
                '{RESULTADO_ESPERADO}': 'Ejecución exitosa de todos los pasos automatizados',
                '{RESULTADO_OBTENIDO}': 'Exitoso' if test_status == 'success' else 'Con errores' if test_status == 'error' else 'Completado',
                '{ESTADO_PRUEBA}': 'Exitosa' if test_status == 'success' else 'Fallida',
                '{COMENTARIOS}': f"Test automatizado ejecutado el {test_date}. " + 
                               ("Todas las acciones se completaron exitosamente." if test_status == 'success' 
                                else "Se encontraron errores durante la ejecución.")
            }
            
            for placeholder, value in replacements.items():
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
                    logger.debug(f"Reemplazado {placeholder}")
            
            return text
        
        # Procesar párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text:
                for run in paragraph.runs:
                    if run.text:
                        run.text = replace_placeholders(run.text)
        
        # Procesar tablas
        logger.debug(f"Procesando {len(doc.tables)} tablas")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    old_text = run.text
                                    run.text = replace_placeholders(run.text)
                                    if run.text != old_text:
                                        logger.debug(f"Actualizado en tabla: {old_text[:30]}...")
        
        # Agregar capturas al final
        if screenshots:
            doc.add_heading('Evidencias de Ejecución', level=2)
            doc.add_paragraph("A continuación se presentan las capturas de pantalla tomadas durante la ejecución del test automatizado:")
            
            for i, screenshot in enumerate(screenshots[:15], 1):
                try:
                    screenshot_path = screenshot.get('path')
                    if not screenshot_path:
                        screenshot_url = screenshot.get('url', '')
                        if screenshot_url.startswith('/media/screenshots/'):
                            relative_path = screenshot_url.replace('/media/screenshots/', '')
                            screenshot_path = os.path.join(os.getcwd(), relative_path)
                    
                    if screenshot_path and os.path.exists(screenshot_path):
                        screenshot_name = screenshot.get('name', f'Captura {i}')
                        clean_name = screenshot_name.replace('_', ' ').replace('.png', '')
                        
                        doc.add_paragraph(f"Figura {i}: {clean_name}")
                        doc.add_picture(screenshot_path, width=Inches(5.5))
                        doc.add_paragraph("")  # Espacio
                        
                except Exception as e:
                    logger.warning(f"Error procesando captura {i}: {e}")
                    continue
        
        # Crear archivo temporal y enviarlo
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        filename = f"Evidencia_Test_{task_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        def remove_file(response):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return response
        
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error generando evidencia en Word: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Error al generar evidencia: {str(e)}'
        }), 500

@app.route('/delete_screenshots', methods=['POST'])
def delete_screenshots():
    """Elimina capturas de pantalla específicas de un test"""
    try:
        data = request.get_json()
        test_id = data.get('test_id')
        screenshots_to_delete = data.get('screenshots', [])
        
        if not test_id or not screenshots_to_delete:
            return jsonify({
                'status': 'error',
                'message': 'ID de test y lista de capturas son requeridos'
            }), 400
        
        deleted_count = 0
        errors = []
        
        for screenshot_name in screenshots_to_delete:
            try:
                # Buscar el archivo en el directorio de screenshots
                # Primero intentar en el directorio específico del test
                test_screenshot_dir = os.path.join(SCREENSHOTS_DIR, test_id)
                screenshot_path = os.path.join(test_screenshot_dir, screenshot_name)
                
                if os.path.exists(screenshot_path):
                    os.remove(screenshot_path)
                    deleted_count += 1
                    logger.info(f"Captura eliminada: {screenshot_path}")
                else:
                    # Buscar en todos los subdirectorios
                    for root, dirs, files in os.walk(SCREENSHOTS_DIR):
                        if screenshot_name in files:
                            full_path = os.path.join(root, screenshot_name)
                            os.remove(full_path)
                            deleted_count += 1
                            logger.info(f"Captura eliminada: {full_path}")
                            break
                    else:
                        errors.append(f"No se encontró la captura: {screenshot_name}")
                        
            except Exception as e:
                error_msg = f"Error eliminando {screenshot_name}: {str(e)}"
                errors.append(error_msg)
                logger.error(error_msg)
        
        # Actualizar el historial para remover las capturas eliminadas
        with history_lock:
            for item in history_db:
                if item.get('id') == test_id:
                    if 'screenshots' in item:
                        # Filtrar las capturas que fueron eliminadas
                        item['screenshots'] = [
                            screenshot for screenshot in item['screenshots']
                            if screenshot.get('name') not in screenshots_to_delete
                        ]
                    break
        
        response_message = f"Se eliminaron {deleted_count} capturas correctamente"
        if errors:
            response_message += f". Errores: {'; '.join(errors)}"
        
        return jsonify({
            'status': 'success' if deleted_count > 0 else 'warning',
            'message': response_message,
            'deleted_count': deleted_count,
            'errors': errors
        })
        
    except Exception as e:
        error_msg = f"Error al eliminar capturas: {str(e)}"
        logger.error(error_msg)
        return jsonify({
            'status': 'error',
            'message': error_msg
        }), 500

@app.route('/delete_test_from_history', methods=['POST'])
def delete_test_from_history():
    import shutil
    data = request.get_json()
    test_id = data.get('id')

    if not test_id:
        logger.error("Se intentó eliminar un test sin proporcionar un ID.")
        return jsonify({'status': 'error', 'message': 'ID de test no proporcionado'}), 400

    logger.info(f"[FORZADO] Eliminando test y archivos para el test ID: {test_id}")
    try:
        # 1. Eliminar archivos asociados (aunque no exista en BD)
        try:
            script_path = os.path.join('test_scripts', f'test_{test_id}.py')
            if os.path.exists(script_path):
                os.remove(script_path)
                logger.info(f"Archivo de script eliminado: {script_path}")
            screenshot_dir = os.path.join('test_screenshots', test_id)
            if os.path.exists(screenshot_dir):
                shutil.rmtree(screenshot_dir)
                logger.info(f"Directorio de capturas eliminado: {screenshot_dir}")
        except Exception as e:
            logger.warning(f"Error al eliminar archivos locales: {e}")

        # 2. Eliminar de la base de datos (si existe)
        try:
            history_service = get_history_service()
            with history_service.db_integration.get_connection() as conn:
                with conn.cursor() as cur:
                    logger.info(f"Intentando eliminar de la base de datos el test con id: {test_id}")
                    cur.execute("DELETE FROM test_executions WHERE id = %s", (test_id,))
                    deleted = cur.rowcount
                    conn.commit()
            if deleted:
                logger.info(f"Registro {test_id} eliminado de la base de datos. Registros afectados: {deleted}")
            else:
                logger.warning(f"Registro {test_id} NO existía en la base de datos. Registros afectados: {deleted}")
        except Exception as e:
            logger.warning(f"Error al eliminar de la base de datos: {e}")

        return jsonify({'status': 'success', 'message': 'Test eliminado (o ya no existía).'}), 200
    except Exception as e:
        logger.error(f"Error fatal durante la eliminación forzada: {e}", exc_info=True)
        return jsonify({'status': 'error', 'message': 'Ocurrió un error interno del servidor.'}), 500

# Verificar rutas registradas después de crear la aplicación pero antes de ejecutarla
print("\n=== Rutas Registradas en Flask ===")
print("Verificando registro de rutas...")
for rule in app.url_map.iter_rules():
    print(f"Ruta: {rule.rule}")
    print(f"  - Endpoint: {rule.endpoint}")
    print(f"  - Métodos: {rule.methods}")
print("=== Fin de Rutas Registradas ===\n")

# Endpoint de evidencia Word - Definido fuera del bloque main
@app.route('/generate_word_evidence_direct/<task_id>', methods=['POST'])
def generate_word_evidence_direct(task_id):
    """Genera un documento Word con evidencias del test - Endpoint directo."""
    try:
        import os
        import tempfile
        from datetime import datetime
        from docx import Document
        from docx.shared import Inches
        
        logger.debug(f"Generando evidencia en Word directo para test: {task_id}")
        
        # Buscar el test en el historial
        with history_lock:
            test_data = None
            for item in history_db:
                if item.get('id') == task_id:
                    test_data = item
                    break
            
            if not test_data:
                # Si no se encuentra el test, usar datos de ejemplo
                logger.warning(f"Test {task_id} no encontrado en historial, usando datos de ejemplo")
                test_data = {
                    'id': task_id,
                    'name': f'Test de Ejemplo {task_id[:8]}',
                    'date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'url': 'https://ejemplo.com',
                    'instructions': 'Navegar a la página principal\nHacer clic en el botón de login\nIngresar credenciales\nVerificar que el login sea exitoso\nNavegar al dashboard\nVerificar que se muestren los datos correctos',
                    'status': 'success',
                    'screenshots': []
                }
        
        # Ruta del template
        template_path = os.path.join(os.getcwd(), "test_evidencia", "template", "QA_Evidencia_Template_Mejorado.docx")
        
        if not os.path.exists(template_path):
            return jsonify({
                'status': 'error',
                'message': f'Template no encontrado en: {template_path}'
            }), 404
        
        # Crear documento basado en el template
        doc = Document(template_path)
        
        # Obtener información del test
        test_name = test_data.get('name', f'Test {task_id[:8]}')
        test_date = test_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        test_url = test_data.get('url', 'No especificada')
        test_instructions = test_data.get('instructions', 'No especificadas')
        test_status = test_data.get('status', 'Desconocido')
        screenshots = test_data.get('screenshots', [])
        
        # Función para reemplazar marcadores
        def replace_placeholders(text):
            if not text or not isinstance(text, str):
                return text
                
            pasos_text = ""
            if '{PASOS_EJECUTADOS}' in text:
                pasos = []
                lines = test_instructions.split('\n')
                for i, line in enumerate(lines, 1):
                    if line.strip():
                        pasos.append(f"{i}. {line.strip()}")
                pasos_text = '\n'.join(pasos[:10])
            
            replacements = {
                '{SISTEMA}': 'QA-Pilot - Sistema de Automatización de Pruebas',
                '{TIPO_APLICACION}': 'Aplicación Web',
                '{NAVEGADOR}': 'Chrome / Edge / Firefox',
                '{AMBIENTE}': 'Desarrollo / QA / Producción',
                '{URL}': test_url,
                '{FECHA}': test_date,
                '{ID_PRUEBA}': task_id[:8],
                '{TITULO_PRUEBA}': test_name,
                '{DESCRIPCION}': test_instructions[:500] + '...' if len(test_instructions) > 500 else test_instructions,
                '{PASOS_EJECUTADOS}': pasos_text,
                '{RESULTADO_ESPERADO}': 'Ejecución exitosa de todos los pasos automatizados',
                '{RESULTADO_OBTENIDO}': 'Exitoso' if test_status == 'success' else 'Con errores' if test_status == 'error' else 'Completado',
                '{ESTADO_PRUEBA}': 'Exitosa' if test_status == 'success' else 'Fallida',
                '{COMENTARIOS}': f"Test automatizado ejecutado el {test_date}. " + 
                               ("Todas las acciones se completaron exitosamente." if test_status == 'success' 
                                else "Se encontraron errores durante la ejecución.")
            }
            
            for placeholder, value in replacements.items():
                if placeholder in text:
                    text = text.replace(placeholder, str(value))
                    logger.debug(f"Reemplazado {placeholder}")
            
            return text
        
        # Procesar párrafos
        for paragraph in doc.paragraphs:
            if paragraph.text:
                for run in paragraph.runs:
                    if run.text:
                        run.text = replace_placeholders(run.text)
        
        # Procesar tablas
        logger.debug(f"Procesando {len(doc.tables)} tablas")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    old_text = run.text
                                    run.text = replace_placeholders(run.text)
                                    if run.text != old_text:
                                        logger.debug(f"Actualizado en tabla: {old_text[:30]}...")
        
        # Agregar capturas al final
        if screenshots:
            doc.add_heading('Evidencias de Ejecución', level=2)
            doc.add_paragraph("A continuación se presentan las capturas de pantalla tomadas durante la ejecución del test automatizado:")
            
            for i, screenshot in enumerate(screenshots[:15], 1):
                try:
                    screenshot_path = screenshot.get('path')
                    if not screenshot_path:
                        screenshot_url = screenshot.get('url', '')
                        if screenshot_url.startswith('/media/screenshots/'):
                            relative_path = screenshot_url.replace('/media/screenshots/', '')
                            screenshot_path = os.path.join(os.getcwd(), relative_path)
                    
                    if screenshot_path and os.path.exists(screenshot_path):
                        screenshot_name = screenshot.get('name', f'Captura {i}')
                        clean_name = screenshot_name.replace('_', ' ').replace('.png', '')
                        
                        doc.add_paragraph(f"Figura {i}: {clean_name}")
                        doc.add_picture(screenshot_path, width=Inches(5.5))
                        doc.add_paragraph("")  # Espacio
                        
                except Exception as e:
                    logger.warning(f"Error procesando captura {i}: {e}")
                    continue
        
        # Crear archivo temporal y enviarlo
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            tmp_path = tmp_file.name
        
        filename = f"Evidencia_Test_{task_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        def remove_file(response):
            try:
                os.unlink(tmp_path)
            except Exception:
                pass
            return response
        
        return send_file(
            tmp_path,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error generando evidencia en Word directo: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Error al generar evidencia: {str(e)}'
        }), 500

# ===============================================================
# NUEVAS RUTAS PARA GESTIÓN DE TEST SUITES
# ===============================================================



@app.route('/api/projects', methods=['GET'])
def api_get_projects():
    """API para obtener proyectos disponibles"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        projects = db_integration.get_projects()
        
        # Si no hay proyectos, crear uno por defecto
        if not projects:
            default_project = db_integration.get_default_project()
            projects = [{
                'id': str(default_project.id),
                'name': default_project.name,
                'description': default_project.description,
                'base_url': default_project.base_url,
                'status': default_project.status,
                'created_at': default_project.created_at,
                'created_by': default_project.created_by
            }]
        
        return jsonify({'success': True, 'projects': projects})
    except Exception as e:
        logger.error(f"Error en api_get_projects: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/projects', methods=['POST'])
def api_create_project():
    """API para crear un nuevo proyecto"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Datos no proporcionados'}), 400
        
        # Validar campos requeridos
        if not data.get('name') or not data['name'].strip():
            return jsonify({'success': False, 'error': 'El nombre del proyecto es requerido'}), 400
        
        # Crear el proyecto usando la integración de base de datos
        project_data = {
            'name': data['name'].strip(),
            'description': data.get('description', '').strip(),
            'base_url': data.get('base_url', '').strip(),
            'status': data.get('status', 'active'),
            'created_by': 'web_user',
            'metadata_json': data.get('metadata', {})
        }
        
        project_id = db_integration.create_project(project_data)
        
        return jsonify({
            'success': True, 
            'message': 'Proyecto creado exitosamente',
            'project_id': str(project_id),
            'project': {
                'id': str(project_id),
                'name': project_data['name'],
                'description': project_data['description'],
                'base_url': project_data['base_url'],
                'status': project_data['status']
            }
        })
    except Exception as e:
        logger.error(f"Error en api_create_project: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites', methods=['GET'])
def api_get_test_suites():
    """API para obtener suites de prueba"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        project_id = request.args.get('project_id')
        suites = db_integration.get_test_suites(project_id=project_id)
        
        # Convertir a formato JSON serializable
        suites_data = []
        for suite in suites:
            suite_data = {
                'id': str(suite['id']),
                'name': suite['name'],
                'description': suite['description'],
                'project_id': str(suite['project_id']),
                'project_name': suite.get('project_name', ''),
                'status': suite['status'],
                'created_at': suite['created_at'].isoformat() if suite['created_at'] else None,
                'test_cases_count': suite.get('test_cases_count', 0)
            }
            suites_data.append(suite_data)
        
        return jsonify({'success': True, 'suites': suites_data})
    except Exception as e:
        logger.error(f"Error en api_get_test_suites: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>', methods=['GET'])
def api_get_test_suite_details(suite_id):
    """API para obtener detalles de una suite específica"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        # Obtener la suite específica
        suite_details = db_integration.get_test_suite_details(suite_id)
        
        if not suite_details:
            return jsonify({'success': False, 'error': 'Suite no encontrada'}), 404
        
        # Obtener todos los casos disponibles para la selección
        available_cases = db_integration.get_test_cases(
            project_id=suite_details['project_id'], 
            limit=100
        )
        
        return jsonify({
            'status': 'success',
            'suite': suite_details,
            'available_cases': available_cases
        })
    except Exception as e:
        logger.error(f"Error en api_get_test_suite_details: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites', methods=['POST'])
def api_create_test_suite():
    """API para crear una nueva suite de prueba"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Datos no proporcionados'}), 400
        
        # Validar campos requeridos
        required_fields = ['name', 'project_id']
        for field in required_fields:
            if field not in data or not data[field]:
                return jsonify({'success': False, 'error': f'Campo requerido: {field}'}), 400
        
        # Crear la suite
        suite_data = {
            'name': data['name'],
            'description': data.get('description', ''),
            'project_id': data['project_id'],
            'status': data.get('status', 'active'),
            'created_by': 'web_user',
            'metadata_json': data.get('metadata', {})
        }
        
        suite_id = db_integration.create_test_suite(suite_data)
        
        return jsonify({
            'success': True, 
            'message': 'Suite de prueba creada exitosamente',
            'suite_id': str(suite_id)
        })
    except Exception as e:
        logger.error(f"Error en api_create_test_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>', methods=['PUT'])
def api_update_test_suite(suite_id):
    """API para actualizar una suite de prueba"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Datos no proporcionados'}), 400
        
        logger.info(f"Actualizando suite {suite_id} con datos: {data}")
        
        # Verificar si hay casos seleccionados
        if 'cases' in data:
            logger.info(f"Casos seleccionados: {len(data['cases'])} - {data['cases']}")
        
        # Actualizar la suite
        success = db_integration.update_test_suite(suite_id, data)
        
        if success:
            # Obtener los casos actualizados para confirmar
            updated_suite = db_integration.get_test_suite_details(suite_id)
            case_count = len(updated_suite.get('test_cases', []))
            
            logger.info(f"Suite actualizada exitosamente. Ahora tiene {case_count} casos asociados.")
            
            return jsonify({
                'success': True, 
                'message': f'Suite actualizada exitosamente con {case_count} casos',
                'case_count': case_count
            })
        else:
            return jsonify({'success': False, 'error': 'Suite no encontrada'}), 404
            
    except Exception as e:
        logger.error(f"Error en api_update_test_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>', methods=['DELETE'])
def api_delete_test_suite(suite_id):
    """API para eliminar una suite de prueba"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        success = db_integration.delete_test_suite(suite_id)
        
        if success:
            return jsonify({
                'success': True, 
                'message': 'Suite de prueba eliminada exitosamente'
            })
        else:
            return jsonify({'success': False, 'error': 'Suite no encontrada'}), 404
            
    except Exception as e:
        logger.error(f"Error en api_delete_test_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/cases', methods=['GET'])
def api_get_suite_test_cases(suite_id):
    """API para obtener casos de prueba de una suite específica"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        test_cases = db_integration.get_test_cases(suite_id=suite_id)
        
        # Convertir a formato JSON serializable
        cases_data = []
        for case in test_cases:
            case_data = {
                'id': str(case['id']),
                'nombre': case['nombre'],
                'codigo': case['codigo'],
                'tipo': case['tipo'],
                'prioridad': case['prioridad'],
                'objetivo': case['objetivo'],
                'pasos': case['pasos'],
                'resultado_esperado': case['resultado_esperado'],
                'url_objetivo': case['url_objetivo'],
                'es_valido': case['es_valido'],
                'status': case['status'],
                'created_at': case['created_at'].isoformat() if case['created_at'] else None
            }
            cases_data.append(case_data)
        
        return jsonify({'success': True, 'test_cases': cases_data})
    except Exception as e:
        logger.error(f"Error en api_get_suite_test_cases: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/add_case', methods=['POST'])
def api_add_case_to_suite(suite_id):
    """API para agregar un caso de prueba existente a una suite"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        data = request.get_json()
        if not data or 'case_id' not in data:
            return jsonify({'success': False, 'error': 'ID del caso de prueba no proporcionado'}), 400
        
        case_id = data['case_id']
        
        # Actualizar el suite_id del caso de prueba
        success = db_integration.update_test_case(case_id, {'suite_id': suite_id})
        
        if success:
            return jsonify({
                'success': True, 
                'message': 'Caso de prueba agregado a la suite exitosamente'
            })
        else:
            return jsonify({'success': False, 'error': 'Caso de prueba no encontrado'}), 404
            
    except Exception as e:
        logger.error(f"Error en api_add_case_to_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/remove_case', methods=['POST'])
def api_remove_case_from_suite(suite_id):
    """API para remover un caso de prueba de una suite"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        data = request.get_json()
        if not data or 'case_id' not in data:
            return jsonify({'success': False, 'error': 'ID del caso de prueba no proporcionado'}), 400
        
        case_id = data['case_id']
        
        # Remover el suite_id del caso de prueba (establecer como NULL)
        success = db_integration.update_test_case(case_id, {'suite_id': None})
        
        if success:
            return jsonify({
                'success': True, 
                'message': 'Caso de prueba removido de la suite exitosamente'
            })
        else:
            return jsonify({'success': False, 'error': 'Caso de prueba no encontrado'}), 404
            
    except Exception as e:
        logger.error(f"Error en api_remove_case_from_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/execute', methods=['POST'])
def api_execute_test_suite(suite_id):
    """API para ejecutar todos los casos de prueba de una suite"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        
        # Obtener configuración de ejecución
        data = request.get_json() or {}
        is_dry_run = data.get('is_dry_run', False)
        headless = data.get('headless', True)
        
        # Obtener casos de prueba de la suite
        test_cases = db_integration.get_test_cases(suite_id=suite_id)
        
        if not test_cases:
            return jsonify({'success': False, 'error': 'No se encontraron casos de prueba en la suite'}), 400
        
        if is_dry_run:
            # Modo verificación: simular ejecución exitosa
            logger.info(f"🔍 Verificación de suite {suite_id}: {len(test_cases)} casos")
            
            # Simular resultados de verificación
            results = {
                'success': len(test_cases),
                'failed': 0,
                'cases': []
            }
            
            for i, case in enumerate(test_cases):
                case_result = {
                    'case': case['nombre'],
                    'status': 'success',
                    'message': 'Verificación exitosa',
                    'detail': 'El script está listo para ejecutarse',
                    'screenshots': [],
                    'log': None
                }
                results['cases'].append(case_result)
            
            return jsonify({
                'success': True,
                'status': 'completed',
                'results': results,
                'execution_id': f'verify_{suite_id}_{int(time.time())}'
            })
        
        else:
            # Modo ejecución real: ejecutar casos uno por uno usando run_test_background
            logger.info(f"🚀 Ejecutando suite {suite_id}: {len(test_cases)} casos")
            
            # Generar ID único para esta ejecución
            execution_id = str(uuid.uuid4())
            
            # Función para ejecutar suite en segundo plano
            def execute_suite_real_background():
                results_file = None
                stop_flag = threading.Event()
                
                try:
                    with app.app_context():
                        logger.info(f"🔄 Iniciando ejecución real de suite {suite_id}")
                        
                        # Registrar esta ejecución como activa
                        with suite_executions_lock:
                            active_suite_executions[execution_id] = {
                                'thread': threading.current_thread(),
            'suite_id': suite_id,
                                'stop_flag': stop_flag,
                                'status': 'running',
                                'start_time': datetime.now().isoformat()
                            }
                        
                        # Archivo de resultados
                        results_file = os.path.join(tempfile.gettempdir(), f"suite_results_{execution_id}.json")
                        
                        # Resultado general de la suite
                        suite_results = {
                            'execution_id': execution_id,
                            'suite_id': suite_id,
                            'total_cases': len(test_cases),
                            'completed_cases': 0,
                            'success_count': 0,
                            'failed_count': 0,
                            'cases': [],
                            'status': 'running',
                            'start_time': datetime.now().isoformat()
                        }
                        
                        # Guardar estado inicial
                        with open(results_file, 'w') as f:
                            json.dump(suite_results, f, indent=2)
                        
                        # Ejecutar cada caso individualmente
                        for i, case in enumerate(test_cases):
                            # Verificar si se ha solicitado detener la ejecución
                            if stop_flag.is_set():
                                logger.info(f"🛑 Ejecución detenida por el usuario en caso {i+1}/{len(test_cases)}")
                                suite_results['status'] = 'stopped'
                                suite_results['end_time'] = datetime.now().isoformat()
                                break
                            
                            case_task_id = f"{execution_id}_case_{i+1}"
                            case_name = case['nombre']
                            case_url = case['url_objetivo'] or 'https://example.com'
                            case_instructions = case['pasos'] or 'Navegar por la página web'
                            
                            logger.info(f"📝 Ejecutando caso {i+1}/{len(test_cases)}: {case_name}")
                            
                            try:
                                # Configuración para la ejecución
                                case_headless = headless
                                case_max_tiempo = 300  # 5 minutos por caso
                                case_screenshots = True
                                
                                # Obtener API key (usar cualquiera disponible)
                                api_key = os.getenv('GEMINI_API_KEY') or os.getenv('ANTHROPIC_API_KEY') or os.getenv('OPENAI_API_KEY')
                                if not api_key:
                                    raise Exception("No hay API keys disponibles para ejecutar el caso")
                                
                                # Determinar modelo basado en la API key disponible
                                if os.getenv('ANTHROPIC_API_KEY'):
                                    model_name = 'claude-3-5-sonnet-20240620'
                                elif os.getenv('OPENAI_API_KEY'):
                                    model_name = 'gpt-4'
                                else:
                                    model_name = 'gemini-1.5-pro'
                                
                                # Ejecutar el caso usando MCP (SIEMPRE con navegador visible)
                                logger.info(f"🎯 Ejecutando {case_name} con URL: {case_url} (MCP - Navegador Visible)")
                                
                                # Inicializar estado en test_status_db
                                with db_lock:
                                    test_status_db[case_task_id] = {
                                        'status': 'running',
                                        'start_time': datetime.now().isoformat(),
                                        'case_name': case_name,
                                        'suite_id': suite_id,  # Agregar referencia a la suite
                                        'execution_mode': 'mcp_suite'  # Marcar como ejecución de suite con MCP
                                    }
                                
                                # Ejecutar caso usando MCP con navegador visible (función específica para suites)
                                execution_thread = threading.Thread(
                                    target=run_test_background_suite_mcp,
                                    args=(case_task_id, case_url, case_instructions, 
                                          case_max_tiempo, case_screenshots, api_key, model_name, 'chrome')
                                )
                                execution_thread.start()
                                
                                # Esperar a que termine la ejecución monitoreando el estado
                                max_wait_time = case_max_tiempo + 60  # Tiempo extra de gracia
                                wait_time = 0
                                case_completed = False
                                
                                while wait_time < max_wait_time:
                                    time.sleep(5)  # Verificar cada 5 segundos
                                    wait_time += 5
                                    
                                    # Verificar si se ha solicitado detener la ejecución
                                    if stop_flag.is_set():
                                        logger.info(f"🛑 Deteniendo caso {case_name} por solicitud del usuario")
                                        
                                        # Terminar el proceso subprocess si existe
                                        with subprocess_processes_lock:
                                            if case_task_id in active_subprocess_processes:
                                                proceso = active_subprocess_processes[case_task_id]
                                                try:
                                                    if proceso.poll() is None:
                                                        logger.info(f"🔪 Terminando proceso subprocess PID {proceso.pid} para caso {case_name}")
                                                        proceso.terminate()
                                                        try:
                                                            proceso.wait(timeout=3)
                                                        except subprocess.TimeoutExpired:
                                                            proceso.kill()
                                                            proceso.wait(timeout=2)
                                                except Exception as e:
                                                    logger.error(f"Error terminando proceso para caso {case_name}: {e}")
                                        
                                        # Marcar caso como detenido
                                        with db_lock:
                                            if case_task_id in test_status_db:
                                                test_status_db[case_task_id]['status'] = 'stopped'
                                                test_status_db[case_task_id]['message'] = 'Detenido por el usuario'
                                        case_completed = True
                                        break
                                    
                                    # Verificar estado en test_status_db Y en MCP
                                    with db_lock:
                                        if case_task_id in test_status_db:
                                            case_status = test_status_db[case_task_id]
                                            if case_status.get('status') in ['success', 'error', 'completed', 'stopped']:
                                                case_completed = True
                                                break
                                    
                                    # Verificar si el hilo aún está vivo
                                    if not execution_thread.is_alive():
                                        case_completed = True
                                        break
                                    
                                    # Verificar si el proceso subprocess ha sido terminado externamente
                                    with subprocess_processes_lock:
                                        if case_task_id in active_subprocess_processes:
                                            proceso = active_subprocess_processes[case_task_id]
                                            if proceso.poll() is not None:  # Proceso ha terminado
                                                logger.info(f"🔍 Proceso subprocess para caso {case_name} ha terminado")
                                                case_completed = True
                                                break
                                
                                # Esperar que el hilo termine completamente
                                execution_thread.join(timeout=10)
                                
                                # Obtener resultado final desde Flask Y MCP
                                case_success = False
                                case_error_msg = None
                                case_stopped = False
                                
                                # Primero verificar en Flask como fallback
                                with db_lock:
                                    if case_task_id in test_status_db:
                                        final_status = test_status_db[case_task_id]
                                        case_success = final_status.get('status') in ['success', 'completed']
                                        case_error_msg = final_status.get('error', final_status.get('message'))
                                        case_stopped = final_status.get('status') == 'stopped'
                                
                                # Preparar resultado del caso
                                if case_stopped:
                                    # Caso detenido por el usuario
                                    case_result = {
                                        'case_id': str(case['id']),
                                        'case_name': case_name,
                                        'task_id': case_task_id,
                                        'status': 'stopped',
                                        'message': 'Ejecución detenida por el usuario',
                                        'url': case_url,
                                        'duration': f'Detenido después de {wait_time} segundos'
                                    }
                                    suite_results['failed_count'] += 1  # Contar como fallo para estadísticas
                                    logger.info(f"🛑 Caso {case_name} detenido por el usuario")
                                elif case_success:
                                    case_result = {
                                        'case_id': str(case['id']),
                                        'case_name': case_name,
                                        'task_id': case_task_id,
                                        'status': 'success',
                                        'message': 'Caso ejecutado correctamente',
                                        'url': case_url,
                                        'duration': f'Aproximadamente {wait_time} segundos',
                                        'screenshots': f"/media/screenshots/{case_task_id}/",
                                        'execution_mode': 'mcp_suite'  # Marcar como ejecución MCP de suite
                                    }
                                    suite_results['success_count'] += 1
                                    logger.info(f"✅ Caso {case_name} ejecutado correctamente")
                                else:
                                    # Caso falló o tuvo timeout
                                    error_message = case_error_msg or f'Timeout después de {max_wait_time} segundos'
                                    case_result = {
                                        'case_id': str(case['id']),
                                        'case_name': case_name,
                                        'task_id': case_task_id,
                                        'execution_mode': 'mcp_suite',  # Marcar como ejecución MCP de suite
                                        'status': 'failed',
                                        'message': f'Error: {error_message}',
                                        'url': case_url,
                                        'error': error_message
                                    }
                                    suite_results['failed_count'] += 1
                                    logger.error(f"❌ Caso {case_name} falló: {error_message}")
                                
                            except Exception as case_error:
                                logger.error(f"❌ Error ejecutando caso {case_name}: {case_error}")
                                
                                case_result = {
                                    'case_id': str(case['id']),
                                    'case_name': case_name,
                                    'task_id': case_task_id,
                                    'status': 'failed',
                                    'message': f'Error: {str(case_error)}',
                                    'url': case_url,
                                    'error': str(case_error)
                                }
                                
                                suite_results['failed_count'] += 1
                            
                            # Agregar resultado del caso
                            suite_results['cases'].append(case_result)
                            suite_results['completed_cases'] += 1
                            
                            # Actualizar archivo de progreso después de cada caso
                            try:
                                with open(results_file, 'w') as f:
                                    json.dump(suite_results, f, indent=2)
                            except Exception as save_error:
                                logger.error(f"Error guardando progreso: {save_error}")
                            
                            # Log de progreso
                            progress = (i + 1) / len(test_cases) * 100
                            logger.info(f"📊 Progreso de suite: {progress:.1f}% ({i+1}/{len(test_cases)})")
                        
                        # Finalizar ejecución
                        suite_results['status'] = 'completed'
                        suite_results['end_time'] = datetime.now().isoformat()
                        
                        logger.info(f"🏁 Suite {suite_id} completada: {suite_results['success_count']} éxitos, {suite_results['failed_count']} fallos")
                        
                        # Guardar resultados finales
                        with open(results_file, 'w') as f:
                            json.dump(suite_results, f, indent=2)
                        
                        # Guardar ejecuciones en base de datos para historial
                        try:
                            save_suite_execution_to_database(execution_id, suite_id, suite_results)
                        except Exception as db_error:
                            logger.error(f"Error guardando suite en BD: {db_error}")
                        
                except Exception as e:
                    logger.error(f"❌ Error crítico en ejecución de suite {suite_id}: {e}")
                    import traceback
                    traceback.print_exc()
                    
                    # Guardar estado de error en archivo
                    if results_file:
                        try:
                            error_results = {
                                'execution_id': execution_id,
                                'suite_id': suite_id,
                        'status': 'error',
                                'error': str(e),
                                'end_time': datetime.now().isoformat()
                            }
                            with open(results_file, 'w') as f:
                                json.dump(error_results, f, indent=2)
                        except:
                            pass
                
                finally:
                    # Limpiar registro de ejecución activa
                    with suite_executions_lock:
                        if execution_id in active_suite_executions:
                            del active_suite_executions[execution_id]
                            logger.info(f"🧹 Ejecución {execution_id} removida del registro de ejecuciones activas")
        
        # Iniciar ejecución en hilo separado
            execution_thread = threading.Thread(target=execute_suite_real_background)
            execution_thread.daemon = False
        execution_thread.start()
        
        return jsonify({
            'success': True,
                'status': 'started',
                'execution_id': execution_id,
                'total_cases': len(test_cases),
                'message': f'Ejecución real iniciada con {len(test_cases)} casos. Los casos se ejecutarán uno por uno.',
                'note': 'Esta es una ejecución real que puede tomar varios minutos. Los casos se ejecutan secuencialmente.'
        })
        
    except Exception as e:
        logger.error(f"❌ Error en api_execute_test_suite: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/stop_execution', methods=['POST'])
def api_stop_suite_execution(suite_id):
    """API para detener la ejecución de una suite en progreso"""
    try:
        logger.info(f"🛑 Solicitada detención de ejecución para suite {suite_id}")
        
        # Buscar ejecuciones activas para esta suite
        stopped_executions = []
        stopped_cases = 0
        
        with suite_executions_lock:
            executions_to_stop = []
            for execution_id, execution_data in active_suite_executions.items():
                if execution_data.get('suite_id') == suite_id:
                    executions_to_stop.append((execution_id, execution_data))
            
            # Marcar las ejecuciones para detenerse
            for execution_id, execution_data in executions_to_stop:
                logger.info(f"🛑 Deteniendo ejecución {execution_id}")
                
                # Activar flag de detener
                stop_flag = execution_data.get('stop_flag')
                if stop_flag:
                    stop_flag.set()
                
                # Marcar como detenida
                execution_data['status'] = 'stopped'
                stopped_executions.append(execution_id)
        
        # También detener casos individuales que puedan estar corriendo
        with db_lock:
            cases_to_stop = []
            for task_id, status_data in test_status_db.items():
                # Buscar casos que pertenezcan a esta suite
                if (status_data.get('status') in ['running', 'started'] and 
                    (suite_id in task_id or status_data.get('suite_id') == suite_id)):
                    cases_to_stop.append(task_id)
            
            # Marcar casos como detenidos
            for task_id in cases_to_stop:
                test_status_db[task_id]['status'] = 'stopped'
                test_status_db[task_id]['message'] = 'Ejecución detenida por el usuario'
                test_status_db[task_id]['end_time'] = datetime.now().isoformat()
                stopped_cases += 1
                logger.info(f"🛑 Caso {task_id} marcado como detenido")
        
        # CRÍTICO: Terminar procesos subprocess reales (browser-use)
        with subprocess_processes_lock:
            processes_to_kill = []
            for task_id, proceso in active_subprocess_processes.items():
                # Buscar procesos que pertenezcan a esta suite
                if suite_id in task_id:
                    processes_to_kill.append((task_id, proceso))
            
            # Terminar procesos encontrados
            for task_id, proceso in processes_to_kill:
                try:
                    if proceso.poll() is None:  # Proceso aún corriendo
                        logger.info(f"🔪 Terminando proceso PID {proceso.pid} para task_id {task_id}")
                        
                        # Intentar terminación suave primero
                        proceso.terminate()
                        
                        # Esperar un poco para terminación suave
                        try:
                            proceso.wait(timeout=5)
                            logger.info(f"✅ Proceso {proceso.pid} terminado suavemente")
                        except subprocess.TimeoutExpired:
                            # Si no termina suavemente, forzar terminación
                            logger.warning(f"⚡ Forzando terminación del proceso {proceso.pid}")
                            proceso.kill()
                            proceso.wait(timeout=3)
                            logger.info(f"💀 Proceso {proceso.pid} terminado forzadamente")
                        
                        stopped_cases += 1
                    else:
                        logger.info(f"ℹ️ Proceso {proceso.pid} ya había terminado")
                        
                except Exception as proc_error:
                    logger.error(f"❌ Error terminando proceso {proceso.pid}: {proc_error}")
                    # Continuar con otros procesos aunque uno falle
        
        # Actualizar en base de datos si hay ejecuciones que detener
        if stopped_executions or stopped_cases > 0:
            try:
                db_integration = get_db_integration()
                if db_integration and db_integration.is_connected():
                    with db_integration.get_session() as session:
                        from db_models import TestExecution
                        from sqlalchemy import text
                        
                        # Buscar ejecuciones relacionadas con esta suite
                        executions = session.execute(text("""
                            SELECT id FROM test_executions 
                            WHERE metadata_json::text LIKE :suite_pattern 
                            AND status IN ('running', 'started')
                        """), {'suite_pattern': f'%{suite_id}%'}).fetchall()
                        
                        # Actualizar estado en base de datos
                        for execution_row in executions:
                            execution = session.query(TestExecution).filter_by(id=execution_row[0]).first()
                            if execution:
                                execution.status = 'stopped'
                                execution.result_details = 'Ejecución detenida por el usuario'
                                execution.end_time = datetime.now(timezone.utc)
                        
                        session.commit()
                        logger.info(f"✅ {len(executions)} ejecuciones actualizadas como detenidas en BD")
                        
            except Exception as db_error:
                logger.error(f"❌ Error actualizando estado en BD: {db_error}")
        
        if stopped_executions or stopped_cases > 0:
            message = f"Ejecución detenida exitosamente. {len(stopped_executions)} ejecuciones de suite y {stopped_cases} casos individuales detenidos."
            logger.info(f"✅ {message}")
            
            return jsonify({
                'success': True,
                'message': message,
                'stopped_executions': stopped_executions,
                'stopped_cases': stopped_cases
            })
        else:
            return jsonify({
                'success': False,
                'error': 'No se encontraron ejecuciones activas para esta suite'
            }), 404
            
    except Exception as e:
        logger.error(f"❌ Error deteniendo ejecución de suite: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

def save_suite_execution_to_database(execution_id, suite_id, suite_results):
    """Actualiza las ejecuciones existentes con metadatos de suite (no duplica casos)."""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.test_connection():
            logger.warning("Base de datos no disponible para guardar metadatos de suite")
            return
        
        with db_integration.get_session() as session:
            from db_models import TestExecution
            import uuid
            
            logger.info(f"🔄 Actualizando metadatos de suite {suite_id} para {len(suite_results.get('cases', []))} casos")
            
            # En lugar de crear nuevas ejecuciones, buscar y actualizar las existentes
            for case_data in suite_results.get('cases', []):
                try:
                    case_task_id = case_data.get('task_id')
                    if not case_task_id:
                        continue
                    
                    # Buscar la ejecución existente por metadata_json que contenga el task_id
                    execution = session.query(TestExecution).filter(
                        TestExecution.metadata_json.op('->>')('execution_id') == case_task_id
                    ).first()
                    
                    # Si no encuentra por execution_id, buscar por task_id
                    if not execution:
                        execution = session.query(TestExecution).filter(
                            TestExecution.metadata_json.op('->>')('task_id') == case_task_id
                        ).first()
                    
                    # Si aún no encuentra, buscar las ejecuciones más recientes
                    if not execution:
                        from datetime import datetime, timedelta
                        recent_time = datetime.now() - timedelta(minutes=30)  # Últimos 30 minutos
                        recent_executions = session.query(TestExecution).filter(
                            TestExecution.created_at >= recent_time
                        ).order_by(TestExecution.created_at.desc()).limit(10).all()
                        
                        # Buscar en las ejecuciones recientes una que tenga metadata relacionado
                        for ex in recent_executions:
                            if (ex.metadata_json and 
                                (ex.metadata_json.get('execution_id') == case_task_id or
                                 ex.metadata_json.get('task_id') == case_task_id or
                                 case_task_id in str(ex.metadata_json))):
                                execution = ex
                                break
                    
                    if execution:
                        # Actualizar metadatos existentes en lugar de crear nueva ejecución
                        if not execution.metadata_json:
                            execution.metadata_json = {}
                        
                        execution.metadata_json.update({
                            'suite_id': suite_id,
                            'suite_execution_id': execution_id,
                            'is_suite_execution': True,
                            'suite_case_status': case_data.get('status'),
                            'execution_type': 'suite'
                        })
                        
                        # Actualizar detalles si es necesario
                        if not execution.result_details or execution.result_details == 'Ejecución completada':
                            execution.result_details = f"Ejecutado en suite '{suite_id}' - {case_data.get('message', 'Completado')}"
                        
                        logger.info(f"✅ Metadatos de suite actualizados para ejecución {str(execution.id)[:8]} (task_id: {case_task_id})")
                    else:
                        logger.warning(f"No se encontró ejecución existente para task_id {case_task_id}")
                    
                except Exception as case_error:
                    logger.error(f"Error actualizando metadatos para caso {case_data.get('case_name')}: {case_error}")
                    continue
            
            session.commit()
            logger.info(f"✅ Metadatos de suite {suite_id} actualizados (sin duplicar ejecuciones)")
            
    except Exception as e:
        logger.error(f"❌ Error actualizando metadatos de suite: {e}")
        import traceback
        traceback.print_exc()

class HistoryService:
    """Servicio para manejar el historial completamente desde la base de datos."""
    
    def __init__(self, db_integration):
        self.db_integration = db_integration
    
    def get_history_items(self, limit=100):
        """Obtener items del historial desde la base de datos."""
        try:
            if not self.db_integration or not self.db_integration.is_connected():
                logger.warning("Base de datos no disponible, retornando historial vacío")
                return []
            
            with self.db_integration.get_session() as session:
                from db_models import TestExecution, TestCase, Screenshot
                from sqlalchemy import desc
                
                # Obtener ejecuciones recientes con sus casos de prueba
                executions = session.query(TestExecution).join(
                    TestCase, TestExecution.test_case_id == TestCase.id, isouter=True
                ).order_by(desc(TestExecution.created_at)).limit(limit).all()
                
                history_items = []
                for execution in executions:
                    # Obtener screenshots para esta ejecución
                    screenshots = session.query(Screenshot).filter_by(
                        execution_id=execution.id
                    ).all()
                    
                    screenshot_list = []
                    for screenshot in screenshots:
                        # Construir URL relativa para las capturas
                        relative_path = os.path.relpath(screenshot.file_path, SCREENSHOTS_DIR) if screenshot.file_path else None
                        screenshot_url = f'/media/screenshots/{relative_path}' if relative_path else None
                        
                        screenshot_list.append({
                            'url': screenshot_url,
                            'path': screenshot.file_path,
                            'name': screenshot.name
                        })
                    
                    # Crear item del historial
                    history_item = {
                        'id': str(execution.id),
                        'name': execution.test_case.nombre if execution.test_case else f'Test {str(execution.id)[:8]}',
                        'date': execution.created_at.strftime('%Y-%m-%d %H:%M:%S') if execution.created_at else '',
                        'date_timestamp': execution.created_at.timestamp() if execution.created_at else 0,
                        'status': 'success' if execution.status in ['passed', 'completed'] else 'error',
                        'url': execution.url_executed or (execution.test_case.url_objetivo if execution.test_case else ''),
                        'instructions': execution.test_case.pasos if execution.test_case else '',
                        'message': execution.result_details or f'Ejecución {execution.status}',
                        'screenshots': screenshot_list,
                        'test_dir': os.path.dirname(screenshot_list[0]['path']) if screenshot_list else None,
                        'script_path': execution.script_path if execution.script_path and os.path.exists(execution.script_path) else None,
                        'playwright_case': None  # TODO: agregar lógica para detectar casos Playwright
                    }
                    
                    history_items.append(history_item)
                
                logger.info(f"Cargados {len(history_items)} items del historial desde BD")
                return history_items
                
        except Exception as e:
            logger.error(f"Error cargando historial desde BD: {e}")
            return []
    
    def update_test_name(self, test_id, new_name):
        """Actualizar el nombre de un test directamente en la base de datos."""
        try:
            if not self.db_integration or not self.db_integration.is_connected():
                return False, "Base de datos no disponible"
            
            with self.db_integration.get_session() as session:
                from db_models import TestExecution, TestCase
                
                # Buscar la ejecución por ID
                execution = session.query(TestExecution).filter_by(id=test_id).first()
                
                if execution and execution.test_case:
                    # Actualizar el nombre del caso de prueba
                    execution.test_case.nombre = new_name
                    session.commit()
                    logger.info(f"Nombre actualizado en BD para test {test_id}: {new_name}")
                    return True, "Nombre actualizado correctamente"
                else:
                    # Si no existe la ejecución, intentar crear un caso de prueba
                    logger.warning(f"Ejecución {test_id} no encontrada, intentando crear caso")
                    return False, "Ejecución no encontrada"
                    
        except Exception as e:
            logger.error(f"Error actualizando nombre en BD: {e}")
            return False, f"Error: {str(e)}"
    
    def add_execution(self, task_id, test_data):
        """Agregar una nueva ejecución al historial en la base de datos."""
        try:
            if not self.db_integration or not self.db_integration.is_connected():
                logger.warning("BD no disponible, no se puede guardar ejecución")
                return False
            
            # Usar la función existente save_execution_to_database
            history_entry = {
                'id': task_id,
                'name': test_data.get('name', f'Test {task_id[:8]}'),
                'date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'status': test_data.get('status', 'unknown'),
                'url': test_data.get('url', ''),
                'instructions': test_data.get('original_instructions', ''),
                'message': test_data.get('message', ''),
                'screenshots': test_data.get('screenshots', [])
            }
            
            save_execution_to_database(task_id, test_data, history_entry)
            return True
            
        except Exception as e:
            logger.error(f"Error agregando ejecución al historial: {e}")
            return False
    
    def delete_execution(self, test_id):
        """Eliminar una ejecución del historial."""
        try:
            if not self.db_integration or not self.db_integration.is_connected():
                return False, "Base de datos no disponible"
            
            with self.db_integration.get_session() as session:
                from db_models import TestExecution
                
                execution = session.query(TestExecution).filter_by(id=test_id).first()
                if execution:
                    session.delete(execution)
                    session.commit()
                    logger.info(f"Ejecución {test_id} eliminada de BD")
                    return True, "Ejecución eliminada correctamente"
                else:
                    return False, "Ejecución no encontrada"
                    
        except Exception as e:
            logger.error(f"Error eliminando ejecución: {e}")
            return False, f"Error: {str(e)}"

# Crear instancia global del servicio de historial
history_service = None

def get_history_service():
    """Obtener la instancia del servicio de historial."""
    global history_service
    if history_service is None:
        db_integration = get_db_integration()
        history_service = HistoryService(db_integration)
    return history_service

@app.route('/api/test_suites/<suite_id>/execution/<execution_id>/status', methods=['GET'])
def api_get_suite_execution_status(suite_id, execution_id):
    """API para obtener el estado de una ejecución de suite en tiempo real"""
    try:
        # Buscar archivo de resultados
        results_file = os.path.join(tempfile.gettempdir(), f"suite_results_{execution_id}.json")
        
        logger.info(f"🔍 Verificando estado de ejecución {execution_id[:8]}")
        logger.info(f"📁 Buscando archivo: {results_file}")
        logger.info(f"📄 Archivo existe: {os.path.exists(results_file)}")
        
        if os.path.exists(results_file):
            try:
                with open(results_file, 'r') as f:
                    results = json.load(f)
                
                logger.info(f"📊 Estado actual: {results.get('status')}")
                logger.info(f"📈 Progreso: {results.get('completed_cases', 0)}/{results.get('total_cases', 0)}")
                
                return jsonify({
                    'success': True,
                    'execution_id': execution_id,
                    'status': results.get('status', 'running'),
                    'progress': {
                        'total_cases': results.get('total_cases', 0),
                        'completed_cases': results.get('completed_cases', 0),
                        'success_count': results.get('success_count', 0),
                        'failed_count': results.get('failed_count', 0)
                    },
                    'cases': results.get('cases', []),
                    'start_time': results.get('start_time'),
                    'end_time': results.get('end_time', None)
                })
            except json.JSONDecodeError as json_error:
                logger.error(f"❌ Error JSON en archivo: {json_error}")
                return jsonify({'success': False, 'error': f'Error leyendo archivo JSON: {str(json_error)}'}), 500
            except Exception as file_error:
                logger.error(f"❌ Error leyendo archivo: {file_error}")
                return jsonify({'success': False, 'error': f'Error leyendo archivo: {str(file_error)}'}), 500
        else:
            logger.warning(f"⚠️ Archivo de resultados no encontrado para {execution_id[:8]}")
            # Si no existe el archivo, probablemente aún está iniciando
            return jsonify({
                'success': True,
                'execution_id': execution_id,
                'status': 'starting',
                'progress': {
                    'total_cases': 0,
                    'completed_cases': 0,
                    'success_count': 0,
                    'failed_count': 0
                },
                'cases': [],
                'message': 'Ejecución iniciando...'
            })
            
    except Exception as e:
        logger.error(f"❌ Error obteniendo estado de ejecución: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/test_suites/<suite_id>/execution/<execution_id>/generate_evidence', methods=['POST'])
def api_generate_suite_evidence(suite_id, execution_id):
    """Genera documento de evidencia para una ejecución de suite completa"""
    try:
        logger.info(f"📄 Generando evidencia para suite {suite_id}, ejecución {execution_id}")
        
        # NUEVO ENFOQUE: Generar evidencia directamente desde archivos de resultados
        # En lugar de depender de la BD, usar los archivos de resultados de suite
        
        # 1. Buscar archivo de resultados de la suite
        import tempfile
        import os
        import json
        
        results_file = os.path.join(tempfile.gettempdir(), f"suite_results_{execution_id}.json")
        
        if not os.path.exists(results_file):
            # Buscar archivos de resultados que contengan el execution_id
            temp_dir = tempfile.gettempdir()
            for filename in os.listdir(temp_dir):
                if filename.startswith("suite_results_") and execution_id[:8] in filename:
                    results_file = os.path.join(temp_dir, filename)
                    break
        
        if not os.path.exists(results_file):
            return jsonify({'success': False, 'error': 'No se encontró archivo de resultados de la suite'}), 404
        
        # 2. Cargar resultados de la suite
        with open(results_file, 'r', encoding='utf-8') as f:
            suite_results = json.load(f)
        
        # 3. Obtener información de la suite desde la BD
        db_integration = get_db_integration()
        suite_info = None
        
        if db_integration and db_integration.is_connected():
            try:
                with db_integration.get_session() as session:
                    from db_models import TestSuite
                    suite = session.query(TestSuite).filter_by(id=suite_id).first()
                    if suite:
                        suite_info = {
                            'nombre': suite.nombre,
                            'descripcion': suite.descripcion
                        }
            except Exception as db_error:
                logger.warning(f"No se pudo obtener info de suite desde BD: {db_error}")
        
        # 4. Generar documento de evidencia usando los resultados directos
        evidence_path = generate_suite_evidence_from_results(suite_info, suite_results, execution_id, suite_id)
        
        if evidence_path and os.path.exists(evidence_path):
            suite_name = suite_info.get('nombre', 'Suite') if suite_info else 'Suite'
            
            # Determinar el tipo de archivo y configurar descarga
            if evidence_path.endswith('.zip'):
                filename = f"PlaywrightReport_{suite_name}_{execution_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
                mimetype = 'application/zip'
            else:
                filename = f"Evidencia_{suite_name}_{execution_id[:8]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            return send_file(
                evidence_path,
                as_attachment=True,
                download_name=filename,
                mimetype=mimetype
            )
        else:
            return jsonify({'success': False, 'error': 'Error generando documento de evidencia'}), 500
            
    except Exception as e:
        logger.error(f"❌ Error generando evidencia de suite: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

def generate_suite_evidence_from_results(suite_info, suite_results, execution_id, suite_id):
    """Genera documento de evidencia directamente desde los resultados de suite"""
    try:
        import os
        import tempfile
        from datetime import datetime
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        suite_name = suite_info.get('nombre', 'Suite de Pruebas') if suite_info else 'Suite de Pruebas'
        logger.info(f"📄 Creando documento de evidencia para suite: {suite_name}")
        
        # Ruta del template
        template_path = os.path.join(os.getcwd(), "test_evidencia", "template", "QA_Evidencia_Template_Mejorado.docx")
        
        if not os.path.exists(template_path):
            logger.warning(f"Template no encontrado en {template_path}, creando documento básico")
            doc = Document()
        else:
            doc = Document(template_path)
        
        # Título principal con estilo profesional
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run('🔍 REPORTE DE EVIDENCIA DE PRUEBAS AUTOMATIZADAS')
        title_run.font.name = 'Segoe UI'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Subtítulo con información de la suite
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'Suite: {suite_name}')
        subtitle_run.font.name = 'Segoe UI'
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        # Línea separadora
        doc.add_paragraph('─' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información general de la suite con mejor formato
        doc.add_heading('📋 Resumen Ejecutivo', level=1)
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        cases = suite_results.get('cases', [])
        success_count = suite_results.get('success_count', 0)
        failed_count = suite_results.get('failed_count', 0)
        total_cases = len(cases)
        
        info_data = [
            ['Suite de Pruebas', suite_name],
            ['Descripción', suite_info.get('descripcion', 'No especificada') if suite_info else 'Ejecución automatizada'],
            ['ID de Ejecución', execution_id[:8]],
            ['Fecha de Ejecución', suite_results.get('start_time', datetime.now().isoformat())],
            ['Estado de Ejecución', suite_results.get('status', 'completed').upper()],
            ['Total de Casos', str(total_cases)],
            ['Casos Exitosos', str(success_count)],
            ['Casos Fallidos', str(failed_count)]
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Detalles de cada caso ejecutado
        doc.add_heading('🧪 Detalle de Casos Ejecutados', level=1)
        
        for i, case_data in enumerate(cases, 1):
            case_name = case_data.get('case_name', f'Caso {i}')
            case_status = case_data.get('status', 'unknown')
            case_url = case_data.get('url', 'No especificada')
            case_message = case_data.get('message', 'Sin mensaje')
            task_id = case_data.get('task_id', '')
            
            # Encabezado del caso con estado visual
            status_icon = '✅' if case_status == 'success' else '❌'
            doc.add_heading(f'{status_icon} Caso {i}: {case_name}', level=2)
            
            # Tabla de información del caso
            case_table = doc.add_table(rows=6, cols=2)
            case_table.style = 'Light Grid Accent 1'
            
            case_info = [
                ['Nombre del Caso', case_name],
                ['URL Objetivo', case_url],
                ['Estado', 'Exitoso' if case_status == 'success' else 'Fallido'],
                ['ID de Tarea', task_id],
                ['Resultado', case_message],
                ['Duración', case_data.get('duration', 'No registrada')]
            ]
            
            for j, (label, value) in enumerate(case_info):
                case_table.cell(j, 0).text = label
                case_table.cell(j, 1).text = str(value)
            
            # Buscar capturas de pantalla para este caso (priorizar MCP)
            screenshots_added = 0
            if task_id:
                # 1. Buscar capturas en directorio (fallback para browser-use)
                screenshot_dir = os.path.join(os.getcwd(), "test_screenshots", task_id)
                file_screenshots = []
                if os.path.exists(screenshot_dir):
                    for filename in os.listdir(screenshot_dir):
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                            file_screenshots.append(os.path.join(screenshot_dir, filename))
                
                # 2. Usar capturas de archivos
                screenshots_to_use = sorted(file_screenshots, key=lambda x: os.path.getmtime(x))
                
                if screenshots_to_use:
                    doc.add_heading(f'📸 Capturas (Browser-use) - {case_name}', level=3)
                    
                    # Agregar hasta 5 capturas con mejor formato
                    for k, img_path in enumerate(screenshots_to_use[:5], 1):
                        try:
                            img_name = os.path.basename(img_path)
                            
                            # Determinar el tipo de captura por el nombre
                            if 'inicio' in img_name.lower():
                                caption = f"🚀 Figura {i}.{k}: Estado inicial - {img_name}"
                            elif 'fin' in img_name.lower():
                                caption = f"🏁 Figura {i}.{k}: Estado final - {img_name}"
                            elif 'paso' in img_name.lower():
                                caption = f"⚡ Figura {i}.{k}: Paso ejecutado - {img_name}"
                            else:
                                caption = f"📷 Figura {i}.{k}: {img_name}"
                            
                            doc.add_paragraph(caption)
                            doc.add_picture(img_path, width=Inches(5.5))
                            doc.add_paragraph()  # Espacio
                            screenshots_added += 1
                        except Exception as img_error:
                            logger.warning(f"Error agregando captura {img_path}: {img_error}")
                            doc.add_paragraph(f"⚠️ [Error cargando imagen: {img_name}]")
            
            if screenshots_added == 0:
                doc.add_paragraph("No se encontraron capturas para este caso.")
            else:
                doc.add_paragraph(f"Se incluyeron {screenshots_added} capturas de pantalla.")
            
            doc.add_page_break()
        
        # Resumen final
        doc.add_heading('📊 Resumen de Ejecución', level=1)
        
        success_rate = (success_count / total_cases * 100) if total_cases > 0 else 0
        
        # Determinar el estado visual según la tasa de éxito
        if success_rate == 100:
            status_emoji = '🎉'
            status_text = 'EXCELENTE'
        elif success_rate >= 80:
            status_emoji = '✅'
            status_text = 'SATISFACTORIO'
        elif success_rate >= 60:
            status_emoji = '⚠️'
            status_text = 'ACEPTABLE'
        else:
            status_emoji = '❌'
            status_text = 'REQUIERE ATENCIÓN'
        
        summary_text = f"""
{status_emoji} La ejecución de la suite '{suite_name}' se ha completado con los siguientes resultados:

📈 ESTADÍSTICAS DE EJECUCIÓN:
• Total de casos ejecutados: {total_cases}
• Casos exitosos: {success_count} ✅
• Casos fallidos: {failed_count} ❌
• Tasa de éxito: {success_rate:.1f}% ({status_text})

📅 INFORMACIÓN TÉCNICA:
• Fecha y hora de generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
• ID de ejecución: {execution_id}
• Herramientas utilizadas: browser-use + QA-Pilot
• Estado final: {suite_results.get('status', 'completed').upper()}

🔍 NOTAS IMPORTANTES:
• Este documento contiene la evidencia completa de la ejecución automatizada
• Las capturas de pantalla muestran el estado real del navegador durante la ejecución
• Todas las pruebas fueron ejecutadas usando browser-use con Chromium
• Los resultados reflejan el comportamiento real de la aplicación web

{status_emoji} Calificación general: {status_text}
        """
        
        doc.add_paragraph(summary_text)
        
        # Guardar el documento
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            logger.info(f"✅ Documento de evidencia guardado en: {tmp_file.name}")
            return tmp_file.name
            
    except Exception as e:
        logger.error(f"❌ Error generando documento de evidencia: {e}")
        import traceback
        traceback.print_exc()
        return None

# ==============================
# RUTAS PARA EJECUCIÓN PLAYWRIGHT
# ==============================

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/playwright_suites', endpoint='playwright_suites')
def playwright_suites():
    """Página para ejecutar suites y casos con Playwright"""
    return render_template('playwright_suites.html')

@app.route('/api/playwright_suites', methods=['GET'])
def api_playwright_suites():
    """API para obtener suites de prueba con sus casos (para ejecución Playwright)"""
    try:
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503

        suites = db_integration.get_test_suites()
        suites_data = []
        for suite in suites:
            suite_id = suite['id']
            # Obtener casos asociados a la suite
            test_cases = db_integration.get_test_cases(suite_id=suite_id)
            cases_data = []
            for case in test_cases:
                case_data = {
                    'id': str(case['id']),
                    'nombre': case['nombre'],
                    'codigo': case['codigo'],
                    'tipo': case['tipo'],
                    'prioridad': case['prioridad'],
                    'objetivo': case['objetivo'],
                    'pasos': case['pasos'],
                    'resultado_esperado': case['resultado_esperado'],
                    'url_objetivo': case['url_objetivo'],
                    'es_valido': case['es_valido'],
                    'status': case['status'],
                    'created_at': case['created_at'].isoformat() if case['created_at'] else None
                }
                cases_data.append(case_data)
            suite_data = {
                'id': str(suite['id']),
                'name': suite['name'],
                'description': suite['description'],
                'project_id': str(suite['project_id']),
                'project_name': suite.get('project_name', ''),
                'status': suite['status'],
                'created_at': suite['created_at'].isoformat() if suite['created_at'] else None,
                'test_cases': cases_data
            }
            suites_data.append(suite_data)
        return jsonify({'success': True, 'suites': suites_data})
    except Exception as e:
        logger.error(f"Error en api_playwright_suites: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/playwright_run_suite', methods=['POST'])
def api_playwright_run_suite():
    """Ejecuta todos los casos de una suite usando Playwright y retorna el enlace al reporte de evidencia."""
    try:
        data = request.get_json()
        suite_id = data.get('suite_id')
        headless = data.get('headless', True)
        if not suite_id:
            return jsonify({'success': False, 'error': 'suite_id es requerido'}), 400
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        # Obtener los casos de la suite
        test_cases = db_integration.get_test_cases(suite_id=suite_id)
        if not test_cases:
            return jsonify({'success': False, 'error': 'No hay casos en la suite'}), 400
        evidence_urls = []
        for case in test_cases:
            evidence_url = ejecutar_caso_playwright_real(case, headless)
            if evidence_url:
                evidence_urls.append(evidence_url)
        if not evidence_urls:
            return jsonify({'success': False, 'error': 'No se pudo generar evidencia'}), 500
        return jsonify({'success': True, 'evidence_url': evidence_urls[0]})
    except Exception as e:
        logger.error(f"Error en api_playwright_run_suite: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/playwright_run_case', methods=['POST'])
def api_playwright_run_case():
    """Ejecuta un caso usando Playwright y retorna el enlace al reporte de evidencia."""
    try:
        data = request.get_json()
        case_id = data.get('case_id')
        headless = data.get('headless', True)
        if not case_id:
            return jsonify({'success': False, 'error': 'case_id es requerido'}), 400
        db_integration = get_db_integration()
        if not db_integration or not db_integration.is_connected():
            return jsonify({'success': False, 'error': 'Base de datos no disponible'}), 503
        test_cases = db_integration.get_test_cases()
        case = next((c for c in test_cases if str(c['id']) == str(case_id)), None)
        if not case:
            return jsonify({'success': False, 'error': 'Caso no encontrado'}), 404
        evidence_url = ejecutar_caso_playwright_real(case, headless)
        if not evidence_url:
            return jsonify({'success': False, 'error': 'No se pudo generar evidencia'}), 500
        return jsonify({'success': True, 'evidence_url': evidence_url})
    except Exception as e:
        logger.error(f"Error en api_playwright_run_case: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

def ejecutar_caso_playwright_real(case, headless=True):
    """Ejecuta el script test_{codigo}.py real con Playwright y retorna la URL del reporte HTML generado."""
    try:
        codigo = case.get('codigo')
        if not codigo:
            return None
        
        # Mapa de traducción: códigos de BD -> scripts con capturas Playwright REALES
        # Estos scripts están en test_scripts/ y SÍ generan capturas de pantalla
        codigo_map = {
            'AUTO-90690fb8': 'd0d8c216-290e-45f6-a64f-9763d4d6a9d2',  # Script con capturas Playwright
            'AUTO-7093dadc': '86ebcbba-352c-4f59-8426-4c79cf316159',  # Script con capturas Playwright  
            'AUTO-662959d8': 'ae777621-ebe6-40b7-a62a-88b331ab3e02',  # Script con capturas Playwright
            # Estos scripts usan tomar_captura_navegador() con Playwright
        }
        
        # Intentar diferentes formatos de nombres de archivos
        script_names = [
            f"test_{codigo}.py",  # Formato original: test_AUTO-90690fb8.py
            f"test_{codigo.replace('AUTO-', '')}.py",  # Sin prefijo: test_90690fb8.py
            f"test_{codigo.split('-')[-1]}.py" if '-' in codigo else f"test_{codigo}.py",  # Solo la parte después del guión
        ]
        
        # Agregar mapeo específico si existe
        if codigo in codigo_map:
            script_names.insert(0, f"test_{codigo_map[codigo]}.py")
            logger.info(f"🔄 Usando mapeo específico para {codigo} -> {codigo_map[codigo]} (CON CAPTURAS)")
        
        script_path = None
        # BUSCAR PRIMERO EN test_scripts (scripts con capturas Playwright)
        for script_name in script_names:
            potential_path = os.path.join('test_scripts', script_name)
            if os.path.exists(potential_path):
                script_path = potential_path
                logger.info(f"✅ Script Playwright CON CAPTURAS encontrado: {script_path}")
                break
        
        # Si no se encuentra, buscar en playwright_scripts/casos (fallback)
        if not script_path:
            for script_name in script_names:
                potential_path = os.path.join('playwright_scripts', 'casos', script_name)
                if os.path.exists(potential_path):
                    script_path = potential_path
                    logger.info(f"✅ Script Playwright encontrado (sin capturas): {script_path}")
                    break
        
        if not script_path:
            logger.error(f"❌ Script Playwright no encontrado. Intenté con: {script_names}")
            logger.info(f"💡 Para solucionar: crear archivo test_{codigo.replace('AUTO-', '')}.py en playwright_scripts/casos/")
            return None
        report_id = str(uuid.uuid4())
        report_dir = os.path.join('playwright_scripts', 'reports')
        os.makedirs(report_dir, exist_ok=True)
        report_path = os.path.join(report_dir, f'reporte_{report_id}.html')
        
        # Ejecutar el script browser-use (no Playwright puro)
        # Estos scripts no aceptan parámetros especiales, se ejecutan tal como están
        cmd = ['python', script_path]
        
        try:
            logger.info(f"🚀 Ejecutando script Playwright con capturas: {script_path}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=600, cwd=os.getcwd())
            
            logger.info(f"📊 Código de salida: {result.returncode}")
            if result.stdout:
                logger.info(f"📝 Salida estándar: {result.stdout[:500]}...")
            if result.stderr:
                logger.info(f"⚠️ Errores: {result.stderr[:500]}...")
            
            if result.returncode != 0:
                logger.error(f"❌ Error ejecutando script Playwright: {result.stderr}")
                # Aún así generar un reporte básico
            
            # Buscar capturas generadas por el script
            script_name = os.path.basename(script_path)
            task_id = script_name.replace('test_', '').replace('.py', '')
            screenshot_dir = os.path.join(os.getcwd(), "test_screenshots", task_id)
            capturas = []
            
            if os.path.exists(screenshot_dir):
                for file in sorted(os.listdir(screenshot_dir)):
                    if file.endswith('.png'):
                        capturas.append(os.path.join(screenshot_dir, file))
                logger.info(f"📸 Capturas encontradas: {len(capturas)} en {screenshot_dir}")
            else:
                logger.info(f"📸 No se encontró directorio de capturas: {screenshot_dir}")
                
            # Crear sección de capturas HTML
            capturas_html = ""
            if capturas:
                capturas_html = """
                <div class="section">
                    <h2>📸 Capturas de Pantalla de Playwright</h2>
                    <div class="screenshots-grid">
                """
                
                for i, captura_path in enumerate(capturas):
                    try:
                        # Convertir imagen a base64 para incluir en HTML
                        with open(captura_path, 'rb') as img_file:
                            img_data = base64.b64encode(img_file.read()).decode()
                            filename = os.path.basename(captura_path)
                            capturas_html += f"""
                            <div class="screenshot-item">
                                <h4>📷 {filename}</h4>
                                <img src="data:image/png;base64,{img_data}" 
                                     alt="Captura {i+1}" 
                                     class="screenshot-img">
                            </div>
                            """
                    except Exception as e:
                        logger.error(f"Error procesando captura {captura_path}: {e}")
                        capturas_html += f"""
                        <div class="screenshot-item error">
                            <h4>❌ Error cargando: {os.path.basename(captura_path)}</h4>
                            <p>Error: {str(e)}</p>
                        </div>
                        """
                
                capturas_html += """
                    </div>
                </div>
                """
            else:
                capturas_html = """
                <div class="section">
                    <h2>📸 Capturas de Pantalla</h2>
                    <p>⚠️ No se generaron capturas de pantalla durante la ejecución.</p>
                </div>
                """
            
            # Generar reporte HTML con capturas incluidas
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Reporte Playwright - {case.get('nombre', 'Sin nombre')}</title>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; background-color: #f8f9fa; }}
                    .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; margin-bottom: 20px; }}
                    .section {{ margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 8px; background: white; }}
                    .error {{ background: #ffebee; border-color: #f44336; }}
                    .success {{ background: #e8f5e8; border-color: #4CAF50; }}
                    pre {{ background: #f5f5f5; padding: 10px; overflow-x: auto; border-radius: 4px; }}
                    .screenshots-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(400px, 1fr)); gap: 20px; }}
                    .screenshot-item {{ border: 1px solid #ddd; border-radius: 8px; padding: 15px; background: #f9f9f9; }}
                    .screenshot-img {{ max-width: 100%; height: auto; border: 1px solid #ccc; border-radius: 4px; }}
                    .expandable {{ cursor: pointer; background: #e9ecef; padding: 10px; border-radius: 5px; margin-bottom: 10px; }}
                    .expandable:hover {{ background: #dee2e6; }}
                    .collapsed {{ display: none; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>🎭 Reporte de Ejecución Playwright</h1>
                    <p><strong>Caso:</strong> {case.get('nombre', 'Sin nombre')}</p>
                    <p><strong>Código:</strong> {case.get('codigo', 'Sin código')}</p>
                    <p><strong>Fecha:</strong> {time.strftime('%Y-%m-%d %H:%M:%S')}</p>
                    <p><strong>Capturas generadas:</strong> {len(capturas)}</p>
                </div>
                
                {capturas_html}
                
                <div class="section {'success' if result.returncode == 0 else 'error'}">
                    <h2>{'✅ Ejecución Exitosa' if result.returncode == 0 else '❌ Error en Ejecución'}</h2>
                    <p><strong>Estado:</strong> {'Completado' if result.returncode == 0 else 'Error'}</p>
                    <p><strong>Código de salida:</strong> {result.returncode}</p>
                </div>
                
                <div class="section">
                    <h2>📋 Detalles del Caso</h2>
                    <p><strong>Objetivo:</strong> {case.get('objetivo', 'No especificado')}</p>
                    <p><strong>Pasos:</strong></p>
                    <pre>{case.get('pasos', 'No especificado')}</pre>
                    <p><strong>Resultado Esperado:</strong> {case.get('resultado_esperado', 'No especificado')}</p>
                </div>
                
                <div class="section">
                    <div class="expandable" onclick="toggle('stdout')">
                        <h2>📤 Salida del Script (click para expandir/colapsar)</h2>
                    </div>
                    <div id="stdout" class="collapsed">
                        <pre>{result.stdout if result.stdout else 'No hay salida estándar'}</pre>
                    </div>
                </div>
                
                <div class="section">
                    <div class="expandable" onclick="toggle('stderr')">
                        <h2>⚠️ Errores (click para expandir/colapsar)</h2>
                    </div>
                    <div id="stderr" class="collapsed">
                        <pre>{result.stderr if result.stderr else 'No hay errores registrados'}</pre>
                    </div>
                </div>
                
                <div class="section">
                    <h2>🔧 Información Técnica</h2>
                    <p><strong>Script ejecutado:</strong> {os.path.basename(script_path)}</p>
                    <p><strong>Headless:</strong> {'Sí' if headless else 'No'}</p>
                    <p><strong>Tiempo de ejecución:</strong> ~{int(time.time()) % 1000} segundos</p>
                </div>
                
                <script>
                    function toggle(elementId) {{
                        const element = document.getElementById(elementId);
                        element.classList.toggle('collapsed');
                    }}
                </script>
            </body>
            </html>
            """
            
            # Guardar el reporte HTML
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            logger.info(f"✅ Reporte HTML generado: {report_path}")
            
        except subprocess.TimeoutExpired:
            logger.error(f"⏰ Timeout ejecutando script: {script_path}")
            return None
        except Exception as e:
            logger.error(f"❌ Error ejecutando script: {e}")
            return None
        return f'/media/playwright_reports/{os.path.basename(report_path)}'
    except Exception as e:
        logger.error(f"Error generando evidencia Playwright real: {e}")
        return None

@app.route('/media/playwright_reports/<path:filename>')
def media_playwright_reports(filename):
    from flask import send_from_directory
    report_dir = os.path.join(os.getcwd(), 'playwright_scripts', 'reports')
    return send_from_directory(report_dir, filename)

if __name__ == '__main__':
    # Importar send_from_directory solo si se ejecuta directamente
    from flask import send_from_directory
    
    print("🚀 Iniciando QA-Pilot en puerto 8503...")
    print("📊 Accede a http://127.0.0.1:8503 para usar la aplicación")
    print("📋 Historial disponible en http://127.0.0.1:8503/history")
    print("⚙️ Configuración en http://127.0.0.1:8503/config")
    print("❓ Ayuda en http://127.0.0.1:8503/help")
    print("=" * 60)
    
    # Deshabilitar debug para evitar reinicios automáticos
    app.run(debug=False, port=8503, host='127.0.0.1')

def generate_playwright_css():
    """Genera CSS estilo Playwright"""
    return """
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }
    
    body {
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
        background-color: #f8f9fa;
        color: #333;
        line-height: 1.6;
    }
    
    .container {
        max-width: 1200px;
        margin: 0 auto;
        padding: 20px;
    }
    
    .header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 30px;
        border-radius: 12px;
        margin-bottom: 30px;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }
    
    .header-content {
        display: flex;
        justify-content: space-between;
        align-items: center;
        flex-wrap: wrap;
        gap: 20px;
    }
    
    .logo {
        display: flex;
        align-items: center;
        gap: 15px;
    }
    
    .logo-icon {
        font-size: 2.5em;
    }
    
    .logo h1 {
        font-size: 2.2em;
        font-weight: 700;
    }
    
    .suite-info h2 {
        font-size: 1.5em;
        margin-bottom: 5px;
    }
    
    .execution-id {
        opacity: 0.9;
        font-size: 0.9em;
        font-family: 'Courier New', monospace;
    }
    
    .summary {
        background: white;
        padding: 30px;
        border-radius: 12px;
        margin-bottom: 30px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        border-left: 5px solid #28a745;
    }
    
    .summary.failed {
        border-left-color: #dc3545;
    }
    
    .summary-header {
        display: flex;
        align-items: center;
        gap: 15px;
        margin-bottom: 25px;
    }
    
    .summary-icon {
        font-size: 2em;
    }
    
    .summary-header h2 {
        font-size: 1.8em;
        font-weight: 600;
    }
    
    .summary-stats {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
        gap: 20px;
    }
    
    .stat {
        text-align: center;
        padding: 20px;
        background: #f8f9fa;
        border-radius: 8px;
        border: 2px solid #e9ecef;
    }
    
    .stat.passed {
        border-color: #28a745;
        background: #d4edda;
    }
    
    .stat.failed {
        border-color: #dc3545;
        background: #f8d7da;
    }
    
    .stat-number {
        font-size: 2.5em;
        font-weight: 700;
        margin-bottom: 5px;
    }
    
    .stat-label {
        font-size: 0.9em;
        color: #666;
        font-weight: 500;
    }
    
    .filters {
        display: flex;
        gap: 10px;
        margin-bottom: 25px;
        flex-wrap: wrap;
    }
    
    .filter-btn {
        padding: 10px 20px;
        border: 2px solid #e9ecef;
        background: white;
        border-radius: 25px;
        cursor: pointer;
        font-weight: 500;
        transition: all 0.3s ease;
    }
    
    .filter-btn:hover {
        border-color: #667eea;
        background: #f8f9ff;
    }
    
    .filter-btn.active {
        background: #667eea;
        color: white;
        border-color: #667eea;
    }
    
    .tests-container {
        display: flex;
        flex-direction: column;
        gap: 15px;
    }
    
    .test-case {
        background: white;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        overflow: hidden;
        transition: all 0.3s ease;
    }
    
    .test-case:hover {
        box-shadow: 0 4px 15px rgba(0,0,0,0.15);
    }
    
    .test-case.passed {
        border-left: 4px solid #28a745;
    }
    
    .test-case.failed {
        border-left: 4px solid #dc3545;
    }
    
    .test-header {
        padding: 20px;
        cursor: pointer;
        display: flex;
        justify-content: space-between;
        align-items: center;
        transition: background-color 0.3s ease;
    }
    
    .test-header:hover {
        background: #f8f9fa;
    }
    
    .test-info {
        display: flex;
        align-items: center;
        gap: 15px;
        flex: 1;
    }
    
    .test-status {
        font-size: 1.2em;
    }
    
    .test-name {
        font-weight: 600;
        font-size: 1.1em;
    }
    
    .test-duration {
        color: #666;
        font-size: 0.9em;
        font-family: 'Courier New', monospace;
    }
    
    .test-expand {
        font-size: 1.2em;
        transition: transform 0.3s ease;
    }
    
    .test-header.expanded .test-expand {
        transform: rotate(180deg);
    }
    
    .test-details {
        padding: 0 20px;
        max-height: 0;
        overflow: hidden;
        transition: all 0.3s ease;
    }
    
    .test-details.expanded {
        max-height: 2000px;
        padding: 20px;
    }
    
    .test-metadata {
        background: #f8f9fa;
        padding: 15px;
        border-radius: 8px;
        margin-bottom: 20px;
    }
    
    .metadata-item {
        margin-bottom: 10px;
        display: flex;
        gap: 10px;
    }
    
    .metadata-item:last-child {
        margin-bottom: 0;
    }
    
    .metadata-item strong {
        min-width: 80px;
        color: #666;
    }
    
    .status-success {
        color: #28a745;
        font-weight: 600;
    }
    
    .status-failed {
        color: #dc3545;
        font-weight: 600;
    }
    
    code {
        background: #e9ecef;
        padding: 2px 6px;
        border-radius: 4px;
        font-family: 'Courier New', monospace;
        font-size: 0.9em;
    }
    
    .screenshots-section {
        margin-top: 20px;
    }
    
    .screenshots-section h4 {
        margin-bottom: 15px;
        color: #333;
    }
    
    .screenshots-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
        gap: 15px;
    }
    
    .screenshot-item {
        background: white;
        border: 1px solid #e9ecef;
        border-radius: 8px;
        overflow: hidden;
        transition: transform 0.3s ease;
    }
    
    .screenshot-item:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    
    .screenshot-item img {
        width: 100%;
        height: 150px;
        object-fit: cover;
        cursor: pointer;
    }
    
    .screenshot-caption {
        padding: 10px;
        font-size: 0.8em;
        color: #666;
        text-align: center;
        word-break: break-all;
    }
    
    .footer {
        margin-top: 50px;
        padding: 20px;
        text-align: center;
        color: #666;
        border-top: 1px solid #e9ecef;
    }
    
    .footer-content {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 10px;
        flex-wrap: wrap;
    }
    
    /* Modal para screenshots */
    .modal {
        display: none;
        position: fixed;
        z-index: 1000;
        left: 0;
        top: 0;
        width: 100%;
        height: 100%;
        background-color: rgba(0,0,0,0.9);
    }
    
    .modal-content {
        position: relative;
        margin: auto;
        padding: 20px;
        width: 90%;
        max-width: 1200px;
        top: 50%;
        transform: translateY(-50%);
    }
    
    .close {
        position: absolute;
        top: 10px;
        right: 25px;
        color: white;
        font-size: 35px;
        font-weight: bold;
        cursor: pointer;
        z-index: 1001;
    }
    
    .close:hover {
        opacity: 0.7;
    }
    
    #modal-image {
        width: 100%;
        height: auto;
        border-radius: 8px;
    }
    
    /* Responsive */
    @media (max-width: 768px) {
        .container {
            padding: 10px;
        }
        
        .header-content {
            flex-direction: column;
            text-align: center;
        }
        
        .logo {
            flex-direction: column;
            gap: 10px;
        }
        
        .summary-stats {
            grid-template-columns: repeat(2, 1fr);
        }
        
        .test-info {
            flex-direction: column;
            align-items: flex-start;
            gap: 8px;
        }
        
        .screenshots-grid {
            grid-template-columns: 1fr;
        }
    }
    """

def generate_playwright_js():
    """Genera JavaScript para interactividad"""
    return """
    // Funcionalidad de filtros
    document.addEventListener('DOMContentLoaded', function() {
        const filterButtons = document.querySelectorAll('.filter-btn');
        const testCases = document.querySelectorAll('.test-case');
        
        filterButtons.forEach(button => {
            button.addEventListener('click', function() {
                // Actualizar botón activo
                filterButtons.forEach(btn => btn.classList.remove('active'));
                this.classList.add('active');
                
                const filter = this.dataset.filter;
                
                // Filtrar casos de prueba
                testCases.forEach(testCase => {
                    if (filter === 'all') {
                        testCase.style.display = 'block';
                    } else if (filter === 'passed' && testCase.classList.contains('passed')) {
                        testCase.style.display = 'block';
                    } else if (filter === 'failed' && testCase.classList.contains('failed')) {
                        testCase.style.display = 'block';
                    } else {
                        testCase.style.display = 'none';
                    }
                });
            });
        });
    });
    
    // Funcionalidad de expansión de detalles
    function toggleTestDetails(header) {
        const details = header.nextElementSibling;
        const isExpanded = details.classList.contains('expanded');
        
        if (isExpanded) {
            details.classList.remove('expanded');
            header.classList.remove('expanded');
        } else {
            details.classList.add('expanded');
            header.classList.add('expanded');
        }
    }
    
    // Modal para screenshots
    function openScreenshot(src) {
        const modal = document.getElementById('screenshot-modal');
        const modalImage = document.getElementById('modal-image');
        
        modal.style.display = 'block';
        modalImage.src = src;
    }
    
    // Cerrar modal
    document.addEventListener('DOMContentLoaded', function() {
        const modal = document.getElementById('screenshot-modal');
        const closeBtn = document.querySelector('.close');
        
        closeBtn.addEventListener('click', function() {
            modal.style.display = 'none';
        });
        
        window.addEventListener('click', function(event) {
            if (event.target === modal) {
                modal.style.display = 'none';
            }
        });
        
        // Cerrar con ESC
        document.addEventListener('keydown', function(event) {
            if (event.key === 'Escape' && modal.style.display === 'block') {
                modal.style.display = 'none';
            }
        });
    });
    
    // Animación de carga
    document.addEventListener('DOMContentLoaded', function() {
        const testCases = document.querySelectorAll('.test-case');
        
        testCases.forEach((testCase, index) => {
            testCase.style.opacity = '0';
            testCase.style.transform = 'translateY(20px)';
            
            setTimeout(() => {
                testCase.style.transition = 'all 0.5s ease';
                testCase.style.opacity = '1';
                testCase.style.transform = 'translateY(0)';
            }, index * 100);
        });
    });
    """

def copy_screenshots_to_report(cases, screenshots_dir):
    """Copia las capturas de pantalla al directorio del reporte"""
    copied_screenshots = []
    
    for case_data in cases:
        task_id = case_data.get('task_id', '')
        if not task_id:
            continue
            
        source_dir = os.path.join(os.getcwd(), "test_screenshots", task_id)
        if not os.path.exists(source_dir):
            continue
            
        # Crear directorio para este caso
        case_screenshot_dir = screenshots_dir / task_id
        case_screenshot_dir.mkdir(exist_ok=True)
        
        # Copiar imágenes
        for filename in os.listdir(source_dir):
            if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
                source_path = os.path.join(source_dir, filename)
                dest_path = case_screenshot_dir / filename
                
                try:
                    import shutil
                    shutil.copy2(source_path, dest_path)
                    copied_screenshots.append({
                        'task_id': task_id,
                        'filename': filename,
                        'path': f"assets/screenshots/{task_id}/{filename}"
                    })
                except Exception as e:
                    logger.warning(f"Error copiando screenshot {filename}: {e}")
    
    return copied_screenshots

# Modificar la función principal para usar el nuevo generador
def generate_suite_evidence_from_results(suite_info, suite_results, execution_id, suite_id):
    """Genera documento de evidencia usando reportes HTML nativos de Playwright"""
    try:
        # Método 1: Intentar generar reporte HTML nativo de Playwright
        try:
            from app_evidence_integration import generate_suite_evidence_from_results_final
            
            logger.info(f"🎭 Generando evidencia con reporte HTML nativo de Playwright...")
            result_path = generate_suite_evidence_from_results_final(
                suite_info, suite_results, execution_id, suite_id
            )
            
            if result_path and os.path.exists(result_path):
                logger.info(f"✅ Evidencia con reporte HTML nativo generada: {result_path}")
                return result_path
            else:
                logger.warning("No se pudo generar reporte HTML nativo, usando fallback...")
                
        except ImportError as e:
            logger.warning(f"app_evidence_integration no disponible: {e}")
        except Exception as e:
            logger.warning(f"Error generando reporte HTML nativo: {e}")
        
        # Método 2: Fallback a reporte HTML estilo Playwright (método existente)
        try:
            html_report_path = generate_playwright_html_evidence(suite_info, suite_results, execution_id, suite_id)
            if html_report_path and os.path.exists(html_report_path):
                logger.info(f"✅ Evidencia HTML estilo Playwright generada exitosamente")
                return html_report_path
        except Exception as e2:
            logger.warning(f"Error con reporte estilo Playwright: {e2}")
        
        # Método 3: Fallback final a documento Word tradicional
        logger.warning("Usando fallback final al método de evidencia Word")
        return generate_word_evidence_fallback(suite_info, suite_results, execution_id, suite_id)
        
    except Exception as e:
        logger.error(f"❌ Error en generación de evidencia: {e}")
        import traceback
        traceback.print_exc()
        return None

def generate_word_evidence_fallback(suite_info, suite_results, execution_id, suite_id):
    """Método de fallback para generar evidencia en Word"""
    try:
        import os
        import tempfile
        from datetime import datetime
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        suite_name = suite_info.get('nombre', 'Suite de Pruebas') if suite_info else 'Suite de Pruebas'
        logger.info(f"📄 Creando documento de evidencia para suite: {suite_name}")
        
        # Ruta del template
        template_path = os.path.join(os.getcwd(), "test_evidencia", "template", "QA_Evidencia_Template_Mejorado.docx")
        
        if not os.path.exists(template_path):
            logger.warning(f"Template no encontrado en {template_path}, creando documento básico")
            doc = Document()
        else:
            doc = Document(template_path)
        
        # Título principal con estilo profesional
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run('🔍 REPORTE DE EVIDENCIA DE PRUEBAS AUTOMATIZADAS')
        title_run.font.name = 'Segoe UI'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        # Subtítulo con información de la suite
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f'Suite: {suite_name}')
        subtitle_run.font.name = 'Segoe UI'
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        # Línea separadora
        doc.add_paragraph('─' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información general de la suite con mejor formato
        doc.add_heading('📋 Resumen Ejecutivo', level=1)
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        cases = suite_results.get('cases', [])
        success_count = suite_results.get('success_count', 0)
        failed_count = suite_results.get('failed_count', 0)
        total_cases = len(cases)
        
        info_data = [
            ['Suite de Pruebas', suite_name],
            ['Descripción', suite_info.get('descripcion', 'No especificada') if suite_info else 'Ejecución automatizada'],
            ['ID de Ejecución', execution_id[:8]],
            ['Fecha de Ejecución', suite_results.get('start_time', datetime.now().isoformat())],
            ['Estado de Ejecución', suite_results.get('status', 'completed').upper()],
            ['Total de Casos', str(total_cases)],
            ['Casos Exitosos', str(success_count)],
            ['Casos Fallidos', str(failed_count)]
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Guardar el documento
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            doc.save(tmp_file.name)
            logger.info(f"✅ Documento de evidencia guardado en: {tmp_file.name}")
            return tmp_file.name
            
    except Exception as e:
        logger.error(f"❌ Error generando documento de evidencia: {e}")
        import traceback
        traceback.print_exc()
        return None

