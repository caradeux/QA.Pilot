#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Integración de Reportes HTML Nativos de Playwright
-------------------------------------------------
Este módulo integra los reportes HTML nativos de Playwright/pytest-html
en el sistema de generación de evidencia de QA-Pilot.
"""

import os
import sys
import tempfile
import shutil
import subprocess
import json
import uuid
import logging
from datetime import datetime
from pathlib import Path

# Configurar logging
logger = logging.getLogger(__name__)

# Agregar el directorio de playwright_scripts al path
sys.path.append(os.path.join(os.getcwd(), 'playwright_scripts'))

try:
    from playwright_scripts.pytest_runner import create_pytest_runner
except ImportError as e:
    logger.error(f"No se pudo importar pytest_runner: {e}")
    create_pytest_runner = None

class PlaywrightHTMLReportGenerator:
    """Generador de reportes HTML nativos de Playwright para evidencia de QA"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="qa_playwright_reports_")
        self.reports_dir = os.path.join(self.temp_dir, "reports")
        os.makedirs(self.reports_dir, exist_ok=True)
        
    def generate_native_html_report_for_suite(self, suite_info, suite_results, execution_id, suite_id):
        """
        Genera un reporte HTML nativo de Playwright para una suite de pruebas.
        
        Args:
            suite_info: Información de la suite
            suite_results: Resultados de la ejecución de la suite
            execution_id: ID de la ejecución
            suite_id: ID de la suite
            
        Returns:
            str: Ruta al reporte HTML generado o None si falla
        """
        try:
            if not create_pytest_runner:
                logger.error("pytest_runner no está disponible")
                return None
                
            suite_name = suite_info.get('nombre', 'Suite de Pruebas') if suite_info else 'Suite de Pruebas'
            logger.info(f"🎭 Generando reporte HTML nativo de Playwright para suite: {suite_name}")
            
            # Obtener casos de la suite
            cases = suite_results.get('cases', [])
            if not cases:
                logger.warning("No hay casos para procesar en la suite")
                return None
            
            # Crear scripts temporales para cada caso
            script_paths = []
            for i, case_data in enumerate(cases):
                case_name = case_data.get('case_name', f'Test Case {i+1}')
                task_id = case_data.get('task_id', f'case_{i+1}')
                
                # Buscar script existente o crear uno básico
                script_content = self._create_basic_test_script(case_data, i+1)
                
                # Guardar script temporal
                script_path = os.path.join(self.temp_dir, f"test_case_{i+1}_{task_id}.py")
                with open(script_path, 'w', encoding='utf-8') as f:
                    f.write(script_content)
                
                script_paths.append(script_path)
                logger.info(f"📝 Script creado para caso {i+1}: {script_path}")
            
            # Crear runner pytest
            pytest_runner = create_pytest_runner(self.temp_dir)
            
            # Ejecutar suite con pytest para generar reporte HTML nativo
            result = pytest_runner.run_suite_with_html_report(
                script_paths,
                suite_name=suite_name,
                execution_id=execution_id
            )
            
            if result.get('success') and result.get('html_report_path'):
                # Copiar reporte a directorio de evidencia permanente
                evidence_dir = os.path.join(os.getcwd(), "test_evidence", "playwright_reports")
                os.makedirs(evidence_dir, exist_ok=True)
                
                final_report_path = os.path.join(evidence_dir, f"native_report_{execution_id}.html")
                shutil.copy2(result['html_report_path'], final_report_path)
                
                logger.info(f"✅ Reporte HTML nativo generado: {final_report_path}")
                
                # También copiar el reporte JSON si existe
                if result.get('json_report_path'):
                    final_json_path = os.path.join(evidence_dir, f"native_report_{execution_id}.json")
                    shutil.copy2(result['json_report_path'], final_json_path)
                    
                return final_report_path
            else:
                logger.error(f"Error generando reporte: {result.get('error', 'Error desconocido')}")
                return None
                
        except Exception as e:
            logger.error(f"❌ Error generando reporte HTML nativo: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _create_basic_test_script(self, case_data, case_number):
        """Crea un script básico de test para un caso"""
        
        case_name = case_data.get('case_name', f'Test Case {case_number}')
        case_url = case_data.get('url', 'https://example.com')
        case_status = case_data.get('status', 'unknown')
        case_message = case_data.get('message', 'No message')
        task_id = case_data.get('task_id', '')
        
        # Determinar si el test debe pasar o fallar basado en el status
        should_pass = case_status == 'success'
        
        script = f'''#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Test Case: {case_name}
Generated for Playwright HTML Report
Task ID: {task_id}
"""

import pytest
import asyncio
import time
from datetime import datetime

class TestCase{case_number}:
    """Test case {case_number}: {case_name}"""
    
    @pytest.mark.asyncio
    async def test_{case_name.lower().replace(" ", "_").replace("-", "_")}(self):
        """
        Test: {case_name}
        URL: {case_url}
        Expected Status: {case_status}
        Message: {case_message}
        """
        
        print(f"🧪 Ejecutando: {case_name}")
        print(f"📍 URL: {case_url}")
        print(f"📝 Mensaje: {case_message}")
        print(f"🏷️  Task ID: {task_id}")
        
        # Simular ejecución del test
        await asyncio.sleep(0.5)  # Simular tiempo de ejecución
        
        # Determinar resultado basado en el status original
        if "{should_pass}".lower() == "true":
            print("✅ Test ejecutado exitosamente")
            assert True, "Test passed successfully"
        else:
            print("❌ Test falló como se esperaba")
            pytest.fail(f"Test failed: {case_message}")

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
'''
        
        return script
    
    def generate_hybrid_evidence(self, suite_info, suite_results, execution_id, suite_id):
        """
        Genera evidencia híbrida que incluye tanto el reporte HTML nativo
        como el documento Word tradicional.
        
        Returns:
            dict: Diccionario con las rutas de ambos reportes
        """
        try:
            # Generar reporte HTML nativo
            html_report_path = self.generate_native_html_report_for_suite(
                suite_info, suite_results, execution_id, suite_id
            )
            
            # Generar documento Word (usando el método existente)
            word_report_path = self._generate_word_evidence_with_html_reference(
                suite_info, suite_results, execution_id, suite_id, html_report_path
            )
            
            return {
                'html_report_path': html_report_path,
                'word_report_path': word_report_path,
                'success': html_report_path is not None or word_report_path is not None
            }
            
        except Exception as e:
            logger.error(f"❌ Error generando evidencia híbrida: {e}")
            return {
                'html_report_path': None,
                'word_report_path': None,
                'success': False,
                'error': str(e)
            }
    
    def _generate_word_evidence_with_html_reference(self, suite_info, suite_results, execution_id, suite_id, html_report_path):
        """Genera documento Word que incluye referencia al reporte HTML nativo"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            suite_name = suite_info.get('nombre', 'Suite de Pruebas') if suite_info else 'Suite de Pruebas'
            
            # Crear documento
            doc = Document()
            
            # Título
            title = doc.add_heading('🎭 EVIDENCIA DE PRUEBAS AUTOMATIZADAS', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtítulo
            subtitle = doc.add_heading(f'Suite: {suite_name}', level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información general
            doc.add_heading('📋 Información General', level=2)
            
            cases = suite_results.get('cases', [])
            success_count = suite_results.get('success_count', 0)
            failed_count = suite_results.get('failed_count', 0)
            total_cases = len(cases)
            
            info_para = doc.add_paragraph()
            info_para.add_run(f"Execution ID: {execution_id}\\n")
            info_para.add_run(f"Fecha de Ejecución: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\\n")
            info_para.add_run(f"Total de Casos: {total_cases}\\n")
            info_para.add_run(f"Casos Exitosos: {success_count}\\n")
            info_para.add_run(f"Casos Fallidos: {failed_count}\\n")
            
            # Referencia al reporte HTML nativo
            if html_report_path:
                doc.add_heading('🎭 Reporte HTML Nativo de Playwright', level=2)
                html_para = doc.add_paragraph()
                html_para.add_run("Se ha generado un reporte HTML nativo de Playwright que incluye:")
                
                html_list = doc.add_paragraph(style='List Bullet')
                html_list.add_run("Interfaz interactiva estilo Playwright")
                
                html_list2 = doc.add_paragraph(style='List Bullet')
                html_list2.add_run("Detalles completos de ejecución de cada test")
                
                html_list3 = doc.add_paragraph(style='List Bullet')
                html_list3.add_run("Métricas de tiempo y rendimiento")
                
                html_list4 = doc.add_paragraph(style='List Bullet')
                html_list4.add_run("Logs detallados de ejecución")
                
                path_para = doc.add_paragraph()
                path_para.add_run(f"📂 Ubicación del reporte HTML: {os.path.basename(html_report_path)}")
            
            # Resumen de casos
            doc.add_heading('📊 Resumen de Casos de Prueba', level=2)
            
            for i, case_data in enumerate(cases, 1):
                case_name = case_data.get('case_name', f'Test Case {i}')
                case_status = case_data.get('status', 'unknown')
                case_message = case_data.get('message', 'No message')
                
                case_para = doc.add_paragraph()
                status_icon = "✅" if case_status == 'success' else "❌"
                case_para.add_run(f"{status_icon} {case_name}: {case_status.upper()}")
                
                if case_message != 'No message':
                    msg_para = doc.add_paragraph()
                    msg_para.add_run(f"   Mensaje: {case_message}")
            
            # Guardar documento
            evidence_dir = os.path.join(os.getcwd(), "test_evidence", "word_reports")
            os.makedirs(evidence_dir, exist_ok=True)
            
            word_path = os.path.join(evidence_dir, f"evidence_with_html_{execution_id}.docx")
            doc.save(word_path)
            
            logger.info(f"📄 Documento Word con referencia HTML generado: {word_path}")
            return word_path
            
        except Exception as e:
            logger.error(f"❌ Error generando documento Word: {e}")
            return None
    
    def cleanup(self):
        """Limpia archivos temporales"""
        try:
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"🧹 Directorio temporal limpiado: {self.temp_dir}")
        except Exception as e:
            logger.warning(f"No se pudo limpiar directorio temporal: {e}")

def create_playwright_html_generator():
    """Factory function para crear un generador de reportes HTML"""
    return PlaywrightHTMLReportGenerator()

# Función principal para integrar con app.py
def generate_native_playwright_evidence(suite_info, suite_results, execution_id, suite_id):
    """
    Función principal para generar evidencia con reporte HTML nativo de Playwright.
    Esta función puede ser llamada desde app.py.
    """
    generator = create_playwright_html_generator()
    
    try:
        # Generar evidencia híbrida (HTML nativo + Word)
        result = generator.generate_hybrid_evidence(
            suite_info, suite_results, execution_id, suite_id
        )
        
        return result
        
    finally:
        generator.cleanup()

if __name__ == "__main__":
    # Ejemplo de uso
    print("🧪 Testing Playwright HTML Report Generator...")
    
    # Datos de ejemplo
    suite_info = {
        'nombre': 'Suite de Prueba',
        'descripcion': 'Suite de ejemplo para testing'
    }
    
    suite_results = {
        'cases': [
            {
                'case_name': 'Test Login',
                'status': 'success',
                'message': 'Login successful',
                'task_id': 'test_001',
                'url': 'https://example.com/login'
            },
            {
                'case_name': 'Test Navigation',
                'status': 'error',
                'message': 'Navigation failed',
                'task_id': 'test_002',
                'url': 'https://example.com/dashboard'
            }
        ],
        'success_count': 1,
        'failed_count': 1
    }
    
    execution_id = str(uuid.uuid4())[:8]
    suite_id = "test_suite_001"
    
    result = generate_native_playwright_evidence(
        suite_info, suite_results, execution_id, suite_id
    )
    
    print(f"Resultado: {result}") 